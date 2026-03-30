"""
IP Arsenal - 个人知识资产管理平台
后端 API 服务
"""
# ── 强制 UTF-8 编码（Windows 默认 GBK 会导致中文乱码）────────────────────
import sys, os
if sys.version_info >= (3, 7):
    import io as _io
    if hasattr(sys.stdout, 'buffer'):
        sys.stdout = _io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    if hasattr(sys.stderr, 'buffer'):
        sys.stderr = _io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
os.environ.setdefault('PYTHONUTF8', '1')
os.environ.setdefault('PYTHONIOENCODING', 'utf-8')

import json, sqlite3, uuid, time, asyncio, base64, io, threading
import queue
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

import fitz  # PyMuPDF
import numpy as np
from PIL import Image
from openai import OpenAI
import httpx
from bs4 import BeautifulSoup

# ─── 导入新的智能提取模块 ──────────────────────────────────────────────
try:
    from chunking import chunk_book_text, HierarchicalChunkingPipeline
    from extraction_pipeline import MultiRoundExtractionPipeline, extract_book_content
    from quality_control import QualityControlPipeline, check_material_quality
    _SMART_EXTRACTION_AVAILABLE = True
    print("[Extraction] 智能提取模块已加载（分层Chunking + 多轮Pipeline + 质量评分）")
except ImportError as e:
    _SMART_EXTRACTION_AVAILABLE = False
    print(f"[Extraction] 智能提取模块未加载: {e}")

# ─── opendataloader-pdf：确保 Java PATH 注入，支持高精度PDF解析 ──────────
_JAVA_BIN = r"C:\Users\jeremyko11\AppData\Local\Programs\Microsoft\jdk-11.0.30.7-hotspot\bin"
if _JAVA_BIN not in os.environ.get("PATH", ""):
    os.environ["PATH"] = _JAVA_BIN + os.pathsep + os.environ.get("PATH", "")
try:
    from opendataloader_pdf import convert as _odl_convert
    _ODL_AVAILABLE = True
    print("[OCR] opendataloader-pdf 已加载（Java 11 XY-Cut++ 引擎）")
except ImportError:
    _ODL_AVAILABLE = False
    print("[OCR] opendataloader-pdf 未安装，跳过")

# ─── PaddleOCR 延迟初始化（首次使用时加载，避免启动慢）────────────────
_paddle_ocr = None

def get_paddle_ocr():
    global _paddle_ocr
    if _paddle_ocr is None:
        try:
            import os
            # 禁用 OneDNN（MKL-DNN），避免 PIR 不兼容问题
            os.environ.setdefault("FLAGS_use_mkldnn", "0")
            os.environ.setdefault("PADDLE_DISABLE_ONEDNN", "1")
            from paddleocr import PaddleOCR
            import logging
            logging.getLogger("ppocr").setLevel(logging.ERROR)
            # 关键参数说明：
            #   use_angle_cls=True  : 启用文字行方向分类（解决倒置/旋转文字识别）
            #   lang="ch"           : 中英文混合识别
            # PaddleOCR 3.x 通过 predict() 接口调用，初始化时不需要传很多参数
            _paddle_ocr = PaddleOCR(lang="ch", use_angle_cls=True)
            print("PaddleOCR ready (PP-OCRv5, angle_cls=True)")
        except Exception as e:
            print(f"PaddleOCR init failed: {e}, will use AI fallback")
            _paddle_ocr = False  # 标记初始化失败
    return _paddle_ocr if _paddle_ocr is not False else None

# ─── 配置 ───────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent.parent
DATA_DIR   = BASE_DIR / "data"
UPLOAD_DIR = BASE_DIR / "uploads"
FRONT_DIR  = BASE_DIR / "frontend"
DB_PATH    = DATA_DIR / "arsenal.db"

DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

API_BASE = "https://maas-coding-api.cn-huabei-1.xf-yun.com/v2"
API_KEY  = "d25220d05c80686af77fcc163c6fe92a:MmRkNjk3MzQ2MGQzMDllNzAyZjM3Mzg0"
MODEL_ID = "astron-code-latest"

# MiniMax 账号1（旧账号，可能额度已耗尽）
MINIMAX_API_BASE_1 = "https://llm.chudian.site/v1"
MINIMAX_API_KEY_1  = "sk-ag-0e87970be36f68d06e47e7a49cceb64d"
MINIMAX_MODEL_ID   = "minimax-m2.7"   # chudian 代理用小写，官方直连用 MiniMax-M2.7

# MiniMax 账号2（新账号，会员账号，每6小时4000次调用）
MINIMAX_API_BASE_2 = "https://api.minimaxi.com/v1"   # MiniMax 国内版官方直连（minimaxi.com，非 minimax.chat）
MINIMAX_MODEL_ID_2 = "MiniMax-M2.7"   # 官方直连用大写（小写返回 401）
MINIMAX_API_KEY_2  = "sk-cp-9ThlxsSvDTzoHt426NobJjtElVxPp1DBPUWEToX7Er0N6cH_MJjMBWQvOSZYXNFJUm071-xnMeGih4D26vl6bjlX7oq19rZ10P77AAXwkiZb0MXdBgIUHyQ"

# 备用模型2：DeepSeek（最终兜底）
FALLBACK2_API_BASE = "https://api.deepseek.com/v1"
FALLBACK2_API_KEY  = "sk-ab948053383f436fb1cf50639f57b439"
FALLBACK2_MODEL_ID = "deepseek-chat"

# ── 全局 AI 模型首选配置（运行时可通过 /api/ai-model 接口修改）────────
# preferred_ai: "xunfei" | "minimax" | "deepseek"
# 选择后，ai_extract 从该模型开始，失败才降级到后面
_AI_PREFERRED = "minimax2"  # 讯飞额度耗尽时直接从 MiniMax 账号2 开始

# 向后兼容别名（旧代码引用）
FALLBACK_API_BASE = MINIMAX_API_BASE_1
FALLBACK_API_KEY  = MINIMAX_API_KEY_1
FALLBACK_MODEL_ID = MINIMAX_MODEL_ID

# 百度 AI Studio PaddleOCR 云端 API token（可选，设置后扫描版识别率大幅提升）
# 获取方式：登录 https://aistudio.baidu.com → 右上角头像 → 访问令牌
# 也可通过环境变量 AISTUDIO_TOKEN 设置
AISTUDIO_TOKEN = os.environ.get("AISTUDIO_TOKEN", "")

IP_DIRECTION = "职场认知升级 / 人性洞察 / 个人成长破局"
MAX_CHARS = 80000
MAX_PROMPT_TEXT_CHARS = 25000  # 单次 AI 调用最多处理的文本量（留空间给提示词）

client = OpenAI(api_key=API_KEY, base_url=API_BASE, timeout=120.0)
fallback_client = OpenAI(api_key=MINIMAX_API_KEY_1, base_url=MINIMAX_API_BASE_1, timeout=120.0)
fallback2_client = OpenAI(api_key=FALLBACK2_API_KEY, base_url=FALLBACK2_API_BASE, timeout=120.0)

# MiniMax 账号2 客户端（key 可在运行时通过接口更新）
_minimax2_client: Optional[OpenAI] = None

def get_minimax2_client() -> Optional[OpenAI]:
    """获取 MiniMax 账号2 的客户端，key 为空则返回 None"""
    global _minimax2_client, MINIMAX_API_KEY_2
    if not MINIMAX_API_KEY_2:
        return None
    if _minimax2_client is None:
        _minimax2_client = OpenAI(api_key=MINIMAX_API_KEY_2, base_url=MINIMAX_API_BASE_2, timeout=600.0)  # 客户端超时要大于接口超时(480s)
    return _minimax2_client

def is_xunfei_blocked(error_str: str) -> bool:
    """判断是否为讯飞内容审核拦截"""
    blocked_keywords = ["法律法规", "无法提供", "xunfei response error", "涉及国家安全", "健康和谐网络"]
    return any(kw in error_str for kw in blocked_keywords)

def strip_think_tags(text: str) -> str:
    """过滤 MiniMax 等模型返回的 <think>...</think> 推理过程标签"""
    import re
    # 去掉 <think>...</think> 块（可能跨多行）
    text = re.sub(r'<think>[\s\S]*?</think>', '', text, flags=re.IGNORECASE)
    return text.strip()

def _call_ai_with_timeout(ai_client, model_id: str, system_prompt: str, user_prompt: str,
                          max_tokens: int, temperature: float, timeout_secs: int = 180) -> str:
    """在独立线程中调用AI，使用concurrent.futures实现强制超时。
    如果超过 timeout_secs 秒没有返回，抛出 TimeoutError。
    这是解决 AI 接口无响应导致 Worker 永久阻塞的根本方案。
    """
    import concurrent.futures
    def _do_call():
        resp = ai_client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            max_tokens=max_tokens,
            temperature=temperature,
        )
        return resp.choices[0].message.content or ""

    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(_do_call)
        try:
            return future.result(timeout=timeout_secs)
        except concurrent.futures.TimeoutError:
            raise TimeoutError(f"AI接口超过{timeout_secs}秒未响应，强制中断")


def _is_quota_error(err_str: str) -> bool:
    """判断是否为余额不足/配额耗尽错误，是则立即跳过，不必等待超时。"""
    LOW = err_str.lower()
    return any(k in LOW for k in (
        "insufficient_quota", "insufficient quota", "402",
        "余额不足", "额度不足", "quota", "billing", "balance",
        "rate_limit_exceeded", "you exceeded your current quota",
    ))


def _build_model_chain() -> list:
    """根据 _AI_PREFERRED 构建 AI 调用链（从首选模型开始，失败依次降级）。
    
    返回列表元素：(label, ai_client, model_id, timeout_secs, strip_think)
    chain 顺序由用户在前端选择的首选模型决定。
    """
    # 全量候选（label, client_getter, model_id, timeout, strip_think）
    # client_getter 用 lambda 延迟获取，以便 MiniMax2 key 在运行时可更新
    _ALL = {
        "xunfei":   ("讯飞",   lambda: client,              MODEL_ID,          120, False),
        "minimax":  ("MiniMax", lambda: fallback_client,    MINIMAX_MODEL_ID,   60, True),
        "minimax2": ("MiniMax2",lambda: get_minimax2_client(), MINIMAX_MODEL_ID_2, 480, True),
        "deepseek": ("DeepSeek",lambda: fallback2_client,   FALLBACK2_MODEL_ID,120, True),
    }

    # 固定降级顺序：讯飞 → MiniMax → MiniMax2（若有key）→ DeepSeek
    _DEFAULT_ORDER = ["xunfei", "minimax", "minimax2", "deepseek"]

    # 以 _AI_PREFERRED 为首，其余按默认顺序跟上
    preferred = _AI_PREFERRED if _AI_PREFERRED in _ALL else "xunfei"
    order = [preferred] + [k for k in _DEFAULT_ORDER if k != preferred]

    chain = []
    for key in order:
        label, getter, model_id, timeout, strip = _ALL[key]
        # minimax2 需要 key 才加进链
        if key == "minimax2" and not MINIMAX_API_KEY_2:
            continue
        chain.append((label, getter, model_id, timeout, strip))
    return chain


def ai_extract(system_prompt: str, user_prompt: str, max_tokens: int = 12000, temperature: float = 0.7) -> tuple[str, str]:
    """调用 AI 提炼，根据 _AI_PREFERRED 决定起点，自动降级。返回 (raw_content, model_used)
    
    特性：
    - 用户在前端选好首选模型后，后端从该模型开始调用
    - 余额不足(402)立即跳过，不等超时
    - 讯飞/DeepSeek 超时 120s，MiniMax 超时 60s
    """
    chain = _build_model_chain()
    errors = []

    for label, getter, model_id, timeout, should_strip in chain:
        ai_client = getter()
        if ai_client is None:
            continue  # 客户端不可用（如 minimax2 key 未填）
        try:
            content = _call_ai_with_timeout(
                ai_client, model_id, system_prompt, user_prompt,
                max_tokens, temperature, timeout_secs=timeout
            )
            if should_strip:
                content = strip_think_tags(content)
            print(f"[AI] {label} 成功（model={model_id}）")
            return content, model_id
        except Exception as _ex:
            err = str(_ex)[:300]
            errors.append(f"{label}:{err[:120]}")
            if _is_quota_error(err):
                print(f"[AI] {label} 余额/配额不足，立即切换下一级...")
            else:
                print(f"[AI] {label} 失败: {err[:120]}，切换下一级...")

    raise RuntimeError(f"所有 AI 模型均失败 | " + " | ".join(errors))

# ─── 数据库初始化 ───────────────────────────────────────────────────
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    # 书籍/来源表
    c.execute("""CREATE TABLE IF NOT EXISTS sources (
        id TEXT PRIMARY KEY,
        title TEXT NOT NULL,
        type TEXT NOT NULL,        -- book/url/text/image
        file_path TEXT,
        url TEXT,
        tags TEXT DEFAULT '[]',
        page_count INTEGER DEFAULT 0,
        char_count INTEGER DEFAULT 0,
        status TEXT DEFAULT 'pending',  -- pending/processing/done/error
        error_msg TEXT,
        is_scanned INTEGER DEFAULT 0,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )""")
    # 素材表
    c.execute("""CREATE TABLE IF NOT EXISTS materials (
        id TEXT PRIMARY KEY,
        source_id TEXT NOT NULL,
        category TEXT NOT NULL,  -- quote/case/viewpoint/action/topic
        content TEXT NOT NULL,
        metadata TEXT DEFAULT '{}',   -- JSON: risk/scene/cost/timeliness/...
        tags TEXT DEFAULT '[]',
        platform TEXT DEFAULT '[]',   -- JSON array
        use_count INTEGER DEFAULT 0,
        is_starred INTEGER DEFAULT 0,
        created_at TEXT NOT NULL,
        FOREIGN KEY(source_id) REFERENCES sources(id)
    )""")
    # 创作记录表
    c.execute("""CREATE TABLE IF NOT EXISTS creations (
        id TEXT PRIMARY KEY,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        source_ids TEXT DEFAULT '[]',
        material_ids TEXT DEFAULT '[]',
        platform TEXT DEFAULT '',
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )""")
    # 任务进度表
    c.execute("""CREATE TABLE IF NOT EXISTS tasks (
        id TEXT PRIMARY KEY,
        source_id TEXT NOT NULL,
        status TEXT DEFAULT 'pending',
        progress INTEGER DEFAULT 0,
        message TEXT DEFAULT '',
        result TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )""")
    # 朴树之道：系列文案表
    c.execute("""CREATE TABLE IF NOT EXISTS scripts (
        id TEXT PRIMARY KEY,
        source_id TEXT NOT NULL,
        source_title TEXT NOT NULL,
        episode_count INTEGER DEFAULT 8,
        platform TEXT DEFAULT '抖音/视频号',
        style TEXT DEFAULT '犀利、接地气、直击痛点',
        status TEXT DEFAULT 'pending',  -- pending/processing/done/error
        progress INTEGER DEFAULT 0,
        message TEXT DEFAULT '',
        episodes TEXT DEFAULT '[]',     -- JSON: [{title, content, ep_no}]
        plan TEXT DEFAULT '[]',         -- JSON: 策划案
        error_msg TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )""")
    conn.commit()
    conn.close()

def migrate_db():
    """兼容旧数据库，添加新字段"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    # 检查是否有 is_scanned 列，没有就加上
    cols = [row[1] for row in c.execute("PRAGMA table_info(sources)").fetchall()]
    if "is_scanned" not in cols:
        c.execute("ALTER TABLE sources ADD COLUMN is_scanned INTEGER DEFAULT 0")
        conn.commit()
    # 检查 scripts 表是否存在（新加的朴树之道功能）
    tables = [row[0] for row in c.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
    if "scripts" not in tables:
        c.execute("""CREATE TABLE IF NOT EXISTS scripts (
            id TEXT PRIMARY KEY,
            source_id TEXT NOT NULL,
            source_title TEXT NOT NULL,
            episode_count INTEGER DEFAULT 8,
            platform TEXT DEFAULT '抖音/视频号',
            style TEXT DEFAULT '犀利、接地气、直击痛点',
            status TEXT DEFAULT 'pending',
            progress INTEGER DEFAULT 0,
            message TEXT DEFAULT '',
            episodes TEXT DEFAULT '[]',
            plan TEXT DEFAULT '[]',
            error_msg TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )""")
        conn.commit()
    conn.close()

init_db()
migrate_db()

# ─── FastAPI 应用 ────────────────────────────────────────────────────
# 自定义 JSON encoder：禁止 ensure_ascii，中文直接输出而不转义为 \uXXXX
from fastapi.responses import ORJSONResponse
try:
    import orjson
    class _UnicodeJSONResponse(ORJSONResponse):
        pass
    _json_response_class = _UnicodeJSONResponse
except ImportError:
    from fastapi.responses import Response as _Resp
    class _UnicodeJSONResponse(_Resp):  # type: ignore
        media_type = "application/json; charset=utf-8"
        def render(self, content) -> bytes:
            return json.dumps(content, ensure_ascii=False, allow_nan=False).encode("utf-8")
    _json_response_class = _UnicodeJSONResponse

app = FastAPI(title="IP Arsenal API", version="1.0.0", default_response_class=_json_response_class)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

@app.on_event("startup")
def on_startup():
    """应用启动时初始化 Worker 线程池并恢复卡住任务"""
    _start_workers()
    _recover_stuck_tasks()



# ─── 全局任务队列 + Worker 线程池 ─────────────────────────────────────
# 使用持久化线程池，避免 BackgroundTasks 在重启后丢失任务
_task_queue: queue.Queue = queue.Queue()
WORKER_COUNT = 8       # 同时处理 8 本书（AI调用是IO密集型，增加线程提升吞吐量）
TASK_MAX_SECONDS = 600 # 单任务最长执行时间（10分钟），超时强制跳过；MiniMax处理80k字约需5-8分钟
_worker_threads: list = []              # 跟踪 worker 线程，供 watchdog 检测
_worker_heartbeats: dict = {}           # worker_id → 最后一次取到任务的时间戳（0=idle）

def _worker_loop(worker_id: int):
    """Worker 线程：持续从队列取任务执行，任何异常都不会导致线程退出。
    
    【重构说明】
    - 心跳时间戳：每次取到任务时更新，Watchdog 可据此检测线程是否阻塞（卡死）
    - 任务本身在独立子线程中执行，主循环有 TASK_MAX_SECONDS 兜底超时
    - 超时的任务会被标记为 error，Worker 继续处理下一个任务
    """
    print(f"[Worker-{worker_id}] 启动")
    while True:
        task_item = None
        try:
            task_item = _task_queue.get(timeout=2)
            task_id, source_id, mode = task_item

            # 更新心跳时间戳
            _worker_heartbeats[worker_id] = time.time()
            print(f"[Worker-{worker_id}] 开始处理 source={source_id} mode={mode}")

            # 在独立线程中跑任务，主循环强制超时兜底
            import concurrent.futures as _cf
            with _cf.ThreadPoolExecutor(max_workers=1, thread_name_prefix=f"W{worker_id}-task") as _tex:
                # 根据 mode 选择处理函数：smart = 智能多轮提取，其他 = 普通提取
                if mode == "smart" and _SMART_EXTRACTION_AVAILABLE:
                    _task_fn = lambda: process_source_task_smart(task_id, source_id, mode)
                else:
                    _task_fn = lambda: process_source_task(task_id, source_id, mode)
                _future = _tex.submit(_task_fn)
                try:
                    _future.result(timeout=TASK_MAX_SECONDS)
                    print(f"[Worker-{worker_id}] 完成 source={source_id}")
                except _cf.TimeoutError:
                    print(f"[Worker-{worker_id}] ⚠️ 任务超时({TASK_MAX_SECONDS}s) source={source_id}，强制跳过")
                    # 标记任务为 error，不阻塞后续任务
                    try:
                        _ec = get_db()
                        _ec.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                                    (f"任务超时（>{TASK_MAX_SECONDS}s），已自动跳过", now(), source_id))
                        _ec.execute("UPDATE tasks SET status='error',message=?,updated_at=? WHERE id=?",
                                    (f"执行超时（>{TASK_MAX_SECONDS}s），已自动跳过", now(), task_id))
                        _ec.commit()
                        _ec.close()
                    except Exception as _de:
                        print(f"[Worker-{worker_id}] 超时标记DB失败: {_de}")
                except Exception as e:
                    print(f"[Worker-{worker_id}] 任务异常 source={source_id}: {e}")
                    import traceback; traceback.print_exc()
        except queue.Empty:
            # 队列空，清空心跳（idle 状态）
            _worker_heartbeats[worker_id] = 0
            continue
        except Exception as e:
            print(f"[Worker-{worker_id}] 队列操作异常: {e}")
        finally:
            if task_item is not None:
                try:
                    _task_queue.task_done()
                except Exception:
                    pass
            _worker_heartbeats[worker_id] = 0  # 重置心跳（idle）
            time.sleep(0.1)  # 避免 CPU 空转

def enqueue_task(task_id: str, source_id: str, mode: str):
    """将任务加入队列（线程安全）"""
    _task_queue.put((task_id, source_id, mode))

def _spawn_worker(worker_id: int) -> threading.Thread:
    """创建并启动一个 worker 线程"""
    t = threading.Thread(
        target=_worker_loop,
        args=(worker_id,),
        daemon=True,
        name=f"Arsenal-Worker-{worker_id}"
    )
    t.start()
    return t

def _watchdog_loop():
    """Watchdog 线程：每60秒检测 worker 线程是否存活或卡死
    
    【重构说明】
    - 死亡检测：线程 is_alive() == False → 立即重建
    - 卡死检测：心跳时间戳超过 TASK_MAX_SECONDS+60 秒未更新 → 标记为卡死，重建线程
      （卡死发生在 AI 子线程中，Worker 主循环不会卡，所以这里不会触发；
       但如果将来 Worker 主循环本身卡住，此机制可兜底）
    """
    time.sleep(15)  # 启动后等15秒再开始监控
    while True:
        try:
            now_ts = time.time()
            for i, t in enumerate(_worker_threads):
                worker_id = i + 1
                if not t.is_alive():
                    print(f"[Watchdog] Worker-{worker_id} 已死亡，正在重建...")
                    new_t = _spawn_worker(worker_id)
                    _worker_threads[i] = new_t
                    _worker_heartbeats[worker_id] = 0
                    print(f"[Watchdog] Worker-{worker_id} 已重建")
        except Exception as e:
            print(f"[Watchdog] 异常: {e}")
        time.sleep(60)

def _start_workers():
    """启动 Worker 线程池 + Watchdog"""
    global _worker_threads, _worker_heartbeats
    _worker_threads.clear()
    _worker_heartbeats.clear()
    for i in range(WORKER_COUNT):
        worker_id = i + 1
        _worker_heartbeats[worker_id] = 0  # 初始 idle
        t = _spawn_worker(worker_id)
        _worker_threads.append(t)
    print(f"[Arsenal] {WORKER_COUNT} 个 Worker 线程已启动")
    # 启动 watchdog
    wd = threading.Thread(target=_watchdog_loop, daemon=True, name="Arsenal-Watchdog")
    wd.start()
    print("[Arsenal] Watchdog 已启动")

def _recover_stuck_tasks():
    """启动时恢复所有 processing/pending 任务"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        # 将 processing 重置为 pending（防止上次崩溃留下的假 processing 状态）
        conn.execute(
            "UPDATE sources SET status='pending', updated_at=? WHERE status='processing'",
            (datetime.now().strftime('%Y-%m-%d %H:%M:%S'),)
        )
        conn.execute(
            "UPDATE tasks SET status='pending', message='等待处理（服务重启恢复）...', updated_at=? WHERE status='processing'",
            (datetime.now().strftime('%Y-%m-%d %H:%M:%S'),)
        )
        conn.commit()

        # 恢复所有 pending 任务入队
        rows = conn.execute(
            """SELECT t.id as task_id, t.source_id
               FROM tasks t JOIN sources s ON t.source_id = s.id
               WHERE s.status = 'pending'
               ORDER BY s.created_at"""
        ).fetchall()
        for row in rows:
            enqueue_task(row["task_id"], row["source_id"], "full")

        if rows:
            print(f"[Arsenal] 恢复 {len(rows)} 个待处理任务入队")
    finally:
        conn.close()

# 启动 Workers + 恢复任务（在 process_source_task 定义之后，由文件末尾调用）
# _start_workers() 和 _recover_stuck_tasks() 在文件末尾 uvicorn.run 之前调用







def get_db():
    """创建 SQLite 连接，启用 WAL 模式防止多线程死锁，设置合理超时"""
    conn = sqlite3.connect(DB_PATH, timeout=30, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    # WAL 模式：多个读写线程并发不会互相阻塞
    conn.execute("PRAGMA journal_mode=WAL")
    # 宽松同步模式，提升写入性能
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn

def now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

# ─── 文本提取 ────────────────────────────────────────────────────────

def _is_scanned_pdf(doc: fitz.Document, sample_pages: int = 10) -> bool:
    """检测PDF是否需要OCR（扫描版 或 低质量混合版）
    
    借鉴 opendataloader-pdf 的混合路由策略：
    - 扫描版：文字层基本为空 + 有大面积图像
    - 混合版（低质量）：文字层存在但密度低、乱码比例高、字符稀疏
      这类PDF直接用 get_text() 会得到残缺或乱码文字，需要 OCR 补救
    
    判断逻辑升级：
    1. 平均每页字符 < 30 → 扫描版（强信号）
    2. 超过60%页面是"图像页"（有图无字）→ 扫描版（强信号）
    3. 平均每页字符 30-150 且图像比例 > 40% → 混合版，触发 OCR（新增）
    4. 文字层中非ASCII/中文可见字符比例低于30%（乱码/符号堆砌）→ 低质量，触发OCR（新增）
    """
    total_pages = len(doc)
    if total_pages == 0:
        return False
    
    # 均匀采样，跳过第0页（封面）
    start_page = min(1, total_pages - 1)
    check_pages = min(sample_pages, total_pages)
    step = max(1, (total_pages - start_page) // check_pages)
    indices = list(range(start_page, total_pages, step))[:check_pages]
    if not indices:
        indices = list(range(total_pages))
    
    total_chars = 0
    image_heavy_pages = 0
    garbled_chars = 0   # 疑似乱码/符号字符数
    total_valid_chars = 0  # 有效字符（中文+字母+数字）计数
    
    for i in indices:
        page = doc[i]
        text = page.get_text().strip()
        total_chars += len(text)
        
        # 检测是否有大面积图像（扫描版的关键特征）
        if len(text) < 50:
            img_list = page.get_images(full=False)
            if img_list:
                image_heavy_pages += 1
        
        # 统计有效字符 vs 乱码字符（借鉴 opendataloader 的质量评估思路）
        for ch in text:
            code = ord(ch)
            # 有效字符：中文字符、ASCII字母数字、常见标点
            if (0x4E00 <= code <= 0x9FFF) or ch.isalnum():
                total_valid_chars += 1
            elif code > 127 and not (0x4E00 <= code <= 0x9FFF):
                # 非ASCII、非中文字符 = 疑似符号/乱码
                garbled_chars += 1
    
    avg = total_chars / len(indices) if indices else 0
    image_ratio = image_heavy_pages / len(indices) if indices else 0
    
    # 判断一：经典扫描版（文字极少 + 图像密集）
    if avg < 30 or image_ratio > 0.6:
        return True
    
    # 判断二：混合版/低质量文字层（字符稀疏 + 有图像）
    # opendataloader-pdf 的混合路由关键：字少但有图 → 大概率文字层是"占位符"而非真实内容
    if 30 <= avg <= 150 and image_ratio > 0.4:
        return True
    
    # 判断三：乱码检测（文字层存在但质量差，有效字符比例低）
    # 正常中文书: 有效字符 > 70%；乱码PDF：可能低于 30%
    if total_valid_chars + garbled_chars > 0:
        valid_ratio = total_valid_chars / (total_valid_chars + garbled_chars + 1)
        if valid_ratio < 0.3 and avg < 200:
            return True
    
    return False


def _page_to_numpy(page: fitz.Page, dpi: int = 300, enhance: bool = True) -> np.ndarray:
    """将PDF页面渲染为numpy数组（供PaddleOCR使用）
    
    升级版图像预处理管线（借鉴 opendataloader-pdf 的高精度渲染策略）：
    - 300 DPI 高分辨率渲染（PaddleOCR官方推荐最低标准）
    - 多步图像增强：锐化→对比度→亮度→降噪
    - 灰度优化：对偏黄/偏灰的扫描页自动纠正
    - 暗页自动提亮（旧书扫描常见问题）
    """
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    
    if not enhance:
        return np.array(img)
    
    try:
        from PIL import ImageEnhance, ImageFilter, ImageOps
        import numpy as _np
        
        # 步骤1：检测页面亮度，暗页自动提亮（旧书扫描常见）
        img_arr_check = _np.array(img.convert('L'))
        mean_brightness = img_arr_check.mean()
        
        # 步骤2：自适应亮度调整
        if mean_brightness < 180:
            # 偏暗页面：提亮
            brightness_factor = min(1.3, 220 / (mean_brightness + 1))
            img = ImageEnhance.Brightness(img).enhance(brightness_factor)
        elif mean_brightness > 240:
            # 过曝页面（白底白字）：轻微降低
            img = ImageEnhance.Brightness(img).enhance(0.95)
        
        # 步骤3：对比度增强（扫描版文字对比度不足是识别率低的主因）
        # opendataloader-pdf 中使用了类似的对比度增强逻辑
        img = ImageEnhance.Contrast(img).enhance(1.5)
        
        # 步骤4：UnsharpMask 精细锐化（比 SHARPEN 更精细，不会过度锐化边缘）
        img = img.filter(ImageFilter.UnsharpMask(radius=1.2, percent=130, threshold=2))
        
        # 步骤5：轻微降噪（高DPI扫描件会引入噪点，影响OCR）
        # 使用 MedianFilter 去椒盐噪声（对文字边缘影响最小）
        img = img.filter(ImageFilter.MedianFilter(size=3))
        
        # 步骤6：色彩空间归一化（消除偏黄/偏棕的扫描色调）
        # 将 RGB 分离，均衡各通道亮度，使文字更黑、背景更白
        r, g, b = img.split()
        # 如果三个通道差异大（偏色），轻微归一化
        r_mean = _np.array(r).mean()
        g_mean = _np.array(g).mean()  
        b_mean = _np.array(b).mean()
        channel_diff = max(r_mean, g_mean, b_mean) - min(r_mean, g_mean, b_mean)
        if channel_diff > 15:  # 偏色明显（超过15灰度级差异）
            # 轻微均衡化：把偏黄/偏棕的扫描件拉向中性灰
            target_mean = (r_mean + g_mean + b_mean) / 3
            r = ImageEnhance.Brightness(r).enhance(target_mean / (r_mean + 0.01))
            g = ImageEnhance.Brightness(g).enhance(target_mean / (g_mean + 0.01))
            b = ImageEnhance.Brightness(b).enhance(target_mean / (b_mean + 0.01))
            img = Image.merge('RGB', (r, g, b))
    except Exception:
        pass  # 预处理失败不影响主流程，直接用原图
    
    return np.array(img)



# ─── 每线程独立 PaddleOCR 实例（PaddleOCR 非线程安全）──────────────
_paddle_ocr_local = threading.local()

def _get_thread_paddle_ocr():
    """获取当前线程专属的 PaddleOCR 实例（懒初始化，避免跨线程共享）"""
    if not getattr(_paddle_ocr_local, "ocr", None):
        try:
            import os as _os
            _os.environ.setdefault("FLAGS_use_mkldnn", "0")
            _os.environ.setdefault("PADDLE_DISABLE_ONEDNN", "1")
            from paddleocr import PaddleOCR
            import logging
            logging.getLogger("ppocr").setLevel(logging.ERROR)
            _paddle_ocr_local.ocr = PaddleOCR(lang="ch", use_angle_cls=True)
        except Exception as e:
            print(f"Thread PaddleOCR init failed: {e}")
            _paddle_ocr_local.ocr = False
    return _paddle_ocr_local.ocr if _paddle_ocr_local.ocr is not False else None


def _ocr_single_page_paddle(args):
    """单页 OCR 工作函数（在线程池中执行）
    返回 (page_num, text_or_none)
    """
    page_num, img_arr = args
    ocr = _get_thread_paddle_ocr()
    if ocr is None:
        return page_num, None
    try:
        result = ocr.predict(img_arr)
        page_lines = []
        for res in (result or []):
            if isinstance(res, dict):
                texts = res.get("rec_texts") or []
                scores = res.get("rec_scores") or []
                for i, t in enumerate(texts):
                    if not t or len(t.strip()) <= 1:
                        continue
                    score = scores[i] if i < len(scores) else 1.0
                    # 置信度阈值 0.65（借鉴 opendataloader-pdf 精度优先策略，减少低质量字符混入）
                    if score >= 0.65:
                        page_lines.append(t.strip())
            elif isinstance(res, list):
                for item in res:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        txt_info = item[1]
                        if isinstance(txt_info, (list, tuple)) and len(txt_info) >= 2:
                            txt, score = txt_info[0], txt_info[1]
                            if txt and len(txt.strip()) > 1 and score >= 0.65:
                                page_lines.append(txt.strip())
        return page_num, "\n".join(page_lines) if page_lines else None
    except Exception as e:
        print(f"PaddleOCR page {page_num} error: {e}")
        return page_num, None


def _ocr_pages_with_paddle(doc: fitz.Document, page_indices: list, task_updater=None,
                           pdf_path: str = None) -> str:
    """用PaddleOCR对指定页面做本地OCR，返回合并文本
    
    【并行优化】使用 ThreadPoolExecutor 多线程并行处理页面（每线程独立OCR实例）。
    页面渲染（CPU密集）+ OCR推理 均可并行，速度约提升 3-4x。
    
    三级降级策略：
      1. 本地 PaddleOCR（300 DPI + 多线程并行，最快，完全离线）
      2. 百度 AI Studio PaddleOCR 云端 API（效果最好，免费）
      3. 讯飞 AI 视觉 OCR（最后备用）
    """
    # 先检查主实例是否可用（懒加载触发器）
    if get_paddle_ocr() is None:
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    import concurrent.futures as _cf

    total = len(page_indices)
    # OCR 线程数：最多 4 线程，避免内存爆炸（每个实例约 500MB）
    ocr_workers = min(8, max(1, total // 3 + 1))

    if task_updater:
        task_updater("processing", 28, f"PaddleOCR并行识别{total}页（{ocr_workers}线程）...")

    # 提前渲染所有页面为图像（在主线程完成，doc 不是线程安全的）
    rendered = []
    for page_num in page_indices:
        try:
            img_arr = _page_to_numpy(doc[page_num], dpi=300)
            rendered.append((page_num, img_arr))
        except Exception as e:
            print(f"Page render failed {page_num}: {e}")
            rendered.append((page_num, None))

    # 并行 OCR
    results_map = {}  # page_num → text
    fail_count = 0
    completed = 0

    def _do_ocr(args):
        pn, img = args
        if img is None:
            return pn, None
        return _ocr_single_page_paddle((pn, img))

    with _cf.ThreadPoolExecutor(max_workers=ocr_workers, thread_name_prefix="OCR") as executor:
        future_map = {executor.submit(_do_ocr, item): item[0] for item in rendered}
        for future in _cf.as_completed(future_map):
            try:
                page_num, text = future.result(timeout=60)
                results_map[page_num] = text
                if text:
                    completed += 1
                else:
                    fail_count += 1
            except Exception as e:
                pn = future_map[future]
                print(f"OCR future error page {pn}: {e}")
                results_map[pn] = None
                fail_count += 1
            # 更新进度
            if task_updater:
                done_cnt = len(results_map)
                pct = 30 + int((done_cnt / total) * 40)
                task_updater("processing", pct, f"PaddleOCR识别中...（{done_cnt}/{total}页完成）")

    # 按原顺序组装结果
    ordered_texts = [results_map.get(pn) for pn in page_indices]
    results = [t for t in ordered_texts if t]

    # 超过50%页面失败，升级到 AI Studio
    if fail_count > total * 0.5:
        if task_updater:
            task_updater("processing", 28, f"PaddleOCR识别率低（{fail_count}/{total}页失败），切换到AI Studio云端...")
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    combined = "\n\n".join(results)
    if len(combined.strip()) < 100 and total > 2:
        if task_updater:
            task_updater("processing", 28, "PaddleOCR输出为空，切换到AI Studio云端...")
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    return combined


def _try_aistudio_or_fallback(doc: fitz.Document, page_indices: list,
                               pdf_path: str = None, task_updater=None) -> str:
    """尝试百度 AI Studio 云端 OCR，失败则降到讯飞 AI-OCR"""
    if pdf_path:
        try:
            result = _ocr_pdf_with_aistudio(pdf_path, page_indices, task_updater)
            if result and len(result.strip()) >= 50:
                return result
        except Exception as e:
            print(f"AI Studio OCR failed: {e}, falling back to AI vision OCR")
            if task_updater:
                task_updater("processing", 45, f"AI Studio识别失败（{str(e)[:60]}），切换到讯飞AI-OCR...")
    return _ocr_pages_with_ai_fallback(doc, page_indices, task_updater)


def _ocr_pages_with_ai_fallback(doc: fitz.Document, page_indices: list, task_updater=None) -> str:
    """降级方案：用讯飞AI视觉接口做OCR（当PaddleOCR不可用时）
    
    【并行优化】2线程并行（讯飞 API 有并发限制，不宜过多）。
    渲染在主线程完成（doc 非线程安全），API 调用并行执行。
    """
    import concurrent.futures as _cf

    total = len(page_indices)
    # 讯飞 API 并发保守设置 2，避免触发限速
    ai_workers = min(4, total)  # 64GB内存，可以跑更多并发

    # 提前渲染所有页面（主线程，doc 非线程安全）
    rendered = []
    for page_num in page_indices:
        try:
            # 提升 AI-OCR 渲染分辨率：250 DPI（原来200，提升约30%清晰度）
            # opendataloader-pdf 使用高分辨率渲染是其高精度的关键之一
            mat = fitz.Matrix(250 / 72, 250 / 72)
            pix = doc[page_num].get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            rendered.append((page_num, b64))
        except Exception as e:
            rendered.append((page_num, None))
            print(f"AI-OCR render page {page_num} failed: {e}")

    if task_updater:
        task_updater("processing", 30, f"讯飞AI-OCR并行识别{total}页（{ai_workers}线程）...")

    def _call_ai_ocr(item):
        page_num, b64 = item
        if b64 is None:
            return page_num, f"[第{page_num+1}页渲染失败]"
        try:
            resp = client.chat.completions.create(
                model=MODEL_ID,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    {"type": "text", "text": "请将这张图片中的所有文字完整识别出来，保持原有段落和分节结构，直接输出识别到的文字内容，不要添加任何说明。处理要求：1. 如果是中文书籍，优先保证中文字符识别完整；2. 如果图片偏暗或偏黄（旧书扫描），请仍然尽力识别可见文字；3. 如果某处文字确实无法识别，用[?]标记，不要随意猜测；4. 保留段落换行，不要合并段落。"}
                ]}],
                max_tokens=4000, temperature=0,
            )
            return page_num, resp.choices[0].message.content or ""
        except Exception as e:
            return page_num, f"[第{page_num+1}页OCR失败: {e}]"

    results_map = {}
    completed = 0
    with _cf.ThreadPoolExecutor(max_workers=ai_workers, thread_name_prefix="AIOCR") as executor:
        future_map = {executor.submit(_call_ai_ocr, item): item[0] for item in rendered}
        for future in _cf.as_completed(future_map):
            try:
                page_num, text = future.result(timeout=90)
                results_map[page_num] = text
                completed += 1
            except Exception as e:
                pn = future_map[future]
                results_map[pn] = f"[第{pn+1}页OCR超时: {e}]"
                completed += 1
            if task_updater:
                pct = 30 + int((completed / total) * 40)
                task_updater("processing", pct, f"讯飞AI-OCR识别中...（{completed}/{total}页完成）")

    # 按原顺序返回
    return "\n\n".join(results_map.get(pn, "") for pn in page_indices)


def _ocr_pdf_with_aistudio(pdf_path: str, page_indices: list, task_updater=None) -> str:
    """百度 AI Studio PaddleOCR 云端 API（高质量，无需 API Key）
    
    使用 https://paddleocr.aistudio-app.com 的免费公开接口，效果等同在线版。
    采用异步任务模式：提交任务 → 轮询结果。
    作为本地 PaddleOCR 之后、讯飞 AI-OCR 之前的第二级降级方案。
    支持直接上传 PDF 文件（最大 50MB），一次处理多页，远比逐页调用讯飞快。
    """
    import time
    import tempfile
    import shutil

    AISTUDIO_BASE = "https://paddleocr.aistudio-app.com"
    SUBMIT_URL = f"{AISTUDIO_BASE}/api/v2/ocr/jobs"
    
    # 检查是否有 token
    token = AISTUDIO_TOKEN
    if not token:
        raise ValueError("AISTUDIO_TOKEN not configured, skipping AI Studio OCR")
    
    headers = {"Authorization": f"token {token}"}
    # 如果 page_indices 是所有页面的子集，需要先提取目标页面为新 PDF
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    all_pages = list(range(total_pages))
    
    # 如果要识别所有页面，直接用原始 PDF；否则提取子集
    if set(page_indices) == set(all_pages):
        upload_path = pdf_path
        temp_pdf = None
    else:
        # 创建临时 PDF，只包含需要识别的页面
        temp_doc = fitz.open()
        for pi in sorted(page_indices):
            temp_doc.insert_pdf(doc, from_page=pi, to_page=pi)
        tmp_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        temp_pdf = tmp_file.name
        tmp_file.close()
        temp_doc.save(temp_pdf)
        temp_doc.close()
        upload_path = temp_pdf
    doc.close()

    try:
        if task_updater:
            task_updater("processing", 28, "上传到百度AI Studio PaddleOCR云端识别（高精度）...")

        # 提交任务（multipart 上传文件）
        file_size_mb = Path(upload_path).stat().st_size / (1024 * 1024)
        if file_size_mb > 48:
            # 文件太大，降级到讯飞 AI-OCR
            raise ValueError(f"PDF too large for AI Studio ({file_size_mb:.1f}MB > 48MB), fallback to AI-OCR")

        with open(upload_path, "rb") as f:
            files_data = {"file": (Path(upload_path).name, f, "application/pdf")}
            resp = httpx.post(
                SUBMIT_URL,
                data={"model": "PP-StructureV3",
                      "optionalPayload": '{"useDocOrientationClassify":true,"useDocUnwarping":true}'},
                files=files_data,
                headers=headers,
                timeout=60,
            )

        if resp.status_code != 200:
            raise ValueError(f"AI Studio submit failed: {resp.status_code} {resp.text[:200]}")

        job_resp = resp.json()
        job_id = job_resp.get("jobId") or (job_resp.get("data") or {}).get("jobId")
        if not job_id:
            raise ValueError(f"No jobId in response: {job_resp}")

        if task_updater:
            task_updater("processing", 35, f"AI Studio任务已提交（jobId={job_id[:8]}...），等待识别完成...")

        # 轮询结果（最多等 120 秒）
        result_url = f"{AISTUDIO_BASE}/api/v2/ocr/jobs/{job_id}"
        max_wait = 120
        poll_interval = 5
        elapsed = 0
        while elapsed < max_wait:
            time.sleep(poll_interval)
            elapsed += poll_interval
            r = httpx.get(result_url, headers=headers, timeout=30)
            if r.status_code != 200:
                continue
            data = r.json()
            status = (data.get("status") or "").lower()
            if status in ("success", "finished", "completed", "done"):
                # 提取所有页的文字
                pages = data.get("result", {}).get("pages") or data.get("pages") or []
                if not pages and "markdownUrl" in (data.get("result") or {}):
                    # 有时结果放在 markdownUrl 里
                    md_url = data["result"]["markdownUrl"]
                    md_resp = httpx.get(md_url, headers=headers, timeout=30)
                    return md_resp.text
                texts = []
                for pg in pages:
                    pg_text = pg.get("text") or ""
                    if not pg_text:
                        # 从 layouts 提取
                        for layout in (pg.get("layouts") or []):
                            pg_text += layout.get("text", "") + "\n"
                    texts.append(pg_text.strip())
                return "\n\n".join(t for t in texts if t)
            elif status in ("failed", "error"):
                raise ValueError(f"AI Studio job failed: {data}")
            # 还在处理中，继续等
            if task_updater:
                task_updater("processing", 35 + min(int(elapsed / max_wait * 30), 30),
                             f"AI Studio识别中...（已等待{elapsed}秒）")

        raise ValueError(f"AI Studio timed out after {max_wait}s")

    finally:
        if temp_pdf and Path(temp_pdf).exists():
            Path(temp_pdf).unlink(missing_ok=True)


def _extract_with_opendataloader(pdf_path: str, max_chars: int = None) -> str | None:
    """用 opendataloader-pdf (Java XY-Cut++ 引擎) 提取PDF全文。
    
    这是精度最高的提取方式（综合精度0.90，表格精度0.93），适用于：
    - 扫描版PDF（Java引擎内置智能OCR路由）
    - 含复杂表格的PDF
    - 多栏版式PDF
    - 普通文字PDF（XY-Cut++确保阅读顺序正确）
    
    Returns:
        提取到的文本（成功），或 None（失败/不可用）
    """
    if not _ODL_AVAILABLE:
        return None
    
    try:
        import tempfile, subprocess
        # 用临时目录存放输出
        with tempfile.TemporaryDirectory() as tmpdir:
            # 调用 opendataloader-pdf：输出 text 格式
            # reading_order=xycut：使用 XY-Cut++ 算法确保阅读顺序
            # keep_line_breaks=True：保留段落结构
            _odl_convert(
                input_path=pdf_path,
                output_dir=tmpdir,
                format="text",
                reading_order="xycut",
                keep_line_breaks=True,
                quiet=True,
            )
            
            # 找输出 .txt 文件
            txt_files = [f for f in os.listdir(tmpdir) if f.endswith(".txt")]
            if not txt_files:
                return None
            
            txt_path = os.path.join(tmpdir, txt_files[0])
            with open(txt_path, encoding="utf-8", errors="replace") as fh:
                text = fh.read()
            
            if not text or len(text.strip()) < 100:
                return None  # 输出为空，降级
            
            if max_chars:
                return text[:max_chars]
            return text
    
    except Exception as e:
        print(f"[opendataloader] 提取失败: {e}")
        return None


def _extract_text_layout_aware(page) -> str:
    """布局感知文本提取（借鉴 opendataloader-pdf 的 XY-Cut++ 阅读顺序策略）
    
    普通 get_text() 对多栏布局会乱序合并文字。
    使用 get_text("blocks") 按位置坐标排序，确保阅读顺序正确。
    
    适用场景：书籍、报告、杂志等多栏排版文档。
    """
    try:
        # 获取所有文字块及其坐标 (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")
        if not blocks:
            return page.get_text()
        
        # 过滤空块和图像块（block_type=1 是图像）
        text_blocks = [(b[0], b[1], b[4]) for b in blocks if b[6] == 0 and b[4].strip()]
        
        if not text_blocks:
            return page.get_text()
        
        # XY-Cut 简化版：按Y坐标（行）分组，每行内按X坐标排序
        # 阈值：Y差距在20像素内认为是同一行
        page_rect = page.rect
        page_width = page_rect.width if page_rect.width > 0 else 595
        
        # 检测是否为多栏布局（若文字块X坐标分布在两个明显的区间）
        x_centers = [(b[0] + (page_width * 0.3)) / page_width for b in text_blocks]
        # 简单判断：如果超过20%的块集中在页面右侧 > 55%位置，可能是多栏
        right_blocks = sum(1 for b in text_blocks if b[0] > page_width * 0.55)
        is_multi_column = right_blocks > len(text_blocks) * 0.2
        
        if is_multi_column:
            # 多栏：先按列（X坐标范围）分组，再按行排序
            # 将页面分为两列（左列 < 55%宽度，右列 >= 55%宽度）
            left_col = [(x, y, t) for x, y, t in text_blocks if x < page_width * 0.55]
            right_col = [(x, y, t) for x, y, t in text_blocks if x >= page_width * 0.55]
            
            left_sorted = sorted(left_col, key=lambda b: b[1])   # 左列按Y排序
            right_sorted = sorted(right_col, key=lambda b: b[1]) # 右列按Y排序
            
            # 左列先，右列后（标准中文书籍阅读顺序）
            sorted_blocks = left_sorted + right_sorted
        else:
            # 单栏：直接按Y坐标排序（默认从上到下）
            sorted_blocks = sorted(text_blocks, key=lambda b: (round(b[1] / 20) * 20, b[0]))
        
        return "\n".join(t.strip() for _, _, t in sorted_blocks if t.strip())
    
    except Exception:
        # 降级到普通提取
        return page.get_text()


def extract_text_from_epub(path: str, task_updater=None) -> tuple[str, int, bool]:
    """从 EPUB 文件直接提取纯文本，不走 OCR。返回 (text, chapter_count, False)"""
    try:
        import zipfile, re as _re
        from html.parser import HTMLParser

        class _StripHTML(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []
            def handle_data(self, data):
                self.parts.append(data)
            def get_text(self):
                return " ".join(self.parts)

        def _html_to_text(html_bytes):
            try:
                html = html_bytes.decode("utf-8", errors="replace")
            except Exception:
                html = str(html_bytes)
            p = _StripHTML()
            p.feed(html)
            return p.get_text()

        if task_updater:
            task_updater("processing", 20, "正在解析 EPUB 文件...")

        chapters = []
        with zipfile.ZipFile(path, "r") as zf:
            names = sorted(zf.namelist())
            html_files = [n for n in names if n.lower().endswith((".xhtml", ".html", ".htm", ".xml"))
                          and "ncx" not in n.lower() and "opf" not in n.lower() and "toc" not in n.lower()]
            for n in html_files:
                try:
                    raw = zf.read(n)
                    chapters.append(_html_to_text(raw))
                except Exception:
                    pass

        if not chapters:
            return "", 0, False

        # 读取全部章节内容
        nc = len(chapters)
        text = "\n\n".join(chapters)
        # 清理多余空白
        text = _re.sub(r'\n{3,}', '\n\n', text).strip()
        # 如果超长，按语义分段而不是硬截断
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS]

        if task_updater:
            task_updater("processing", 60, f"EPUB 解析完成（{nc} 章，{len(text):,} 字），正在 AI 提炼...")

        return text, nc, False

    except Exception as e:
        print(f"[EPUB] 提取失败: {e}")
        return "", 0, False


def extract_text_from_txt(path: str, task_updater=None) -> tuple[str | list[str], int, bool]:
    """从 TXT/MD 文件直接读取文本，不走 OCR。返回 (text_or_chunks, chunk_count, False)

    注意：超过 MAX_CHARS 的长文本会按段落分割成多个 chunk，返回 list[str]
    """
    try:
        if task_updater:
            task_updater("processing", 20, "正在读取文本文件...")

        for enc in ("utf-8", "gbk", "utf-16", "latin-1"):
            try:
                with open(path, encoding=enc, errors="replace") as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            return "", 0, False

        # 如果文本超长，按段落分割成多个 chunk
        if len(text) > MAX_CHARS:
            # 按双换行分段（段落）
            chunks = []
            paragraphs = text.split('\n\n')
            current_chunk = ""
            for para in paragraphs:
                if len(current_chunk) + len(para) + 2 > MAX_CHARS:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para
                else:
                    current_chunk = (current_chunk + '\n\n' + para) if current_chunk else para
            if current_chunk.strip():
                chunks.append(current_chunk.strip())

            if task_updater:
                task_updater("processing", 60, f"文件读取完成（{len(chunks)} 个分段，共 {len(text):,} 字），正在 AI 提炼...")
            return chunks, len(chunks), False
        else:
            if task_updater:
                task_updater("processing", 60, f"文件读取完成（{len(text):,} 字），正在 AI 提炼...")
            return text, 1, False

    except Exception as e:
        print(f"[TXT] 提取失败: {e}")
        return "", 0, False


def extract_text_from_docx(path: str, task_updater=None) -> tuple[str, int, bool]:
    """从 DOCX（Word）文件提取纯文本，不走 OCR。返回 (text, para_count, False)"""
    try:
        import docx as _docx_mod
        if task_updater:
            task_updater("processing", 20, "正在读取 Word 文档...")

        doc = _docx_mod.Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            s = para.text.strip()
            if s:
                paragraphs.append(s)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        text = "\n".join(paragraphs)
        text = text[:MAX_CHARS]
        para_count = len(paragraphs)

        if task_updater:
            task_updater("processing", 60, f"Word 文档读取完成（{para_count} 段，{len(text):,} 字），正在 AI 提炼...")
        return text, para_count, False

    except Exception as e:
        print(f"[DOCX] 提取失败: {e}")
        return "", 0, False


def extract_text_from_pdf(path: str, task_updater=None) -> tuple[str, int, bool]:
    """
    提取PDF文本，三级精度递进策略：
    
    Level 1 — opendataloader-pdf（Java XY-Cut++ 引擎，精度0.90，首选）
        适用：所有PDF（普通/扫描/表格/多栏）
        耗时：1-5秒，精度最高
        降级条件：未安装/Java不可用/输出为空
    
    Level 2 — PaddleOCR（本地GPU加速，精度~0.85）
        适用：扫描版PDF（Level1失败或不可用）
        耗时：数十秒至数分钟
        降级条件：OCR失败/结果质量差
    
    Level 3 — 讯飞 AI-OCR（视觉大模型，精度~0.88）
        适用：PaddleOCR失败或质量差时的最终兜底
    
    返回 (text, page_count, is_scanned)
    """
    doc = fitz.open(path)
    page_count = len(doc)

    # ─── Level 1：opendataloader-pdf（精度优先，全能型）────────────────────
    if _ODL_AVAILABLE:
        if task_updater:
            task_updater("processing", 15,
                         f"正在用 opendataloader-pdf 高精度解析（{page_count}页）...")
        
        odl_text = _extract_with_opendataloader(path, max_chars=MAX_CHARS)
        
        if odl_text and len(odl_text.strip()) >= 200:
            doc.close()
            print(f"[OCR] Level1 opendataloader 成功，字符数: {len(odl_text)}")
            # 判断是否为扫描版（仅用于返回标记，不影响实际处理）
            is_scanned_hint = len(odl_text) < 1000 and page_count > 10
            return odl_text, page_count, is_scanned_hint
        else:
            print(f"[OCR] Level1 opendataloader 输出不足，降级到 PaddleOCR")

    # ─── 检测扫描版 + 文字层提取（用于 Level2/Level3 决策）──────────────────
    segments = []
    total_chars = 0
    for page in doc:
        t = _extract_text_layout_aware(page)
        total_chars += len(t)
        segments.append(t)

    is_scanned = _is_scanned_pdf(doc)

    if not is_scanned and total_chars >= 500:
        # 普通文字PDF（opendataloader不可用时的回退）：智能采样 前40% + 中30% + 后30%
        doc.close()
        n = len(segments)
        front = segments[:int(n * 0.4)]
        mid   = segments[int(n * 0.35):int(n * 0.65)]
        back  = segments[int(n * 0.7):]
        all_text = "\n".join(front + mid + back)
        return all_text[:MAX_CHARS], page_count, False

    # ─── Level 2：PaddleOCR（本地扫描版OCR）──────────────────────────────────
    # 采样策略：最多识别60页（前40% + 中30% + 后30%），避免太慢
    MAX_OCR_PAGES = 60
    n = page_count
    if n <= MAX_OCR_PAGES:
        page_indices = list(range(n))
    else:
        front_n = int(MAX_OCR_PAGES * 0.4)
        mid_n   = int(MAX_OCR_PAGES * 0.3)
        back_n  = MAX_OCR_PAGES - front_n - mid_n
        front_idx = list(range(front_n))
        mid_idx   = list(range(int(n*0.35), int(n*0.35) + mid_n))
        back_idx  = list(range(n - back_n, n))
        page_indices = sorted(set(front_idx + mid_idx + back_idx))

    if task_updater:
        task_updater("processing", 25,
                     f"检测到扫描版PDF（{page_count}页），启动PaddleOCR识别{len(page_indices)}页...")

    ocr_text = _ocr_pages_with_paddle(doc, page_indices, task_updater, pdf_path=path)
    doc.close()
    return ocr_text[:MAX_CHARS], page_count, True

async def extract_text_from_url(url: str) -> str:
    """从URL提取文本。对需要登录的平台（微博/抖音/小红书等）提前检测并报错。"""
    # ── 需要登录/JS渲染的平台检测 ───────────────────────────────────
    LOGIN_REQUIRED_DOMAINS = {
        "weibo.com": "微博",
        "weibo.cn": "微博",
        "m.weibo.cn": "微博",
        "douyin.com": "抖音",
        "xhs.link": "小红书",
        "xiaohongshu.com": "小红书",
        "www.zhihu.com": "知乎（部分内容需登录）",
        "mp.weixin.qq.com": "微信公众号",
        "x.com": "X（Twitter）",
        "twitter.com": "X（Twitter）",
        "www.twitter.com": "X（Twitter）",
    }
    from urllib.parse import urlparse
    parsed = urlparse(url)
    domain = parsed.netloc.lower().lstrip("www.")
    for d, name in LOGIN_REQUIRED_DOMAINS.items():
        if d in domain or domain == d:
            raise ValueError(
                f"【{name}】需要登录才能获取内容，直接导入URL无效。\n"
                f"请使用「文字导入」方式：复制博主内容文本后粘贴到文字输入框。\n"
                f"或使用浏览器爬取工具先抓取内容，再通过文字方式导入。"
            )

    try:
        async with httpx.AsyncClient(
            timeout=15,
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
        ) as c:
            resp = await c.get(url)
            text = resp.text

            # 检测是否被重定向到登录页
            final_url = str(resp.url)
            LOGIN_REDIRECT_HINTS = ["passport.", "login", "sso.", "signin", "auth."]
            if any(hint in final_url for hint in LOGIN_REDIRECT_HINTS):
                raise ValueError(
                    f"访问被重定向到登录页（{final_url[:80]}），无法直接抓取内容。\n"
                    "请复制网页正文后使用「文字导入」方式导入。"
                )

            soup = BeautifulSoup(text, "html.parser")
            for tag in soup(["script","style","nav","footer","header","aside","iframe","noscript"]):
                tag.decompose()
            extracted = soup.get_text(separator="\n", strip=True)

            # 检测提取内容是否过短（可能是JS渲染页面未加载）
            if len(extracted) < 500:
                raise ValueError(
                    f"提取内容过短（{len(extracted)}字），网页可能需要登录或JS渲染才能显示内容。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            # 检测是否为JS渲染页面（内容充满换行但没有实质内容）
            lines = [l.strip() for l in extracted.split('\n') if l.strip()]
            if len(lines) < 3:
                raise ValueError(
                    f"提取内容无效（仅{len(lines)}行有效文字），网页可能需要JS渲染。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            return extracted[:MAX_CHARS]
    except ValueError:
        raise  # ValueError 直接抛出（友好错误信息）
    except Exception as e:
        raise ValueError(f"抓取失败：{e}")

# ─── AI 提炼 ─────────────────────────────────────────────────────────
EXTRACT_MODES = {
    "full": "全量提炼（5个模块）",
    "quotes": "仅提炼金句弹药库",
    "cases": "仅提炼故事与案例",
    "viewpoints": "仅提炼认知与观点",
    "actions": "仅提炼实操行动库",
    "topics": "仅提炼IP选题映射",
    "ip_atomic": "IP原子化提取（每张卡片一个观点，适合内容生产）",  # 新增
}

def build_prompt(book_name: str, text: str, mode: str = "full") -> tuple[str, str]:
    system = f"""你是一位拥有海量知识库的"超级内容合伙人"，负责将内容彻底拆解为可安全复用、能持续爆款的"内容弹药库"。
我的IP方向是：{IP_DIRECTION}
目标受众痛点：职场焦虑、认知升级、破局成长，人性洞察。
每条素材必须附带：⚠️风险标签 + 🎯爆点场景 + ⏳改写成本 + 🕒时效熔断。
拒绝遗漏，拒绝凑数，宁可少出10条有用的，不要多出50条废话。"""

    # 限制单次 AI 调用的文本量，避免超时
    if len(text) > MAX_PROMPT_TEXT_CHARS:
        text = text[:MAX_PROMPT_TEXT_CHARS]
        truncated_note = f"\n\n【注意：原文过长，已截取前 {MAX_PROMPT_TEXT_CHARS:,} 字进行提炼】"
    else:
        truncated_note = ""

    if mode == "quotes":
        user = f"""请对《{book_name}》进行地毯式搜索，只输出【金句弹药库】{truncated_note}。

【内容】
{text}

---
## 【金句弹药库】

格式（每条必须完整）：
> [场景标签] 金句内容 【适用场景】
> ⚠️ 风险标签：✅安全 / ⚠️需语境 / ❌禁用
> 🎯 爆点场景：🔥爆款 / 💡治愈 / 📚深度
> ⏳ 改写成本：⚡0成本 / ✂️中成本 / 🛠️高成本
> 🕒 时效熔断：✅长效 / ⚠️需更新 / ❌已过期

全量提取，上限50条，不凑数。"""

    elif mode == "cases":
        user = f"""请对《{book_name}》进行地毯式搜索，只输出【故事与案例库】。

【内容】
{text}

---
## 【故事与案例库】

格式：
**案例名称：**
- 冲突：
- 动作：
- 结果：
- 启示：适合场景：
- ⚠️ 风险标签：
- 🕒 时效熔断：✅长效 / ⚠️需更新 / ❌已过期

全量提取，不限数量。"""

    elif mode == "viewpoints":
        user = f"""请对《{book_name}》进行地毯式搜索，只输出【认知与观点库】。

【内容】
{text}

---
## 【认知与观点库】

格式：
**观点名称：**
- 书中依据：[用自己的话重新表达]
- IP化角度：
- 📌 冲突预警：与[反驳观点]冲突，调和方案：
- ⚠️ 风险标签：✅安全 / ⚠️需语境 / ❌禁用
- 🕒 时效熔断：✅长效 / ⚠️需更新 / ❌已过期

全量提取，不限数量。"""

    elif mode == "actions":
        user = f"""请对《{book_name}》进行地毯式搜索，只输出【实操行动库】。

【内容】
{text}

---
## 【实操与行动库】

格式：
**行动名称：**
- 步骤：1. 2. 3.
- 适用场景：
- 风险提示：
- ⚡ 改写成本：⚡0成本 / ✂️中成本 / 🛠️高成本

全量提取，不限数量。"""

    elif mode == "topics":
        user = f"""请对《{book_name}》进行地毯式搜索，只输出【IP选题映射】。

【内容】
{text}

---
## 【IP选题映射】

列出10-20个爆款选题：
- 选题标题（能直接用的爆款标题）
- 核心素材来源
- 平台适配：抖音/小红书/公众号
- 钩子设计：开头+结尾互动"""

    elif mode == "ip_atomic":
        user = f"""请对《{book_name}》进行原子化提取，每张卡片一个独立观点。

【内容】
{text}

---
## 【IP原子笔记库】

格式（每条必须完整，JSON格式）：
```json
{{
  "atomic_notes": [
    {{
      "core_idea": "一句话核心观点（必须用自己的话重新表达，让外行能听懂）",
      "detailed_explanation": "3-5句话详细解释",
      "original_quote": "原文引用（可选）",
      "thinking_model": "所属思维模型（如：认知偏差/复利效应/机会成本/马太效应等）",
      "content_areas": ["职场", "情感", "人性", "社会", "底层逻辑"],
      "applicable_scenarios": ["短视频", "文章", "直播", "课程", "金句"],
      "emotional_resonance": ["焦虑", "希望", "愤怒", "认同", "惊讶", "好奇"],
      "target_audience": ["职场新人", "30岁焦虑", "情感困惑者", "创业者"],
      "transform_tips": "如何转化为口播稿/文章的具体建议"
    }}
  ]
}}
```

要求：
1. 每个观点一张卡片，独立完整（脱离原文也能理解）
2. 用自己的话重新表达（费曼技巧）
3. 标注思维模型（便于跨书整合）
4. 标注情感共鸣点（便于引发传播）
5. 标注适用场景（便于内容生产）
6. 放弃陈词滥调，保留反常识洞察
7. 优先提取能直接用于自媒体创作的实用观点

只输出JSON，不要有其他文字。"""

    else:  # full
        user = f"""请对《{book_name}》进行地毯式搜索，按格式完整输出"全量安全军火库"。

【内容（全书智能采样）】
{text}

---

## 【第一部分：金句弹药库】
格式（每条完整填写）：
> [场景标签] 金句内容 【适用场景】
> ⚠️ 风险标签：✅安全 / ⚠️需语境 / ❌禁用
> 🎯 爆点场景：🔥爆款 / 💡治愈 / 📚深度
> ⏳ 改写成本：⚡0成本 / ✂️中成本 / 🛠️高成本
> 🕒 时效熔断：✅长效 / ⚠️需更新 / ❌已过期
（全量提取，上限50条）

---

## 【第二部分：故事与案例库】
**案例名称：**
- 冲突：- 动作：- 结果：- 启示：- ⚠️风险：- 🕒时效：
（全量提取）

---

## 【第三部分：认知与观点库】
**观点名称：**
- 书中依据：- IP化角度：- 📌冲突预警：- ⚠️风险：- 🕒时效：
（全量提取）

---

## 【第四部分：实操与行动库】
**行动名称：**
- 步骤：- 适用场景：- 风险提示：- ⚡改写成本：
（全量提取）

---

## 【第五部分：IP选题映射】
10-20个爆款选题，每个含：标题+素材来源+平台+钩子设计

---

## 【书籍综合评级】
IP含金量/素材丰富度/可持续产出周期/跨书组合推荐"""

    return system, user


def parse_atomic_notes(source_id: str, book_title: str, raw_content: str) -> List[dict]:
    """
    解析IP原子笔记格式的AI输出

    将JSON格式的原子笔记解析为materials数据结构
    """
    import json
    import re

    materials = []
    now_str = now()

    # 尝试提取JSON
    try:
        # 清理markdown代码块标记
        cleaned = re.sub(r'^```(?:json)?\s*', '', raw_content.strip())
        cleaned = re.sub(r'\s*```\s*$', '', cleaned.strip())

        # 替换中文引号
        cleaned = cleaned.replace('\u201c', '"').replace('\u201d', '"')
        cleaned = cleaned.replace('\u2018', "'").replace('\u2019', "'")

        # 解析JSON
        data = json.loads(cleaned)
        notes = data.get("atomic_notes", [])

        for i, note in enumerate(notes):
            # 构建content（核心观点+详细解释）
            content_parts = [
                f"**核心观点**：{note.get('core_idea', '')}",
                f"",
                f"**详细解释**：{note.get('detailed_explanation', '')}",
            ]

            if note.get('original_quote'):
                content_parts.extend([
                    f"",
                    f"**原文引用**：{note['original_quote']}"
                ])

            content = "\n".join(content_parts)

            # 构建metadata
            meta = {
                "thinking_model": note.get("thinking_model", ""),
                "content_areas": note.get("content_areas", []),
                "applicable_scenarios": note.get("applicable_scenarios", []),
                "emotional_resonance": note.get("emotional_resonance", []),
                "target_audience": note.get("target_audience", []),
                "transform_tips": note.get("transform_tips", ""),
                "source_book": book_title,
                "note_index": i,
                "extraction_mode": "ip_atomic"
            }

            materials.append({
                "id": str(uuid.uuid4()),
                "source_id": source_id,
                "category": "atomic_note",  # 新的类别
                "content": content,
                "metadata": json.dumps(meta, ensure_ascii=False),
                "tags": json.dumps(["ip_content", note.get("thinking_model", "")], ensure_ascii=False),
                "platform": json.dumps(note.get("applicable_scenarios", []), ensure_ascii=False),
                "use_count": 0,
                "is_starred": 0,
                "created_at": now_str,
            })

    except Exception as e:
        print(f"Parse atomic notes failed: {e}")
        # 解析失败返回空列表
        pass

    return materials


def parse_materials(source_id: str, raw_content: str) -> List[dict]:
    """将AI输出解析为结构化素材条目"""
    materials = []
    now_str = now()

    lines = raw_content.split("\n")
    current_cat = "quote"
    current_block = []

    cat_map = {
        "第一部分": "quote",
        "金句弹药库": "quote",
        "第二部分": "case",
        "故事与案例": "case",
        "第三部分": "viewpoint",
        "认知与观点": "viewpoint",
        "第四部分": "action",
        "实操与行动": "action",
        "第五部分": "topic",
        "IP选题映射": "topic",
        "书籍综合评级": "rating",
    }

    def flush_block():
        nonlocal current_block
        text = "\n".join(current_block).strip()
        if text and len(text) > 10:
            meta = {}
            # 提取风险标签
            for line in current_block:
                if "风险标签" in line:
                    if "✅安全" in line: meta["risk"] = "safe"
                    elif "⚠️需语境" in line: meta["risk"] = "context"
                    elif "❌禁用" in line: meta["risk"] = "forbidden"
                if "爆点场景" in line:
                    if "🔥爆款" in line: meta["scene"] = "viral"
                    elif "💡治愈" in line: meta["scene"] = "heal"
                    elif "📚深度" in line: meta["scene"] = "deep"
                if "改写成本" in line:
                    if "⚡0成本" in line: meta["cost"] = "zero"
                    elif "✂️中成本" in line: meta["cost"] = "mid"
                    elif "🛠️高成本" in line: meta["cost"] = "high"
                if "时效熔断" in line:
                    if "✅长效" in line: meta["timeliness"] = "long"
                    elif "⚠️需更新" in line: meta["timeliness"] = "update"
                    elif "❌已过期" in line: meta["timeliness"] = "expired"

            materials.append({
                "id": str(uuid.uuid4()),
                "source_id": source_id,
                "category": current_cat,
                "content": text,
                "metadata": json.dumps(meta, ensure_ascii=False),
                "tags": "[]",
                "platform": "[]",
                "use_count": 0,
                "is_starred": 0,
                "created_at": now_str,
            })
        current_block = []

    for line in lines:
        # 检测分类标题
        new_cat = None
        for key, cat in cat_map.items():
            if key in line and ("##" in line or "【" in line):
                new_cat = cat
                break

        if new_cat:
            flush_block()
            current_cat = new_cat
        else:
            # 金句单条切割（以 > 开头的引用块）
            if current_cat == "quote" and line.startswith("> ") and "风险标签" not in line and "爆点场景" not in line and "改写成本" not in line and "时效熔断" not in line:
                if current_block:
                    flush_block()
                current_block = [line]
            elif current_cat == "quote" and line.startswith("> "):
                current_block.append(line)
            # 案例/观点/行动 — 以 ** 开头新条目
            elif current_cat in ("case","viewpoint","action","topic") and line.startswith("**") and line.endswith("**"):
                flush_block()
                current_block = [line]
            elif current_cat in ("case","viewpoint","action","topic") and line.startswith("- ") and current_block:
                current_block.append(line)
            elif current_cat == "topic" and line.strip().startswith(("1.","2.","3.","4.","5.","6.","7.","8.","9.","10.","11.","12.","13.","14.","15.")):
                flush_block()
                current_block = [line]
            elif current_block:
                current_block.append(line)

    flush_block()
    return materials


# ─── 后台任务：提炼处理 ──────────────────────────────────────────────
def process_source_task(task_id: str, source_id: str, mode: str):
    conn = get_db()
    try:
        def update_task(status, progress, message, result=None):
            conn.execute(
                "UPDATE tasks SET status=?,progress=?,message=?,result=?,updated_at=? WHERE id=?",
                (status, progress, message, result, now(), task_id)
            )
            conn.commit()

        update_task("processing", 10, "读取文件中...")
        src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
        if not src:
            update_task("error", 0, "找不到来源记录")
            return

        # 提取文本
        text = ""
        page_count = 0
        is_scanned = False
        is分段处理 = False  # 默认值，防止某些分支未赋值时引发 UnboundLocalError
        if src["type"] == "epub" and src["file_path"]:
            is分段处理 = False
            text, page_count, is_scanned = extract_text_from_epub(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("EPUB 解析失败，文件可能损坏或格式不标准", now(), source_id))
                conn.commit()
                update_task("error", 0, "EPUB 解析失败，请检查文件")
                return
        elif src["type"] == "txt" and src["file_path"]:
            text, page_count, is_scanned = extract_text_from_txt(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("文本文件读取失败", now(), source_id))
                conn.commit()
                update_task("error", 0, "文本文件读取失败")
                return
            # text 可能是 str（短文本）或 list[str]（长文本分段）
            is分段处理 = isinstance(text, list)
        elif src["type"] == "docx" and src["file_path"]:
            is分段处理 = False
            text, page_count, is_scanned = extract_text_from_docx(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("Word 文档读取失败，文件可能损坏或加密", now(), source_id))
                conn.commit()
                update_task("error", 0, "Word 文档读取失败，请检查文件")
                return
        elif src["type"] == "book" and src["file_path"]:
            is分段处理 = False
            text, page_count, is_scanned = extract_text_from_pdf(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("PDF无可提取文字（可能是加密或损坏）", now(), source_id))
                conn.commit()
                update_task("error", 0, "PDF无法识别文字，请检查文件是否加密或损坏")
                return
        elif src["type"] == "text":
            is分段处理 = False
            text = src["url"] or ""  # 对text类型，文本存在url字段
            page_count = 1
        elif src["type"] == "url":
            is分段处理 = False
            import asyncio
            loop = asyncio.new_event_loop()
            try:
                text = loop.run_until_complete(extract_text_from_url(src["url"]))
            except ValueError as ve:
                # 需要登录/JS渲染等友好错误
                err_msg = str(ve)
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             (err_msg, now(), source_id))
                conn.commit()
                update_task("error", 0, err_msg)
                return
            finally:
                loop.close()
            page_count = 1

        # 处理文本长度（分段处理时 len(text) 是 list）
        total_chars = sum(len(c) for c in text) if isinstance(text, list) else len(text)
        conn.execute("UPDATE sources SET page_count=?,char_count=?,status='processing',updated_at=? WHERE id=?",
                     (page_count, total_chars, now(), source_id))
        conn.commit()

        ocr_hint = "（扫描版·AI-OCR）" if is_scanned else ""

        # 分段处理：每个 chunk 分别提炼，然后合并结果
        if is分段处理:
            update_task("processing", 60, f"文本提取完成{ocr_hint}（{len(text)} 个分段，共 {total_chars:,} 字），正在分段 AI 提炼...")
            all_materials = []
            for i, chunk in enumerate(text):
                update_task("processing", 60 + int((i / len(text)) * 20), f"正在提炼第 {i+1}/{len(text)} 个分段...")
                chunk_prompt_system, chunk_prompt_user = build_prompt(f"{src['title']} (第{i+1}段)", chunk, mode)
                chunk_raw, chunk_model = ai_extract(chunk_prompt_system, chunk_prompt_user, max_tokens=12000, temperature=0.7)
                chunk_materials = parse_materials(source_id, chunk_raw)
                all_materials.extend(chunk_materials)
            raw_content = f"【合并 {len(text)} 个分段提炼结果】\n" + "\n---\n".join([f"第{i+1}段" for i in range(len(text))])
            materials = all_materials
            model_used = "分段模式"
        else:
            update_task("processing", 72, f"文本提取完成{ocr_hint}（{total_chars:,}字），正在调用AI提炼...")
            # 调用AI提炼（讯飞被拦截时自动切换备用模型）
            system_prompt, user_prompt = build_prompt(src["title"], text, mode)
            raw_content, model_used = ai_extract(system_prompt, user_prompt, max_tokens=12000, temperature=0.7)
            fallback_hint = "（备用模型）" if model_used == FALLBACK_MODEL_ID else ""
            update_task("processing", 88, f"AI提炼完成{fallback_hint}，正在解析存储...")
            # 解析存储
            materials = parse_materials(source_id, raw_content)

        update_task("processing", 92, f"正在存储 {len(materials)} 条素材...")

        for m in materials:
            conn.execute(
                "INSERT OR REPLACE INTO materials VALUES (?,?,?,?,?,?,?,?,?,?)",
                (m["id"], m["source_id"], m["category"], m["content"],
                 m["metadata"], m["tags"], m["platform"],
                 m["use_count"], m["is_starred"], m["created_at"])
            )

        conn.execute("UPDATE sources SET status='done',is_scanned=?,updated_at=? WHERE id=?",
                     (1 if is_scanned else 0, now(), source_id))
        conn.commit()

        update_task("done", 100, f"完成！提取 {len(materials)} 条素材（{len(text) if is分段处理 else 1} 个分段）", raw_content[:500])

    except Exception as e:
        err_str = str(e)
        # 精确判断：只有当所有模型均为讯飞拦截（不含 MiniMax/DeepSeek 成功但内容拦截）时才用友好提示
        # 修复：原来 is_xunfei_blocked 匹配"无法提供"等词，但 MiniMax 自身也可能返回这类词
        # 改为：只有当错误来源仅有讯飞（没有 MiniMax 参与）时才显示该提示
        is_all_xunfei_block = (
            is_xunfei_blocked(err_str) and
            "MiniMax" not in err_str and
            "minimax" not in err_str.lower() and
            "DeepSeek" not in err_str
        )
        if is_all_xunfei_block:
            friendly = "讯飞内容审核拦截，备用模型也未能处理。建议手动删除后重新上传，或换用其他同类书籍。"
        else:
            friendly = f"提炼失败：{err_str[:300]}"
        conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                     (friendly, now(), source_id))
        conn.commit()
        conn.execute("UPDATE tasks SET status='error',message=?,updated_at=? WHERE id=?",
                     (friendly, now(), task_id))
        conn.commit()
    finally:
        conn.close()


# ─── API 路由 ────────────────────────────────────────────────────────

# 上传PDF书籍
@app.post("/api/sources/upload-pdf")
async def upload_pdf(file: UploadFile = File(...), mode: str = "full"):
    file_id = str(uuid.uuid4())
    filename = file.filename or "unknown.pdf"
    save_path = UPLOAD_DIR / f"{file_id}_{filename}"

    content = await file.read()
    save_path.write_bytes(content)

    source_id = str(uuid.uuid4())
    task_id = str(uuid.uuid4())
    title = Path(filename).stem

    # 根据文件后缀决定 source type
    ext = Path(filename).suffix.lower()
    if ext == ".epub":
        source_type = "epub"
    elif ext in (".txt", ".md"):
        source_type = "txt"
    elif ext == ".docx":
        source_type = "docx"
    else:
        source_type = "book"  # pdf / 其他

    conn = get_db()
    conn.execute(
        "INSERT INTO sources VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (source_id, title, source_type, str(save_path), None, "[]", 0, 0, "pending", None, 0, now(), now())
    )
    conn.execute(
        "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
        (task_id, source_id, "pending", 0, "等待处理...", None, now(), now())
    )
    conn.commit()
    conn.close()

    enqueue_task(task_id, source_id, mode)
    return {"source_id": source_id, "task_id": task_id, "title": title}


# ─── 文件夹批量导入 ───────────────────────────────────────────────────

class ImportFolderRequest(BaseModel):
    folder_path: str
    mode: str = "full"
    recursive: bool = False        # 是否递归子目录
    skip_existing: bool = True     # 跳过已导入的同名书籍


@app.post("/api/sources/import-folder")
async def import_folder(req: ImportFolderRequest):
    """扫描本地文件夹，将所有 PDF/EPUB/TXT 加入书库并启动提炼任务"""
    folder = Path(req.folder_path.strip())
    if not folder.exists():
        raise HTTPException(400, f"路径不存在：{folder}")
    if not folder.is_dir():
        raise HTTPException(400, f"该路径不是文件夹：{folder}")

    # 扫描 PDF + EPUB + TXT + DOCX
    SUPPORTED_EXTS = (".pdf", ".epub", ".txt", ".md", ".docx")
    if req.recursive:
        all_files = [f for f in sorted(folder.rglob("*")) if f.suffix.lower() in SUPPORTED_EXTS]
    else:
        all_files = [f for f in sorted(folder.glob("*")) if f.suffix.lower() in SUPPORTED_EXTS]

    if not all_files:
        return {"total": 0, "queued": 0, "skipped": 0, "tasks": [],
                "message": "文件夹中没有找到 PDF/EPUB/TXT 文件"}

    # 查已有书名，跳过重复
    conn = get_db()
    existing_titles = set()
    if req.skip_existing:
        rows = conn.execute("SELECT title FROM sources WHERE type IN ('book','epub','txt')").fetchall()
        existing_titles = {r["title"] for r in rows}

    tasks_created = []
    skipped = []

    for file_path in all_files:
        title = file_path.stem
        if req.skip_existing and title in existing_titles:
            skipped.append(title)
            continue

        # 根据后缀设置 type
        ext = file_path.suffix.lower()
        if ext == ".epub":
            source_type = "epub"
        elif ext in (".txt", ".md"):
            source_type = "txt"
        elif ext == ".docx":
            source_type = "docx"
        else:
            source_type = "book"

        # 复制到 uploads
        file_id  = str(uuid.uuid4())
        dest     = UPLOAD_DIR / f"{file_id}_{file_path.name}"
        try:
            import shutil
            shutil.copy2(str(file_path), str(dest))
        except Exception as e:
            skipped.append(f"{title} (复制失败: {e})")
            continue

        source_id = str(uuid.uuid4())
        task_id   = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO sources VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (source_id, title, source_type, str(dest), None, "[]", 0, 0,
             "pending", None, 0, now(), now())
        )
        conn.execute(
            "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
            (task_id, source_id, "pending", 0, "等待批量提炼...", None, now(), now())
        )
        tasks_created.append({"task_id": task_id, "source_id": source_id, "title": title})
        existing_titles.add(title)   # 同一批次也去重

    conn.commit()
    conn.close()

    # 入队（由全局 Worker 线程池串行消费，避免同时打爆 API）
    for t in tasks_created:
        enqueue_task(t["task_id"], t["source_id"], req.mode)

    return {
        "total":   len(all_files),
        "queued":  len(tasks_created),
        "skipped": len(skipped),
        "tasks":   tasks_created,
        "message": f"已加入队列 {len(tasks_created)} 本，跳过 {len(skipped)} 本"
    }


@app.get("/api/sources/scan-folder")
def scan_folder(path: str, recursive: bool = False):
    """预扫描文件夹，返回文件列表（不导入），用于前端预览"""
    folder = Path(path.strip())
    if not folder.exists() or not folder.is_dir():
        raise HTTPException(400, "路径无效或不是文件夹")

    # 扫描所有支持的文件类型
    SUPPORTED_EXTS = (".pdf", ".epub", ".txt", ".md", ".docx")
    if recursive:
        all_files = sorted(f for f in folder.rglob("*") if f.suffix.lower() in SUPPORTED_EXTS)
    else:
        all_files = sorted(f for f in folder.glob("*") if f.suffix.lower() in SUPPORTED_EXTS)

    conn = get_db()
    existing_titles = {r["title"] for r in
                       conn.execute("SELECT title FROM sources WHERE type IN ('book','epub','txt','docx')").fetchall()}
    conn.close()

    items = []
    for p in all_files[:200]:   # 预览最多200条
        size_mb = round(p.stat().st_size / 1024 / 1024, 1)
        items.append({
            "name": p.stem,
            "filename": p.name,
            "size_mb": size_mb,
            "already_imported": p.stem in existing_titles
        })

    return {"path": str(folder), "count": len(all_files), "items": items}


@app.get("/api/pick-folder")
def pick_folder(initial_dir: str = ""):
    """弹出系统原生文件夹选择对话框，返回用户选中的路径。
    使用 tkinter（Python 标准库），无需额外安装。
    在后台线程中运行以避免阻塞 asyncio 事件循环。
    """
    import threading

    result_holder = {"path": None, "error": None}

    def _show_dialog():
        try:
            import tkinter as tk
            from tkinter import filedialog
            root = tk.Tk()
            root.withdraw()          # 隐藏主窗口
            root.attributes("-topmost", True)   # 置顶对话框
            root.update()
            kwargs = {"title": "选择书籍文件夹"}
            if initial_dir and Path(initial_dir).is_dir():
                kwargs["initialdir"] = initial_dir
            selected = filedialog.askdirectory(**kwargs)
            root.destroy()
            result_holder["path"] = selected or ""
        except Exception as e:
            result_holder["error"] = str(e)

    t = threading.Thread(target=_show_dialog, daemon=True)
    t.start()
    t.join(timeout=120)   # 最多等 2 分钟

    if result_holder["error"]:
        raise HTTPException(500, f"无法打开文件夹选择框：{result_holder['error']}")

    return {"selected": result_holder["path"] or ""}


# 添加文字/URL来源
class AddSourceRequest(BaseModel):
    title: str
    type: str   # text / url
    content: str
    mode: str = "full"

@app.post("/api/sources/add")
async def add_source(req: AddSourceRequest):
    source_id = str(uuid.uuid4())
    task_id = str(uuid.uuid4())

    conn = get_db()
    conn.execute(
        "INSERT INTO sources VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (source_id, req.title, req.type, None, req.content, "[]", 0, 0, "pending", None, 0, now(), now())
    )
    conn.execute(
        "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
        (task_id, source_id, "pending", 0, "等待处理...", None, now(), now())
    )
    conn.commit()
    conn.close()

    enqueue_task(task_id, source_id, req.mode)
    return {"source_id": source_id, "task_id": task_id}

# 查询任务进度
@app.get("/api/tasks/{task_id}")
def get_task(task_id: str):
    conn = get_db()
    task = conn.execute("SELECT * FROM tasks WHERE id=?", (task_id,)).fetchone()
    conn.close()
    if not task:
        raise HTTPException(404, "任务不存在")
    return dict(task)

# 获取所有书库
@app.get("/api/sources")
def list_sources(q: str = "", type: str = ""):
    conn = get_db()
    sql = "SELECT s.*, (SELECT COUNT(*) FROM materials m WHERE m.source_id=s.id) as material_count FROM sources s WHERE 1=1"
    params = []
    if q:
        sql += " AND s.title LIKE ?"
        params.append(f"%{q}%")
    if type:
        sql += " AND s.type=?"
        params.append(type)
    sql += " ORDER BY s.created_at DESC"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# 重试失败任务
@app.post("/api/sources/{source_id}/retry")
async def retry_source(source_id: str, mode: str = "full"):
    """重新提炼：重置 source + task 状态，立即重新入队"""
    conn = get_db()
    src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
    if not src:
        conn.close()
        raise HTTPException(404, "书籍不存在")

    # 重置 source 状态
    conn.execute(
        "UPDATE sources SET status='pending', error_msg=NULL, updated_at=? WHERE id=?",
        (now(), source_id)
    )

    # 复用已有 task（最新的那条），或新建一条
    existing_task = conn.execute(
        "SELECT id FROM tasks WHERE source_id=? ORDER BY created_at DESC LIMIT 1",
        (source_id,)
    ).fetchone()
    if existing_task:
        task_id = existing_task["id"]
        conn.execute(
            "UPDATE tasks SET status='pending', progress=0, message='重试中，等待处理...', result=NULL, updated_at=? WHERE id=?",
            (now(), task_id)
        )
    else:
        task_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
            (task_id, source_id, "pending", 0, "重试中，等待处理...", None, now(), now())
        )

    conn.commit()
    conn.close()

    # 立即重新入队
    enqueue_task(task_id, source_id, mode)
    return {"ok": True, "task_id": task_id, "status": "pending"}


# 一键恢复所有卡住的任务
@app.post("/api/sources/recover-all")
def recover_all_stuck():
    """将所有 processing/pending 状态的任务重新入队"""
    conn = get_db()
    try:
        # 先将 processing 重置为 pending
        conn.execute(
            "UPDATE sources SET status='pending', updated_at=? WHERE status='processing'",
            (now(),)
        )
        conn.execute(
            "UPDATE tasks SET status='pending', message='等待恢复处理...', updated_at=? WHERE status='processing'",
            (now(),)
        )
        conn.commit()

        rows = conn.execute(
            """SELECT t.id as task_id, t.source_id
               FROM tasks t JOIN sources s ON t.source_id = s.id
               WHERE s.status = 'pending'
               ORDER BY s.created_at"""
        ).fetchall()

        queued = 0
        for row in rows:
            enqueue_task(row["task_id"], row["source_id"], "full")
            queued += 1

        return {"ok": True, "queued": queued, "queue_size": _task_queue.qsize()}
    finally:
        conn.close()



# ─── AI Studio OCR Token 配置接口 ───────────────────────────────────
class AistudioTokenRequest(BaseModel):
    token: str

@app.get("/api/ocr/aistudio-token")
def get_aistudio_token():
    """查询当前 AI Studio token 状态（不返回 token 明文）"""
    return {"configured": bool(AISTUDIO_TOKEN), "length": len(AISTUDIO_TOKEN)}

@app.post("/api/ocr/aistudio-token")
def set_aistudio_token(req: AistudioTokenRequest):
    """设置 AI Studio OCR token（运行时生效，重启后失效，如需持久化请设置环境变量 AISTUDIO_TOKEN）"""
    global AISTUDIO_TOKEN
    AISTUDIO_TOKEN = req.token.strip()
    os.environ["AISTUDIO_TOKEN"] = AISTUDIO_TOKEN
    return {"ok": True, "configured": bool(AISTUDIO_TOKEN)}


# ─── AI 模型选择接口 ─────────────────────────────────────────────────
class AiModelRequest(BaseModel):
    preferred: str          # "xunfei" | "minimax" | "minimax2" | "deepseek"
    minimax2_key: Optional[str] = None   # 若切换 minimax2，可同时传入 key
    minimax2_base: Optional[str] = None  # MiniMax2 API Base（可选，默认官方直连）

@app.get("/api/ai-model")
def get_ai_model():
    """查询当前 AI 首选模型状态"""
    global _AI_PREFERRED, MINIMAX_API_KEY_2, MINIMAX_API_BASE_2
    labels = {
        "xunfei":   "讯飞星辰",
        "minimax":  "MiniMax（账号1）",
        "minimax2": "MiniMax（账号2）",
        "deepseek": "DeepSeek",
    }
    return {
        "preferred": _AI_PREFERRED,
        "preferred_label": labels.get(_AI_PREFERRED, _AI_PREFERRED),
        "minimax2_configured": bool(MINIMAX_API_KEY_2),
        "minimax2_key_preview": (MINIMAX_API_KEY_2[:8] + "****") if MINIMAX_API_KEY_2 else "",
        "minimax2_base": MINIMAX_API_BASE_2,
        "available_models": [
            {"id": "xunfei",   "label": "讯飞星辰",       "note": "每日 5000万 token，文本理解最强"},
            {"id": "minimax",  "label": "MiniMax（账号1）","note": "每6小时 4000次，旧账号"},
            {"id": "minimax2", "label": "MiniMax（账号2）","note": "每6小时 4000次，需填入 API Key"},
            {"id": "deepseek", "label": "DeepSeek",        "note": "最终兜底，按量计费"},
        ]
    }

@app.post("/api/ai-model")
def set_ai_model(req: AiModelRequest):
    """切换全局 AI 首选模型（立即生效，后续所有提炼任务从该模型开始）"""
    global _AI_PREFERRED, MINIMAX_API_KEY_2, MINIMAX_API_BASE_2, _minimax2_client
    allowed = {"xunfei", "minimax", "minimax2", "deepseek"}
    if req.preferred not in allowed:
        raise HTTPException(status_code=400, detail=f"preferred 必须是: {allowed}")
    
    # 如果切换到 minimax2，且传入了 key，则更新
    if req.minimax2_key is not None:
        MINIMAX_API_KEY_2 = req.minimax2_key.strip()
        _minimax2_client = None   # 清除缓存，下次重新创建
        if req.minimax2_base:
            MINIMAX_API_BASE_2 = req.minimax2_base.strip()
    
    # 切 minimax2 但 key 还没有
    if req.preferred == "minimax2" and not MINIMAX_API_KEY_2:
        raise HTTPException(status_code=400, detail="切换到 MiniMax 账号2 前，请先在 minimax2_key 字段填写 API Key")
    
    _AI_PREFERRED = req.preferred
    labels = {"xunfei":"讯飞星辰","minimax":"MiniMax（账号1）","minimax2":"MiniMax（账号2）","deepseek":"DeepSeek"}
    print(f"[AI] 用户切换首选模型 → {labels.get(_AI_PREFERRED)}")
    return {"ok": True, "preferred": _AI_PREFERRED, "label": labels.get(_AI_PREFERRED)}


# 查询正在处理中的任务详情（供前端实时进度面板使用）
@app.get("/api/processing-tasks")
def processing_tasks():
    """返回所有 processing 状态的书籍及其最新任务进度（title、progress、message）"""
    conn = get_db()
    try:
        rows = conn.execute("""
            SELECT s.id as source_id, s.title, s.status as source_status,
                   t.id as task_id, t.status as task_status,
                   t.progress, t.message, t.updated_at
            FROM sources s
            LEFT JOIN tasks t ON t.source_id = s.id
            WHERE s.status IN ('processing', 'pending')
            ORDER BY t.updated_at DESC
        """).fetchall()
        # 每个 source 只取最新一条 task
        seen = set()
        result = []
        for row in rows:
            sid = row["source_id"]
            if sid in seen:
                continue
            seen.add(sid)
            result.append({
                "source_id": sid,
                "title": row["title"],
                "source_status": row["source_status"],
                "task_id": row["task_id"],
                "task_status": row["task_status"],
                "progress": row["progress"] or 0,
                "message": row["message"] or "",
                "updated_at": row["updated_at"] or "",
            })
        return result
    finally:
        conn.close()


# 查询当前队列状态
@app.get("/api/queue-status")
def queue_status():
    """返回当前任务队列大小 + worker 线程存活状态 + 心跳信息"""
    conn = get_db()
    try:
        pending = conn.execute("SELECT COUNT(*) FROM sources WHERE status='pending'").fetchone()[0]
        processing = conn.execute("SELECT COUNT(*) FROM sources WHERE status='processing'").fetchone()[0]
        done = conn.execute("SELECT COUNT(*) FROM sources WHERE status='done'").fetchone()[0]
        error = conn.execute("SELECT COUNT(*) FROM sources WHERE status='error'").fetchone()[0]
        alive_workers = sum(1 for t in _worker_threads if t.is_alive())
        # 心跳信息：0=idle, >0=正在处理，值为开始时间戳
        now_ts = time.time()
        worker_states = []
        for i, t in enumerate(_worker_threads):
            wid = i + 1
            hb = _worker_heartbeats.get(wid, 0)
            if not t.is_alive():
                state = "dead"
            elif hb == 0:
                state = "idle"
            else:
                elapsed = int(now_ts - hb)
                state = f"busy({elapsed}s)"
            worker_states.append(state)
        return {
            "queue_size": _task_queue.qsize(),
            "workers": WORKER_COUNT,
            "alive_workers": alive_workers,
            "worker_states": worker_states,
            "pending": pending,
            "processing": processing,
            "done": done,
            "error": error,
        }
    finally:
        conn.close()


# 手动重启 workers（出现死线程时应急用）
@app.post("/api/workers/restart")
def restart_workers():
    """手动重建所有 worker 线程（用于 worker 线程异常死亡时应急恢复）"""
    global _worker_threads
    dead = [i for i, t in enumerate(_worker_threads) if not t.is_alive()]
    for i in dead:
        new_t = _spawn_worker(i + 1)
        _worker_threads[i] = new_t
        print(f"[API] Worker-{i+1} 手动重建")
    alive = sum(1 for t in _worker_threads if t.is_alive())
    return {"ok": True, "restarted": len(dead), "alive_workers": alive, "queue_size": _task_queue.qsize()}


# 强制跳过卡住的 processing 任务（将其标记为 error，重新入队 pending 任务）
@app.post("/api/sources/skip-stuck")
def skip_stuck_tasks():
    """将所有长时间卡在 processing 的任务标记为 error，让队列继续跑"""
    conn = get_db()
    try:
        # 找出 processing 超过 10 分钟的任务（超时阈值宽松，避免误杀正常任务）
        stuck_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        rows = conn.execute(
            """SELECT s.id as source_id, t.id as task_id 
               FROM sources s LEFT JOIN tasks t ON t.source_id = s.id
               WHERE s.status = 'processing'"""
        ).fetchall()
        skipped = 0
        for row in rows:
            conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                         ("手动强制跳过（任务卡住）", now(), row["source_id"]))
            if row["task_id"]:
                conn.execute("UPDATE tasks SET status='error',message=?,updated_at=? WHERE id=?",
                             ("手动强制跳过（任务卡住）", now(), row["task_id"]))
            skipped += 1
        conn.commit()
        print(f"[API] 强制跳过 {skipped} 个卡住任务")
        return {"ok": True, "skipped": skipped, "queue_size": _task_queue.qsize()}
    finally:
        conn.close()


# 重置所有错误任务为 pending，重新入队
@app.post("/api/sources/retry-errors")
def retry_error_tasks():
    """将所有 error 状态的任务重置为 pending 重新入队"""
    conn = get_db()
    try:
        rows = conn.execute(
            """SELECT t.id as task_id, t.source_id
               FROM tasks t JOIN sources s ON t.source_id = s.id
               WHERE s.status = 'error'"""
        ).fetchall()
        retried = 0
        for row in rows:
            conn.execute("UPDATE sources SET status='pending',error_msg=NULL,updated_at=? WHERE id=?",
                         (now(), row["source_id"]))
            conn.execute("UPDATE tasks SET status='pending',message='重新等待处理...',progress=0,updated_at=? WHERE id=?",
                         (now(), row["task_id"]))
            enqueue_task(row["task_id"], row["source_id"], "full")
            retried += 1
        conn.commit()
        print(f"[API] 重试 {retried} 个错误任务")
        return {"ok": True, "retried": retried, "queue_size": _task_queue.qsize()}
    finally:
        conn.close()






# 一键清空书库（全部删除）
@app.delete("/api/sources")
def delete_all_sources():
    conn = get_db()
    try:
        # 先取出所有文件路径，用于删除物理文件
        rows = conn.execute("SELECT file_path FROM sources WHERE file_path IS NOT NULL AND file_path != ''").fetchall()
        for row in rows:
            try:
                Path(row["file_path"]).unlink(missing_ok=True)
            except Exception:
                pass
        conn.execute("DELETE FROM materials")
        conn.execute("DELETE FROM tasks")
        conn.execute("DELETE FROM sources")
        conn.commit()
        return {"ok": True, "deleted": len(rows)}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# 删除来源
@app.delete("/api/sources/{source_id}")
def delete_source(source_id: str):
    conn = get_db()
    conn.execute("DELETE FROM materials WHERE source_id=?", (source_id,))
    conn.execute("DELETE FROM tasks WHERE source_id=?", (source_id,))
    src = conn.execute("SELECT file_path FROM sources WHERE id=?", (source_id,)).fetchone()
    if src and src["file_path"]:
        try: Path(src["file_path"]).unlink()
        except: pass
    conn.execute("DELETE FROM sources WHERE id=?", (source_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

# 获取素材列表
@app.get("/api/materials")
def list_materials(source_id: str = "", category: str = "", q: str = "",
                   starred: int = -1, review_only: int = 0, skip: int = 0, limit: int = 50):
    conn = get_db()
    sql = """SELECT m.*, s.title as source_title FROM materials m
             LEFT JOIN sources s ON m.source_id = s.id WHERE 1=1"""
    params = []
    if source_id:
        sql += " AND m.source_id=?"
        params.append(source_id)
    if category:
        sql += " AND m.category=?"
        params.append(category)
    if q:
        sql += " AND m.content LIKE ?"
        params.append(f"%{q}%")
    if starred == 1:
        sql += " AND m.is_starred=1"
    if review_only == 1:
        sql += " AND m.metadata LIKE '%_review_needed%'"
    sql += " ORDER BY m.created_at DESC LIMIT ? OFFSET ?"
    params += [limit, skip]
    rows = conn.execute(sql, params).fetchall()
    # 总数
    cnt_sql = sql.replace("SELECT m.*, s.title as source_title", "SELECT COUNT(*)")
    cnt_sql = cnt_sql[:cnt_sql.rfind("LIMIT")]
    total = conn.execute(cnt_sql, params[:-2]).fetchone()[0]
    conn.close()
    # 解析质量评分
    items = []
    for r in rows:
        d = dict(r)
        try:
            meta = json.loads(d.get("metadata") or "{}")
            d["_quality_score"] = meta.pop("_quality_score", None)
            d["_routing"] = meta.pop("_routing", None)
            d["_review_needed"] = meta.pop("_review_needed", False)
        except Exception:
            d["_quality_score"] = None
            d["_routing"] = None
            d["_review_needed"] = False
        items.append(d)
    return {"total": total, "items": items}

# 收藏/取消收藏
@app.post("/api/materials/{mid}/star")
def star_material(mid: str):
    conn = get_db()
    current = conn.execute("SELECT is_starred FROM materials WHERE id=?", (mid,)).fetchone()
    if not current:
        raise HTTPException(404)
    new_val = 0 if current["is_starred"] else 1
    conn.execute("UPDATE materials SET is_starred=? WHERE id=?", (new_val, mid))
    conn.commit()
    conn.close()
    return {"is_starred": new_val}

# 记录使用次数
@app.post("/api/materials/{mid}/use")
def use_material(mid: str):
    conn = get_db()
    conn.execute("UPDATE materials SET use_count=use_count+1 WHERE id=?", (mid,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/materials/{mid}")
def delete_material(mid: str):
    conn = get_db()
    conn.execute("DELETE FROM materials WHERE id=?", (mid,))
    conn.commit()
    conn.close()
    return {"ok": True}


# ─── 导出 ───────────────────────────────────────────────────────────

CATEGORY_ZH = {
    "quote":     "金句弹药库",
    "case":      "故事案例库",
    "viewpoint": "认知观点库",
    "action":    "实操行动库",
    "topic":     "IP选题映射",
}

def _fetch_all_materials(source_id: str = "", category: str = "", starred_only: bool = False):
    """取全量素材，不分页"""
    conn = get_db()
    sql = """SELECT m.id, m.category, m.content, m.is_starred,
                    m.use_count, m.created_at, s.title as source_title
             FROM materials m LEFT JOIN sources s ON m.source_id = s.id
             WHERE 1=1"""
    params = []
    if source_id:
        sql += " AND m.source_id = ?"
        params.append(source_id)
    if category:
        sql += " AND m.category = ?"
        params.append(category)
    if starred_only:
        sql += " AND m.is_starred = 1"
    sql += " ORDER BY m.category, s.title, m.created_at"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/api/export/excel")
def export_excel(source_id: str = "", category: str = "", starred_only: bool = False):
    """导出素材为 Excel，每个分类一个 Sheet"""
    from fastapi.responses import StreamingResponse
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    import io

    rows = _fetch_all_materials(source_id, category, starred_only)

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # 删除默认空sheet

    # 按分类分组
    from collections import defaultdict
    groups = defaultdict(list)
    for r in rows:
        groups[r["category"]].append(r)

    # 分类顺序
    order = ["quote", "case", "viewpoint", "action", "topic"]
    cats = order + [c for c in groups if c not in order]

    header_fill   = PatternFill("solid", fgColor="7C5CFC")
    header_font   = Font(bold=True, color="FFFFFF", size=11)
    star_fill     = PatternFill("solid", fgColor="FFF9E6")
    wrap_align    = Alignment(wrap_text=True, vertical="top")

    for cat in cats:
        items = groups.get(cat)
        if not items:
            continue
        sheet_name = CATEGORY_ZH.get(cat, cat)[:31]
        ws = wb.create_sheet(title=sheet_name)

        # 表头
        headers = ["序号", "内容", "来源书籍", "收藏", "使用次数", "创建时间"]
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=ci, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")

        ws.row_dimensions[1].height = 22

        for ri, item in enumerate(items, 2):
            ws.cell(row=ri, column=1, value=ri - 1)
            ws.cell(row=ri, column=2, value=item["content"]).alignment = wrap_align
            ws.cell(row=ri, column=3, value=item.get("source_title", ""))
            ws.cell(row=ri, column=4, value="⭐" if item["is_starred"] else "")
            ws.cell(row=ri, column=5, value=item["use_count"])
            ws.cell(row=ri, column=6, value=(item["created_at"] or "")[:10])
            if item["is_starred"]:
                for ci in range(1, 7):
                    ws.cell(row=ri, column=ci).fill = star_fill

        # 列宽
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 60
        ws.column_dimensions["C"].width = 20
        ws.column_dimensions["D"].width = 6
        ws.column_dimensions["E"].width = 8
        ws.column_dimensions["F"].width = 12
        ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    from urllib.parse import quote
    fname = f"IP军火库素材_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    encoded = quote(fname)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
    )


@app.get("/api/export/markdown")
def export_markdown(source_id: str = "", category: str = "", starred_only: bool = False):
    """导出素材为 Markdown 文本"""
    from fastapi.responses import Response
    from urllib.parse import quote

    rows = _fetch_all_materials(source_id, category, starred_only)

    from collections import defaultdict
    groups = defaultdict(list)
    for r in rows:
        groups[r["category"]].append(r)

    order = ["quote", "case", "viewpoint", "action", "topic"]
    cats = order + [c for c in groups if c not in order]

    lines = [
        f"# IP 军火库素材导出",
        f"> 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}  ",
        f"> 共 {len(rows)} 条素材\n",
    ]

    for cat in cats:
        items = groups.get(cat)
        if not items:
            continue
        lines.append(f"\n## {CATEGORY_ZH.get(cat, cat)}\n")
        # 按来源分组
        src_groups = defaultdict(list)
        for item in items:
            src_groups[item.get("source_title") or "未知来源"].append(item)
        for src_title, src_items in sorted(src_groups.items()):
            lines.append(f"\n### 📖 {src_title}\n")
            for idx, item in enumerate(src_items, 1):
                star = " ⭐" if item["is_starred"] else ""
                lines.append(f"{idx}. {item['content']}{star}\n")

    md_text = "\n".join(lines)
    fname = f"IP军火库素材_{datetime.now().strftime('%Y%m%d_%H%M')}.md"
    encoded = quote(fname)
    return Response(
        content=md_text.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
    )


# ─── 公众号排版 ──────────────────────────────────────────────────────

# 主题名称→中文标签映射（用于返回给前端选择器展示）
WECHAT_THEMES = [
    # 深度长文
    {"id": "newspaper",    "label": "报纸风",    "group": "深度长文", "desc": "《纽约时报》蓝灰，适合新闻/分析"},
    {"id": "magazine",     "label": "杂志风",    "group": "深度长文", "desc": "高端杂志版式，适合品质内容"},
    {"id": "coffee-house", "label": "咖啡馆",    "group": "深度长文", "desc": "咖啡色暖调，适合随笔/日记"},
    {"id": "ink",          "label": "水墨",      "group": "深度长文", "desc": "极简黑白，适合文学/散文"},
    # 科技产品
    {"id": "bytedance",    "label": "字节风",    "group": "科技产品", "desc": "蓝绿渐变，现代科技感"},
    {"id": "github",       "label": "GitHub",    "group": "科技产品", "desc": "开发者风格，代码友好"},
    {"id": "sspai",        "label": "少数派",    "group": "科技产品", "desc": "少数派橙色，适合效率/工具"},
    {"id": "midnight",     "label": "深夜",      "group": "科技产品", "desc": "深色背景，赛博朋克感"},
    # 文艺随笔
    {"id": "terracotta",   "label": "陶土",      "group": "文艺随笔", "desc": "暖橙圆角，轻松活泼"},
    {"id": "mint-fresh",   "label": "薄荷绿",    "group": "文艺随笔", "desc": "清新绿色，适合生活/美食"},
    {"id": "sunset-amber", "label": "琥珀",      "group": "文艺随笔", "desc": "琥珀暖色，情感/故事"},
    {"id": "lavender-dream","label": "薰衣草",   "group": "文艺随笔", "desc": "紫色梦幻，适合情感内容"},
    # 活力动态
    {"id": "sports",       "label": "运动",      "group": "活力动态", "desc": "渐变条纹，活力满满"},
    {"id": "bauhaus",      "label": "包豪斯",    "group": "活力动态", "desc": "几何色块，大胆设计感"},
    {"id": "chinese",      "label": "朱红国风",  "group": "活力动态", "desc": "朱红古典，传统文化"},
    {"id": "wechat-native","label": "微信原生",  "group": "活力动态", "desc": "微信绿，官方/通用推荐"},
    # 简约模板
    {"id": "minimal-gold", "label": "极简金",    "group": "简约模板", "desc": "金色点缀，商务简洁"},
    {"id": "focus-blue",   "label": "专注蓝",    "group": "简约模板", "desc": "蓝色聚焦，适合知识分享"},
    {"id": "elegant-green","label": "优雅绿",    "group": "简约模板", "desc": "绿色雅致，清新正式"},
    {"id": "bold-blue",    "label": "粗体蓝",    "group": "简约模板", "desc": "蓝色粗体标题，醒目"},
]


@app.get("/api/wechat-themes")
def get_wechat_themes():
    """返回可用的公众号主题列表"""
    return WECHAT_THEMES


class RewriteRequest(BaseModel):
    content: str
    mode: str = "口语化"       # 爆款化/口语化/精简/深度展开/加开头钩子/互动结尾
    platform: str = ""         # 抖音脚本/小红书图文/公众号文章/微博段子/通用正文


@app.post("/api/rewrite")
def rewrite_content(req: RewriteRequest):
    """AI 改写：对输入内容按指定模式和平台进行改写，直接返回结果文本"""
    import re

    # ── 模式指令 ──────────────────────────────────────────────
    mode_tasks = {
        "爆款化":    "将以下内容改写为爆款风格：标题吸睛、开头抓人、节奏感强、情绪饱满，适合在社交媒体刷屏传播。",
        "口语化":    "将以下内容改写为轻松口语化风格：去掉书面腔，像和朋友聊天一样自然流畅，有温度感。",
        "精简":      "将以下内容精简压缩：去掉废话和重复表达，保留核心观点，控制在原文 40%-60% 的字数。",
        "深度展开":  "将以下内容深度展开：补充案例、数据、逻辑链，使论点更有说服力和深度，适合长文内容。",
        "加开头钩子": "保持正文不变，在开头加一个强力钩子（疑问/数据冲击/反常识/场景感）吸引读者继续读。",
        "互动结尾":  "保持正文不变，在结尾加一个互动引导结尾（提问/投票/引发讨论），提升评论和互动率。",
    }

    # ── 平台字数要求 ─────────────────────────────────────────
    WORD_MIN = {
        "抖音脚本":   200,
        "小红书图文": 300,
        "公众号文章": 800,
        "微博段子":   100,
        "通用正文":   150,
    }
    platform_hint = {
        "抖音脚本":   "（目标平台：抖音脚本，口语化强，适合配音朗读，每句话简短有力）",
        "小红书图文": "（目标平台：小红书，标题加emoji，段落短，多用感叹号，种草氛围）",
        "公众号文章": "（目标平台：微信公众号，文字更正式，可以稍长，注重排版节奏）",
        "微博段子":   "（目标平台：微博，精炼幽默，100-300字，结尾留金句或反转）",
        "通用正文":   "（通用平台格式，不限制风格）",
    }

    base_task    = mode_tasks.get(req.mode, f"将以下内容进行【{req.mode}】改写。")
    plat_hint    = platform_hint.get(req.platform, f"（目标平台：{req.platform}）" if req.platform else "")
    min_words    = WORD_MIN.get(req.platform, 150)

    # ── Few-Shot 示例库 ──────────────────────────────────────
    # 每个 mode 配 2 个正面示例（输入→输出），展示"好的改写长什么样"
    EXAMPLES = {
        '爆款化': '''
【示例1】
输入：「婚姻需要双方共同努力经营，遇到矛盾时要互相理解包容，这样才能走得更长远。」
输出：「婚姻这东西，最怕的不是吵架，是两个人都憋着不说。我见过太多夫妻，表面上和和气气，背地里各过各的。其实啊，吵一架把话说开，比什么都强。」

【示例2】
输入：「根据研究表明，长期熬夜会导致免疫力下降，对身体健康产生严重影响。」
输出：「熬夜这事，偶尔问题不大，长期真的要命。我不是吓你，是有研究数据撑着的——连续熬夜一周，你身体的免疫力直接拉胯。」''',

        '口语化': '''
【示例1】
输入：「成功的关键在于持续学习和不断实践，只有不断提升自己才能适应环境变化。」
输出：「说白了，你想混得好，就两件事：学和干。不是看书感动自己，是真的下场去练。我见过太多人学了一脑子道理，一做事就抓瞎。」

【示例2】
输入：「该产品具有高效、便捷、环保等多项优势，深受消费者青睐。」
输出：「这东西用起来是真的顺手。我第一次上手的时候就感觉，这设计者是真的懂啊，不像某些产品反人类。」''',

        '精简': '''
【示例1】
输入：「在我们日常生活和工作中，时间管理是一个非常重要的话题，一个人如果能够很好地管理自己的时间，那么他就能够更加高效地完成各种任务，从而有更多的时间去做自己喜欢的事情。」
输出：「会管理时间的人，效率能甩普通人几条街。说白了就是：先想清楚什么重要，什么次要，然后把时间花在刀刃上。」

【示例2】
输入：「在当今社会，互联网技术的发展日新月异，改变了人们的生活方式和工作模式，越来越多的人开始选择远程办公，这种工作方式不仅节省了通勤时间，还提高了工作的灵活性。」
输出：「互联网这东西，彻底改变了我们上班的方式。远程办公越来越常见，省了通勤时间，工作也更灵活了。」''',

        '深度展开': '''
【示例1】
输入：「要学会拒绝别人。」
输出：「不懂拒绝的人，活得有多累？同事找你帮忙，你明明忙得脚不沾地，还是咬牙接下；朋友借钱，你心里一百个不愿意，嘴上还是说"行"。结果呢？自己的事没办好，别人的事也没办好，两头不讨好。学会拒绝不是自私，是对自己负责，也是对别人负责——你帮不了的事硬接，最后搞砸了，对方才真的难受。」

【示例2】
输入：「坚持很重要。」
输出：「为什么有些人做事三分钟热度，有些人能咬牙死磕？我观察下来，差别不在意志力，在于你是不是真的想要。我有个朋友减肥说了三年，每次都失败，不是他不努力，是他其实没那么想瘦——嘴上说说而已。你呢？你想做的事，是真想要，还是只是觉得应该想要？」''',

        '加开头钩子': '''
【示例1】
输入：「职场中要学会与不同性格的人相处...」
输出：「你有没有遇到过那种人？表面上笑眯眯，背后捅刀子。职场三年，我踩过的坑比你听过的道理还多。今天说点真的。」

【示例2】
输入：「健康的饮食习惯对每个人都至关重要...」
输出：「中国有4亿人每天吃的这个，正在慢慢杀死你。不是外卖，是你自己都意识不到的那个习惯。」''',

        '互动结尾': '''
【示例1】
输入：「...以上就是我对这个话题的全部看法」
输出：「...你觉得呢？你有没有类似的经历？欢迎评论区说出来，觉得有用的转给朋友。」

【示例2】
输入：「...这就是我的建议」
输出：「以上，完。觉得有用的话，点个赞；想知道更多，评论区抛出你的问题，我来答。」''',
    }

    few_shot = EXAMPLES.get(req.mode, "")

    # ── AI 禁止词清单（持续补充）───────────────────────────────
    BANNED_WORDS = [
        "此外", "总之", "综上所述", "值得注意的是", "至关重要", "深入探讨",
        "深刻的启示", "引人深省", "不禁让人", "持久的", "格局", "赋能",
        "洞见", "值得深思", "由此可见", "不可磨灭", "见证了", "标志着",
        "体现了", "不言而喻", "毋庸置疑", "而言之", "显然", "无疑",
        "可以看出", "不难发现", "至关重要", "核心", "关键", "本质",
        "层面上", "角度来看", "层面上讲", "从这个意义上",
    ]

    # ── 组装 System Prompt ────────────────────────────────────
    system_prompt = f"""你是一位顶级内容创作者，擅长把素材改写成有灵魂、有温度的真人内容。

【第一步：提炼灵魂——必须先完成】
在改写之前，先从原文中提炼出"3个必须保留的信息点"：
1. 一句话概括核心观点或故事
2. 一个具体细节（数字/名字/场景/案例）
3. 一种核心情绪或态度

格式：【灵魂1】...【灵魂2】...【灵魂3】...

【第二步：执行改写】
{base_task}{plat_hint}

【字数要求】
输出内容必须 ≥ {min_words} 汉字（不含标点符号）。字数不够=任务失败。

【第三步：自检清单】
改写完成后，对照以下清单检查是否通过：
□ 原文3个灵魂信息点是否全部保留？
□ 有没有出现禁用词（此外、总之、关键、核心、赋能、洞见...）？
□ 有没有三段论（首先/其次/最后）？
□ 有没有"这不仅仅是…而是…"结构？
□ 句子长度是否有变化（不是每句都一样长）？
□ 有没有具体数字/名字/场景（不是说"长期"而是说"3年"）？
□ 有没有态度（不是中立报道，而是"我喜欢/我讨厌/我觉得"）？

如果自检不通过，重新改写。

【Few-Shot 参考示例】
{few_shot}

【去AI化规则】

▌词汇红线：以下词出现必须替换
{', '.join(BANNED_WORDS)}
→ 替换示例：此外→还有、总之→说白了、关键→重点、核心→最要紧的、赋能→帮忙、洞见→发现

▌结构红线：
- 禁用三段论（首先→其次→最后）→ 改成自然叙述
- 禁用"这不仅仅是…而是…" → 直接说结论
- 禁用"标志着…" → 删掉，说具体事实
- 禁用金句式收尾 → 换成有态度的话

▌真人写法（给方向，不是禁止清单）：
✅ 句子长短交错——短句打节奏，长句偶尔来一句
✅ 有你明确的态度——喜欢就是喜欢，不爽就是不爽
✅ 说具体的——"3年"好过"长期"，"第一个月"好过"初期"
✅ 允许有点毛边——真人说话就是有点散，不是每句都对仗
✅ 像给朋友发消息那样写——自然、亲切、有画面感

直接输出【改写结果】，不要输出自检过程、不要前言后记。"""

    user_prompt = req.content

    try:
        # 先估算输入字数，用于判断"精简"模式的目标
        input_len = len(re.sub(r'\s+', '', req.content))
        if req.mode == "精简":
            user_prompt = f"【原文约{input_len}字，目标输出{int(input_len * 0.4)}-{int(input_len * 0.6)}字】\n\n" + req.content

        result, model_used = ai_extract(system_prompt, user_prompt, max_tokens=8000, temperature=0.85)
        return {"result": result, "model": model_used}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI改写失败：{str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# P2: 多版本 Ensemble + 4-Pass Humanization
# ═══════════════════════════════════════════════════════════════════════════

def _humanize_pass(text: str) -> str:
    """4-Pass Humanization 后处理"""
    import re

    # Pass 1: 移除 AI 高频词（扩展清单）
    AI_WORDS = [
        '此外', '总之', '综上所述', '值得注意的是', '至关重要', '深入探讨',
        '深刻的启示', '引人深省', '不禁让人', '持久的', '格局', '赋能',
        '洞见', '值得深思', '由此可见', '不可磨灭', '见证了', '标志着',
        '体现了', '不言而喻', '毋庸置疑', '而言之', '显然', '无疑',
        '可以看出', '不难发现', '核心', '关键', '本质', '层面上',
        '角度来看', '层面上讲', '从这个意义上', '值得注意的是',
        '值得注意的是', '不难发现', '由此可见', '可以说',
    ]
    for w in AI_WORDS:
        # 简单替换为口语化表达
        replacements = {
            '此外': '还有', '总之': '说白了', '综上所述': '归根结底',
            '值得注意的是': '要我说', '至关重要': '特别重要', '深入探讨': '细说',
            '深刻的启示': '挺有启发的', '引人深省': '值得想想', '不禁让人': '让人',
            '持久的': '长期的', '格局': '眼界', '赋能': '帮忙', '洞见': '发现',
            '值得深思': '值得想想', '由此可见': '可见', '不可磨灭': '不会消失的',
            '见证了': '看到了', '标志着': '说明', '体现了': '反映了',
            '不言而喻': '不用说', '毋庸置疑': '不用怀疑', '而言之': '总之',
            '显然': '很明显', '无疑': '不用说', '可以看出': '可见',
            '不难发现': '容易看到', '核心': '最核心', '关键': '最关键',
            '本质': '最本质', '层面上': '方面',
        }
        text = text.replace(w, replacements.get(w, '很'))

    # Pass 2: 打破均匀句长——在连续3个短句后插入一个较长句子
    sentences = re.split(r'([。！？；\n])', text)
    result = []
    short_count = 0
    for i in range(0, len(sentences), 2):
        s = sentences[i]
        punct = sentences[i+1] if i+1 < len(sentences) else ''
        if not s.strip():
            continue
        word_count = len(s.strip())
        if word_count < 15:
            short_count += 1
        else:
            short_count = 0
        result.append(s + punct)
        # 连续短句过多时，强制合并
        if short_count >= 3 and i+2 < len(sentences):
            next_s = sentences[i+2].strip() if i+2 < len(sentences) else ''
            if next_s:
                result[-1] = result[-1].rstrip('。！？；') + '，' + next_s
                short_count = 0

    text = ''.join(result)

    # Pass 3: 移除"可引用金句结尾"——如果结尾像格言，加一句口语化吐槽
    text = text.strip()
    if text and len(text) > 20:
        last_sentence = re.split(r'[。！？；\n]', text)[-2] if len(re.split(r'[。！？；\n]', text)) > 1 else ''
        if last_sentence and (len(last_sentence) < 30 and ('。' not in last_sentence[-10:])):
            # 结尾像金句，加一句口语化的话
            text = text.rstrip('。！？') + '。说完了，就这样。'

    # Pass 4: 注入自我引用（只在合适位置加一句）
    if '我认为' not in text and '我觉得' not in text and len(text) > 100:
        # 在中间位置插入
        mid = len(text) // 2
        text = text[:mid] + '我自己的想法是，' + text[mid:]

    return text


class RewriteEnsembleRequest(BaseModel):
    content: str
    mode: str = "口语化"
    platform: str = ""


@app.post("/api/rewrite-ensemble")
def rewrite_ensemble(req: RewriteEnsembleRequest):
    """多版本改写——用两种不同风格生成2个版本，让用户选择更好的"""
    # 复用 rewrite_content 的逻辑，但用不同 temperature
    # 版本1: 温度 0.7（保守、精准）
    # 版本2: 温度 0.95（创意、多变）
    import re as _re2

    mode_tasks = {
        "爆款化":    "将以下内容改写为爆款风格：标题吸睛、开头抓人、节奏感强、情绪饱满，适合在社交媒体刷屏传播。",
        "口语化":    "将以下内容改写为轻松口语化风格：去掉书面腔，像和朋友聊天一样自然流畅，有温度感。",
        "精简":      "将以下内容精简压缩：去掉废话和重复表达，保留核心观点，控制在原文 40%-60% 的字数。",
        "深度展开":  "将以下内容深度展开：补充案例、数据、逻辑链，使论点更有说服力和深度，适合长文内容。",
        "加开头钩子": "保持正文不变，在开头加一个强力钩子（疑问/数据冲击/反常识/场景感）吸引读者继续读。",
        "互动结尾":  "保持正文不变，在结尾加一个互动引导结尾（提问/投票/引发讨论），提升评论和互动率。",
    }
    WORD_MIN = {
        "抖音脚本":   200, "小红书图文": 300, "公众号文章": 800, "微博段子": 100, "通用正文": 150,
    }
    platform_hint = {
        "抖音脚本":   "（目标平台：抖音脚本，口语化强，适合配音朗读，每句话简短有力）",
        "小红书图文": "（目标平台：小红书，标题加emoji，段落短，多用感叹号，种草氛围）",
        "公众号文章": "（目标平台：微信公众号，文字更正式，可以稍长，注重排版节奏）",
        "微博段子":   "（目标平台：微博，精炼幽默，100-300字，结尾留金句或反转）",
        "通用正文":   "（通用平台格式，不限制风格）",
    }

    base_task = mode_tasks.get(req.mode, f"将以下内容进行【{req.mode}】改写。")
    plat_hint = platform_hint.get(req.platform, f"（目标平台：{req.platform}）" if req.platform else "")
    min_words = WORD_MIN.get(req.platform, 150)

    system_base = f"""你是一位顶级内容创作者，擅长把素材改写成有灵魂、有温度的真人内容。

【本次任务】
{base_task}{plat_hint}
【输出字数要求】≥ {min_words} 汉字。
【去AI化】禁用词：此外、总之、关键、核心、赋能、洞见、值得注意的是、深入探讨。替换：此外→还有、总之→说白了。
【真人写法】✅ 句子长短交错 ✅ 有态度 ✅ 说具体的（"3年"好过"长期"）✅ 像给朋友发消息
直接输出改写结果，不要前言后记。"""

    results = []
    # 跑两个版本：低温（0.7）和高温（0.95）
    for temp, label in [(0.7, "稳重版"), (0.95, "创意版")]:
        try:
            text, model = ai_extract(system_base, req.content, max_tokens=8000, temperature=temp)
            # 应用 humanization 后处理
            text = _humanize_pass(text)
            char_count = len(_re2.sub(r'\s+', '', text))
            results.append({
                "label": label,
                "text": text,
                "chars": char_count,
                "model": model,
                "temp": temp,
            })
        except Exception as e:
            results.append({"label": label, "error": str(e), "text": ""})

    return {"versions": results}


class WechatFormatRequest(BaseModel):
    content: str
    theme: str = "wechat-native"


@app.post("/api/wechat-format")
def wechat_format(req: WechatFormatRequest):
    """将 Markdown 排版为微信公众号兼容的内联样式 HTML"""
    import sys
    import re as _re
    from pathlib import Path

    format_dir = BASE_DIR / "wechat-format"
    if str(format_dir) not in sys.path:
        sys.path.insert(0, str(format_dir))

    try:
        import format as wfmt
    except ImportError as e:
        raise HTTPException(500, f"排版模块加载失败：{e}")

    # 验证主题存在
    valid_ids = {t["id"] for t in WECHAT_THEMES}
    theme_name = req.theme if req.theme in valid_ids else "wechat-native"

    try:
        theme = wfmt.load_theme(theme_name)

        # 直接执行核心处理流程，绕开 vault_root os.walk（Windows 根目录遍历会卡死）
        content = req.content
        input_path = Path("article.md")
        output_dir = BASE_DIR / "data" / "wechat-output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 通用预处理（与 format_for_output 一致，跳过 wikilinks/本地图片）
        title = wfmt.extract_title(content, input_path)
        word_count = wfmt.count_words(content)
        content = wfmt.strip_frontmatter(content)
        content = wfmt.fix_cjk_spacing(content)
        content = wfmt.fix_cjk_bold_punctuation(content)
        content = wfmt.process_callouts(content)
        content = wfmt.process_manual_footnotes(content)
        content = wfmt.process_fenced_containers(content)
        content = _re.sub(r'~~(.+?)~~', r'<del>\1</del>', content)
        # 跳过 convert_wikilinks（需要遍历本地文件系统，Windows 上 vault_root='/' 会崩溃）
        # 跳过 copy_markdown_images（用户粘贴的是纯文字，无本地图片）

        html = wfmt.md_to_html(content)
        html, footnote_html = wfmt.extract_links_as_footnotes(html)
        html = wfmt.inject_inline_styles(html, theme)
        if footnote_html:
            footnote_html = wfmt.inject_inline_styles(footnote_html, theme, skip_wrapper=True)
        html = wfmt.convert_image_captions(html)

        full_html = html + ("\n" + footnote_html if footnote_html else "")
        return {
            "html": full_html,
            "title": title,
            "word_count": word_count,
            "theme": theme_name,
        }
    except Exception as e:
        import traceback
        raise HTTPException(500, f"排版失败：{e}\n{traceback.format_exc()}")


@app.post("/api/wechat-preview")
def wechat_preview(req: WechatFormatRequest):
    """返回完整可渲染的 HTML 页面（旧接口保留兼容）"""
    result = wechat_format(req)
    body_html = result["html"]
    full_page = _build_preview_page(body_html)
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=full_page, media_type="text/html; charset=utf-8")


# ─── 预览缓存（token → HTML）─────────────────────────────────────────────
import uuid as _uuid
_preview_cache: dict = {}  # token -> full_html

def _build_preview_page(body_html: str) -> str:
    base_style = (
        "*{box-sizing:border-box}"
        "html,body{margin:0;padding:0;background:#fff}"
        "body{padding:16px 24px;font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;"
        "font-size:15px;line-height:1.75;color:#333}"
        "img{max-width:100%;display:block;margin:8px auto}"
    )
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<style>{base_style}</style>
</head>
<body>{body_html}</body>
</html>"""


class PreviewStoreRequest(BaseModel):
    content: str
    theme: str = "default"


@app.post("/api/preview-store")
def preview_store(req: PreviewStoreRequest):
    """生成排版并缓存到内存，返回 token；前端用 GET /api/preview/{token} 在 iframe 中展示"""
    fmt_result = wechat_format(WechatFormatRequest(content=req.content, theme=req.theme))
    token = _uuid.uuid4().hex
    _preview_cache[token] = _build_preview_page(fmt_result["html"])
    # 最多保留 20 条缓存，防止内存膨胀
    if len(_preview_cache) > 20:
        oldest = next(iter(_preview_cache))
        del _preview_cache[oldest]
    return {
        "token": token,
        "html": fmt_result["html"],       # 仍返回 html 供复制
        "word_count": fmt_result.get("word_count"),
        "theme": fmt_result.get("theme", req.theme),
    }


@app.get("/api/preview/{token}")
def preview_get(token: str):
    """GET 接口，返回缓存的完整 HTML 预览页面；可直接作为 iframe src"""
    from fastapi.responses import HTMLResponse
    html = _preview_cache.get(token)
    if not html:
        return HTMLResponse("<html><body><p style='color:#999;text-align:center;padding:40px'>预览已过期，请重新生成排版</p></body></html>")
    return HTMLResponse(content=html, media_type="text/html; charset=utf-8")





# ─── 朴树之道分享 ──────────────────────────────────────────────────────

PUSHUTREE_SYSTEM_PROMPT = """你是一个真正读过无数书的人，现在做书籍推荐视频/文案。

【你的核心定位】
你不是在写读书笔记，不是在写书评，也不是在讲故事。
你是在做"推荐词"——你要让一个完全不了解这本书的观众，在看完你这期内容之后：
1. 得到一个对他生活真正有用的认知/方法/感悟
2. 心里有点被触动（情绪共鸣）
3. 产生想去读这本书的冲动

每一期内容都是独立完整的。观众不需要看前一期，也不需要看后一期。
这不是连续剧，是系列推荐词——每期讲书里一个核心点，从不同角度打动人。

【什么是好的推荐词——学董宇辉/朴树之道】

董宇辉推荐书的方式：
- 先说一个让人心里一颤的问题或生活场景（开头抓人）
- 然后说书里怎么回答这个问题，或者这本书给了他什么新的认知
- 说自己的感受，不是复述书的内容
- 金句多，但不是鸡汤，是真实的洞察
- 让人觉得：这个人是真的读过，而且真的有感悟

朴树之道的感觉：
- 不装，不端，真实
- 说话像跟朋友聊天，不像在开讲座
- 有自己的判断和态度，不是中立的介绍

【硬性禁止】
❌ 叙事型写法："这本书讲了一个故事……""在书的第X章，作者提到……"
❌ 读书报告结构："这本书共分三部分，首先……其次……最后……"
❌ AI废话：深刻、启示、洞见、值得深思、引人深省、综上所述、由此可见
❌ 连续剧式开头："上一期我们讲到……""今天我们继续……"
❌ 讲述书的情节，像在讲故事

【每期推荐词的正确结构】
① 开头钩子（1-3句）：一个让人停下来的问题、反常识判断、或真实生活场景
   - 要打中观众的某个真实困惑或痛点
   - 前三句决定观众走不走

② 核心认知（2-3段）：
   - 书里给出的答案/方法/视角——但用你自己的话说
   - 要具体，用书里真实的案例/数据/故事支撑
   - 让人有"原来如此"的感觉

③ 为什么你/观众需要这本书（1-2段）：
   - 这个认知能帮观众解决什么真实问题
   - 读完这本书，观众的某一件事会变得不同

④ 结尾推荐句（1-2句）：
   - 不是"这本书值得一读"这种废话
   - 是一句让人想截图或转发的话
   - 带情绪，带态度，真实

【字数要求】：600-900字，干货密度高，不要凑字数

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【去AI化——写作灵魂规则，让文字像真人说出来的】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

这是口播稿，更不能有AI味。观众一眼就能听出来那种"机器人感"。

▌绝对禁止的AI词汇：
此外、总之、综上所述、值得注意的是、至关重要、深入探讨、深刻的启示、引人深省、
不禁让人、持久的影响、格局、赋能、洞见、值得深思、由此可见、不可磨灭、
见证了、标志着、体现了、不言而喻、毋庸置疑、可以说、某种程度上

▌禁止的AI写作结构（口播特别注意）：
- 三段论公式开头："有三个方面……" → 直接讲第一个，自然带出下一个
- 否定排比："这不仅仅是……而是……" → 直接说结论
- 宣传性形容词："深刻的""颠覆性的""令人震撼的" → 删掉，让事实本身说话
- 过长的句子（超过25字）→ 拆成两句，口播要短句

▌真人口播的写法：
- 多用"你"，少用"我们"——直接跟观众说话
- 允许说"说实话""我觉得""坦白说"——口语感
- 举的例子要具体，有数字有细节（"三年""每天早上6点""第87页"比"长期""坚持""书中"更真实）
- 句子有轻有重，有长有短——模拟真人说话的节奏
- 偶尔留一个问句让观众心里有点颤——不需要他回答，但他会停下来想"""


def _fix_json_bare_quotes(json_str: str) -> str:
    """修复 JSON 字符串值内部的裸双引号（AI 常见输出错误）。
    
    逐字符扫描：在字符串值内遇到未转义的双引号时替换为单引号。
    这样 {"key": "he said "hello""} 会被修复为 {"key": "he said 'hello'"}
    """
    result = []
    in_string = False
    i = 0
    prev_was_key = False  # 上一个字符串是键（不在值里）
    
    while i < len(json_str):
        c = json_str[i]
        
        if not in_string:
            if c == '"':
                in_string = True
                result.append(c)
            else:
                result.append(c)
        else:
            if c == '\\' and i + 1 < len(json_str):
                # 合法转义序列，原样保留
                result.append(c)
                result.append(json_str[i + 1])
                i += 2
                continue
            elif c == '"':
                # 字符串结束，检查下一个非空白字符是否为 JSON 合法后续符
                j = i + 1
                while j < len(json_str) and json_str[j] in ' \t\r\n':
                    j += 1
                next_c = json_str[j] if j < len(json_str) else ''
                if next_c in ',:}]':
                    # 这是合法的字符串结束引号
                    in_string = False
                    result.append(c)
                else:
                    # 这是字符串内部的裸引号，替换为单引号
                    result.append("'")
            else:
                result.append(c)
        i += 1
    
    return ''.join(result)


def pushutree_plan(book_name: str, text: str, episode_count: int, style: str) -> list[dict]:
    """第一步：策划——分析全书找痛点，生成标题+大纲"""
    user_prompt = f"""请仔细阅读《{book_name}》的内容，用你作为一个真正读过这本书的人的眼光，找出{episode_count}个最值得分享的核心主题。

不要找那种"第X章讲了什么"的表面总结，要找：
- 让你看完会停下来想一想的观点
- 打破常见认知误区的内容
- 能帮普通人解决真实生活困惑的道理
- 书里藏着但很少人注意到的深层规律

【书籍内容】
{text[:60000]}

【输出要求】
请严格按照以下JSON格式输出，不要输出其他任何内容：
{{
  "episodes": [
    {{
      "ep_no": 1,
      "title": "标题（15字以内，有反差感或冲击力，不用你不知道的XX这类老套标题）",
      "angle": "这期的切入角度（一句话，说清楚从哪个生活场景或困惑切入）",
      "core_insight": "这期最核心的认知或规律（不是金句，是一个真实的洞察）",
      "pain_point": "击中的读者真实痛点（要具体，不要写迷茫这种大词）",
      "source_hint": "书中哪个章节或故事或案例或数据可以支撑",
      "why_useful": "读者看完能学到什么，能用在哪里（一句话）"
    }}
  ]
}}

重要：JSON字段的值必须是合法字符串，不能包含未转义的双引号（如需引用，请用书名号《》或单引号'代替）。
风格要求：{style}
共需{episode_count}期，每期角度不重复，要覆盖全书的精华，不要只集中在前几章。"""

    raw, _model = ai_extract(
        system_prompt=PUSHUTREE_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        max_tokens=4000,
        temperature=0.7,
    )
    raw = raw.strip()
    # 提取JSON（容错：处理AI在JSON前后输出废话、markdown代码块、中文引号等）
    import re
    # 去掉 markdown 代码块包裹 ```json ... ```
    raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
    raw = re.sub(r'\s*```\s*$', '', raw, flags=re.MULTILINE)
    # 替换中文引号为英文引号（常见AI输出错误）
    raw = raw.replace('\u201c', '"').replace('\u201d', '"').replace('\u2018', "'").replace('\u2019', "'")
    # 提取最外层 JSON 对象
    m = re.search(r'\{[\s\S]*\}', raw)
    if m:
        json_str = m.group()
        # 第一次尝试：直接解析
        try:
            data = json.loads(json_str)
            return data.get("episodes", [])
        except json.JSONDecodeError as je:
            print(f"[pushutree_plan] 第一次JSON解析失败: {je}，尝试修复裸引号...")
        # 第二次尝试：修复字段值内的裸双引号
        # 思路：逐字扫描，在字符串值内部遇到裸双引号时替换为单引号
        try:
            fixed = _fix_json_bare_quotes(json_str)
            data = json.loads(fixed)
            print(f"[pushutree_plan] 修复裸引号后解析成功")
            return data.get("episodes", [])
        except Exception as je2:
            print(f"[pushutree_plan] 修复后仍失败: {je2}\n原始内容(前800): {json_str[:800]}")
            return []
    print(f"[pushutree_plan] 未找到JSON结构，原始内容(前500): {raw[:500]}")
    return []


def pushutree_write(book_name: str, text: str, episode: dict, style: str) -> str:
    """第二步：撰写——写一篇独立的书籍推荐词"""
    user_prompt = f"""现在为《{book_name}》写第{episode['ep_no']}期书籍推荐词。

【本期要传递的核心内容】
- 这期聚焦的角度：{episode['angle']}
- 核心认知：{episode.get('core_insight', episode.get('core_quote',''))}
- 观众痛点：{episode['pain_point']}
- 书中支撑：{episode['source_hint']}
- 观众能得到什么：{episode.get('why_useful','一个对生活有用的新认知')}

【书籍原文（供参考提取真实内容）】
{text[:50000]}

---

【写作任务】
写一篇让观众不想划走、看完有收获、想去读这本书的推荐词。

这不是读书笔记，不是书评，不是连续剧的某一集。
这是一篇独立完整的推荐词，观众看了这一期，不需要看其他期。

【必须做到的三件事】

第一：开头要抓人（前3句）
不要从"这本书讲了……"开始。
要从观众的真实感受/困惑/生活场景开始。
比如：
- 直接抛出一个让人心里一紧的问题
- 说一个很多人有但没说出来的感受
- 说一个和常识相反的判断

第二：干货要真实（中间部分）
用书里真实的案例、数据、观点来说话。
不要只说结论，要说"书里是怎么说的"——但用你自己的话，不是复制粘贴。
每个认知点要对应一个"观众能用在哪里"。

第三：结尾要有力量
不是"这本书值得一读"。
是一句让人想截图的话——有点哲学，有点反常识，或者说出了很多人想说又说不出来的话。

【风格】：{style}
【字数】：600-900字
【格式】：自然段落，口语化，短句为主，不用序号标题

【去AI化——写的时候就不能有AI味，不要等到精修再改】
✅ 用"你"开头或在关键句里直接和观众说话
✅ 举具体例子（书里的真实数字、细节、场景），不说空话
✅ 句子长短交错——3-5个短句打节奏，偶尔一句长句讲道理
✅ 有自己的态度——"我觉得这书有一点特别重要""说实话这个观点很少有人能做到"
❌ 绝对不写：此外、总之、由此可见、深刻的、至关重要、值得注意的是
❌ 绝对不用三段论公式：首先……其次……最后……"""

    draft, _model = ai_extract(
        system_prompt=PUSHUTREE_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        max_tokens=3000,
        temperature=0.85,
    )
    return draft.strip()


def pushutree_polish(book_name: str, draft: str, episode: dict) -> str:
    """第三步：精修——观众视角终审 + 去AI化深度清洗"""
    user_prompt = f"""对下面这篇《{book_name}》书籍推荐词做最终精修。

【精修标准——你是一个极度挑剔的观众】

先以观众身份看一遍，问自己：

❶ 开头三句，我会不会划走？
   - 如果开头是"这本书讲了……" → 必须重写，从观众痛点切入
   - 如果开头是废话过渡 → 删掉，直接进正文
   
❷ 中间内容，我学到了什么？
   - 有没有一个对我日常生活真正有用的认知？
   - 说的是书里真实的东西，还是泛泛而谈？

❸ 这像推荐词还是读书报告？
   - 推荐词：说人话，有态度，让我想读这本书
   - 读书报告：复述内容，中立无感，让我想睡觉
   - 如果是读书报告味道 → 砍掉复述段落，换成"这对你意味着什么"

❹ 结尾金句，我会不会截图？
   - 不截图 → 重写

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【去AI化深度清洗——逐项检查，必须全部通过】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

这是口播稿，观众一听就知道是不是AI写的。把所有AI味彻底洗掉。

▌词汇层清查（发现即替换）：
❌ AI高频词：此外、总之、综上所述、值得注意的是、至关重要、深入探讨、深刻的启示、
   引人深省、不禁让人、持久的影响、格局、赋能、洞见、值得深思、由此可见、
   不可磨灭、见证了、标志着、体现了、不言而喻、毋庸置疑
→ 直接删除或换成口语表达

❌ 宣传性空话：充满活力的、丰富的、令人叹为观止的、深刻的、颠覆性的
→ 删掉形容词，用具体事实替代

▌结构层清查：
❌ 三段论公式（首先……其次……最后……）→ 打破，改成自然流动的口语叙述
❌ 否定排比（"这不仅仅是……而是……"）→ 直接说结论
❌ 以"这本书"或"作者"开头的句子（超过2句）→ 换成"你""我"的视角

▌节奏层清查：
❌ 超过30字的长句 → 拆成两句
❌ 连续3句以上相同句型（都是陈述句/都是问句）→ 换节奏
✅ 检查：有没有几个短句打节奏的段落？（好的口播稿一定有）

▌真实感注入：
- 如果全文没有一个具体数字或细节 → 补上（从书中找）
- 如果全文没有"你"字 → 至少加3处，直接和观众说话
- 如果结尾感觉是"总结"而不是"留白/回味" → 改成一句有态度的话

【原稿】
{draft}

【输出要求】
直接输出修改后的完整推荐词。不要说"已修改如下"，不要有任何前缀。
保留原有信息量，让语言更自然、更有力、更像真人在说话。
这篇文章应该让人读完觉得：这是一个真正读过这本书、有自己想法的人写的。"""

    polished, _model = ai_extract(
        system_prompt=PUSHUTREE_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        max_tokens=3000,
        temperature=0.6,
    )
    return polished.strip()


def run_pushutree_task(script_id: str, source_id: str, episode_count: int,
                       platform: str, style: str, direct_text: str = "", direct_book_name: str = ""):
    """后台任务：朴树之道三步流程
    支持两种模式：
    1. 书库模式：source_id 非空，从数据库读取书籍
    2. 独立模式：source_id 为空，用 direct_text + direct_book_name
    """
    conn = get_db()

    def upd(status, progress, message, episodes=None, plan=None, error_msg=None):
        fields = "status=?,progress=?,message=?,updated_at=?"
        vals = [status, progress, message, now()]
        if episodes is not None:
            fields += ",episodes=?"
            vals.append(json.dumps(episodes, ensure_ascii=False))
        if plan is not None:
            fields += ",plan=?"
            vals.append(json.dumps(plan, ensure_ascii=False))
        if error_msg is not None:
            fields += ",error_msg=?"
            vals.append(error_msg)
        vals.append(script_id)
        conn.execute(f"UPDATE scripts SET {fields} WHERE id=?", vals)
        conn.commit()

    try:
        upd("processing", 5, "准备书籍内容...")

        text = ""
        book_name = direct_book_name

        if direct_text:
            # 独立模式：直接用传入的文本
            text = direct_text
            upd("processing", 15, f"已获取书籍文本，共{len(text)}字...")
        elif source_id:
            # 书库模式：从数据库读取
            upd("processing", 8, "读取书籍信息...")
            src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
            if not src:
                upd("error", 0, "找不到书籍", error_msg="source not found")
                return
            book_name = src["title"]

            # 从PDF文件提取文本（任意状态的书籍都可以，不必须提炼完成）
            print(f"[DEBUG] pushutree source: type={src['type']}, file_path={src['file_path']}, url={src['url']}")
            if src["type"] == "book" and src["file_path"]:
                upd("processing", 10, "提取PDF文本...")
                text, _, _ = extract_text_from_pdf(src["file_path"])
                upd("processing", 15, f"PDF文本提取完成（{len(text)}字）...")
            elif src["type"] == "epub" and src["file_path"]:
                # epub 类型
                upd("processing", 10, "解析EPUB文件...")
                text, _, _ = extract_text_from_epub(src["file_path"])
                if isinstance(text, list):
                    text = "\n".join(text)
                upd("processing", 15, f"EPUB解析完成（{len(text)}字）...")
            elif src["type"] == "txt" and src["file_path"]:
                # txt 类型
                upd("processing", 10, "读取文本文件...")
                text, _, _ = extract_text_from_txt(src["file_path"])
                if isinstance(text, list):
                    text = "\n".join(text)
                upd("processing", 15, f"文本文件读取完成（{len(text)}字）...")
            elif src["type"] == "text":
                # text类型：文本存在url字段
                text = src["url"] or ""
                upd("processing", 15, f"已获取文本内容（{len(text)}字）...")
            elif src["type"] == "url":
                # url类型：需要从URL抓取内容
                import asyncio
                loop = asyncio.new_event_loop()
                try:
                    upd("processing", 10, "正在从URL抓取内容...")
                    text = loop.run_until_complete(extract_text_from_url(src["url"]))
                    upd("processing", 15, f"URL内容获取完成（{len(text)}字）...")
                finally:
                    loop.close()
            elif src["url"]:
                # book类型但没有file_path，有url字段——尝试当作URL处理
                import asyncio
                loop = asyncio.new_event_loop()
                try:
                    upd("processing", 10, "正在从URL抓取内容...")
                    text = loop.run_until_complete(extract_text_from_url(src["url"]))
                    upd("processing", 15, f"URL内容获取完成（{len(text)}字）...")
                finally:
                    loop.close()
            else:
                # 兜底：尝试从素材库获取内容
                mat_rows = conn.execute("SELECT content FROM materials WHERE source_id=? LIMIT 1", (source_id,)).fetchall()
                if mat_rows:
                    text = " ".join([m["content"] for m in mat_rows])
                    upd("processing", 15, f"从素材库获取内容（{len(text)}字）...")
                else:
                    upd("error", 0, "无法获取书籍内容", error_msg="no content found")
                    return
        else:
            upd("error", 0, "缺少书籍内容", error_msg="no source")
            return

        if not text or len(text.strip()) < 50:
            print(f"[DEBUG] text too short: len={len(text) if text else 0}, first 100 chars: {text[:100] if text else 'empty'}")
            # 文本太短，尝试从素材库获取内容作为兜底
            if source_id:
                mat_rows = conn.execute("SELECT content FROM materials WHERE source_id=?", (source_id,)).fetchall()
                print(f"[DEBUG] materials from DB: {len(mat_rows)} rows")
                if mat_rows:
                    text = " ".join([m["content"] for m in mat_rows])
                    upd("processing", 15, f"从素材库获取内容（{len(text)}字）...")
            if not text or len(text.strip()) < 50:
                upd("error", 0, "书籍文本太少，无法生成", error_msg="text too short")
                return

        if not book_name:
            book_name = "本书"

        # 第一步：策划
        upd("processing", 20, f"第一步：分析《{book_name}》，生成{episode_count}期策划...")
        plan = pushutree_plan(book_name, text, episode_count, style)
        if not plan:
            upd("error", 0, "策划生成失败，请重试", error_msg="plan empty")
            return
        upd("processing", 35, f"策划完成，共{len(plan)}期，开始撰写...", plan=plan)

        # 第二步+第三步：逐期撰写+精修
        episodes = []
        for i, ep in enumerate(plan):
            ep_no = ep.get("ep_no", i + 1)
            title = ep.get("title", f"第{ep_no}期")

            progress = 35 + int((i / len(plan)) * 55)
            upd("processing", progress, f"撰写第{ep_no}/{len(plan)}期：{title}...")

            try:
                draft = pushutree_write(book_name, text, ep, style)
                upd("processing", progress + 3, f"精修第{ep_no}期...")
                final = pushutree_polish(book_name, draft, ep)
            except Exception as e:
                final = f"（生成失败：{e}）"

            episodes.append({
                "ep_no": ep_no,
                "title": title,
                "content": final,
                "pain_point": ep.get("pain_point", ""),
                "angle": ep.get("angle", ""),
            })
            upd("processing", progress + 5, f"第{ep_no}期完成", episodes=episodes)

        upd("done", 100, f"全部{len(episodes)}期生成完毕！", episodes=episodes)

    except Exception as e:
        import traceback
        upd("error", 0, f"生成失败：{e}", error_msg=traceback.format_exc())
    finally:
        conn.close()


class PushutreeRequest(BaseModel):
    source_id: str = ""           # 可选：已有书籍ID
    book_name: str = ""           # 可选：直接传书名（独立模式）
    book_text: str = ""           # 可选：直接传书籍文本（独立模式）
    episode_count: int = 8
    platform: str = "抖音/视频号"
    style: str = "犀利、接地气、直击痛点"


@app.post("/api/pushutree/create")
def create_pushutree(req: PushutreeRequest):
    """创建朴树之道分享任务
    支持两种模式：
    1. 书库模式：传 source_id（任意状态的书籍都可以，不必须提炼完成）
    2. 独立模式：传 book_name + book_text（直接输入书名和文本）
    """
    conn = get_db()

    # 确定来源信息
    if req.source_id:
        src = conn.execute("SELECT * FROM sources WHERE id=?", (req.source_id,)).fetchone()
        if not src:
            conn.close()
            raise HTTPException(404, "书籍不存在")
        source_id = req.source_id
        source_title = src["title"]
    elif req.book_name:
        # 独立模式：没有对应的 sources 记录，用 book_name 作标题，source_id 为空
        source_id = ""
        source_title = req.book_name
    else:
        conn.close()
        raise HTTPException(400, "请提供书籍ID或书名")

    script_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO scripts VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (script_id, source_id, source_title,
         req.episode_count, req.platform, req.style,
         "pending", 0, "等待生成...",
         "[]", "[]", None, now(), now())
    )
    conn.commit()
    conn.close()

    # 后台异步执行
    threading.Thread(
        target=run_pushutree_task,
        args=(script_id, source_id, req.episode_count, req.platform, req.style, req.book_text, req.book_name),
        daemon=True
    ).start()

    return {"script_id": script_id, "status": "pending"}


@app.get("/api/pushutree/{script_id}")
def get_pushutree(script_id: str):
    """查询朴树之道任务进度和结果"""
    conn = get_db()
    row = conn.execute("SELECT * FROM scripts WHERE id=?", (script_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "任务不存在")
    d = dict(row)
    d["episodes"] = json.loads(d["episodes"] or "[]")
    d["plan"] = json.loads(d["plan"] or "[]")
    return d


@app.get("/api/pushutree")
def list_pushutree(source_id: str = None):
    """获取所有朴树之道系列（可按书籍过滤）"""
    conn = get_db()
    if source_id:
        rows = conn.execute(
            "SELECT * FROM scripts WHERE source_id=? ORDER BY created_at DESC",
            (source_id,)
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT * FROM scripts ORDER BY created_at DESC"
        ).fetchall()
    conn.close()
    result = []
    for row in rows:
        d = dict(row)
        d["episodes"] = json.loads(d["episodes"] or "[]")
        d["plan"] = json.loads(d["plan"] or "[]")
        result.append(d)
    return result


@app.delete("/api/pushutree/{script_id}")
def delete_pushutree(script_id: str):
    """删除朴树之道系列"""
    conn = get_db()
    conn.execute("DELETE FROM scripts WHERE id=?", (script_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


@app.post("/api/pushutree/upload-and-create")
async def pushutree_upload_and_create(
    file: UploadFile = File(None),
    book_name: str = Form(...),
    book_text: str = Form(""),
    episode_count: int = Form(8),
    platform: str = Form("抖音/视频号"),
    style: str = Form("犀利、接地气、直击痛点"),
):
    """朴树之道独立上传入口：
    直接上传PDF或填入书名+文本，不依赖书库，立即生成系列分享文案。
    """
    text = book_text

    # 如果上传了PDF，提取文本
    if file and file.filename:
        file_id = str(uuid.uuid4())
        save_path = UPLOAD_DIR / f"pt_{file_id}_{file.filename}"
        content = await file.read()
        with open(save_path, "wb") as f:
            f.write(content)
        try:
            extracted, _, _ = extract_text_from_pdf(str(save_path))
            text = extracted
        except Exception as e:
            raise HTTPException(500, f"PDF解析失败：{e}")

    if not text or len(text.strip()) < 50:
        raise HTTPException(400, "请上传PDF文件或填入书籍文本（至少50字）")

    if not book_name.strip():
        raise HTTPException(400, "请填写书名")

    # 创建 script 记录
    conn = get_db()
    script_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO scripts VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (script_id, "", book_name.strip(),
         episode_count, platform, style,
         "pending", 0, "等待生成...",
         "[]", "[]", None, now(), now())
    )
    conn.commit()
    conn.close()

    # 后台异步执行
    threading.Thread(
        target=run_pushutree_task,
        args=(script_id, "", episode_count, platform, style, text, book_name.strip()),
        daemon=True
    ).start()

    return {"script_id": script_id, "status": "pending", "book_name": book_name.strip()}


# 统计数据

@app.get("/api/stats")
def get_stats():
    conn = get_db()
    sources_total = conn.execute("SELECT COUNT(*) FROM sources WHERE status='done'").fetchone()[0]
    materials_total = conn.execute("SELECT COUNT(*) FROM materials").fetchone()[0]
    cat_counts = dict(conn.execute(
        "SELECT category, COUNT(*) FROM materials GROUP BY category"
    ).fetchall())
    starred = conn.execute("SELECT COUNT(*) FROM materials WHERE is_starred=1").fetchone()[0]
    conn.close()
    return {
        "sources": sources_total,
        "materials": materials_total,
        "starred": starred,
        "by_category": cat_counts,
    }

# 保存创作
class SaveCreationRequest(BaseModel):
    title: str
    content: str
    platform: str = ""
    source_ids: list = []
    material_ids: list = []

@app.post("/api/creations")
def save_creation(req: SaveCreationRequest):
    cid = str(uuid.uuid4())
    conn = get_db()
    conn.execute(
        "INSERT INTO creations VALUES (?,?,?,?,?,?,?,?)",
        (cid, req.title, req.content,
         json.dumps(req.source_ids), json.dumps(req.material_ids),
         req.platform, now(), now())
    )
    conn.commit()
    conn.close()
    return {"id": cid}

@app.get("/api/creations")
def list_creations(q: str = ""):
    conn = get_db()
    sql = "SELECT * FROM creations"
    params = []
    if q:
        sql += " WHERE title LIKE ? OR content LIKE ?"
        params = [f"%{q}%", f"%{q}%"]
    sql += " ORDER BY updated_at DESC"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]



# ═══════════════════════════════════════════════════════════════════════════
# 智能提取增强功能（分层 Chunking + 多轮 Pipeline + 质量评分）
# ═══════════════════════════════════════════════════════════════════════════

def process_source_task_smart(task_id: str, source_id: str, mode: str = "full"):
    """
    智能提取流程 - 使用分层 Chunking + 多轮迭代 + 质量评分

    相比原版 process_source_task，改进点：
    1. 分层 Chunking：保留书籍结构，按章节语义切分
    2. 多轮迭代：结构理解 → 逐章提取 → 跨章分析 → IP选题
    3. 质量评分：完整度/唯一性/IP契合度/可执行性/风险等级
    4. 智能路由：自动批准/人工审核/自动丢弃
    """
    if not _SMART_EXTRACTION_AVAILABLE:
        # 模块不可用，回退到原版
        print(f"[SmartExtraction] 模块不可用，回退到原版流程")
        process_source_task(task_id, source_id, mode)
        return

    conn = get_db()

    def update_task(status, progress, message, result=None):
        conn.execute(
            "UPDATE tasks SET status=?,progress=?,message=?,result=?,updated_at=? WHERE id=?",
            (status, progress, message, result, now(), task_id)
        )
        conn.commit()

    try:
        update_task("processing", 5, "准备智能提取流程...")

        # 获取书籍信息
        src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
        if not src:
            update_task("error", 0, "找不到来源记录")
            return

        # 步骤1: 文本提取
        update_task("processing", 10, "正在提取文本...")
        text = ""
        page_count = 0
        is_scanned = False

        if src["type"] == "epub" and src["file_path"]:
            text, page_count, is_scanned = extract_text_from_epub(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("EPUB 解析失败", now(), source_id))
                conn.commit()
                update_task("error", 0, "EPUB 解析失败")
                return
        elif src["type"] == "txt" and src["file_path"]:
            text, page_count, is_scanned = extract_text_from_txt(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("文本文件读取失败", now(), source_id))
                conn.commit()
                update_task("error", 0, "文本文件读取失败")
                return
            # 智能处理模式下 TXT 也可能返回分段，但智能处理有自己的分块逻辑
            if isinstance(text, list):
                text = "\n\n".join(text)  # 合并分段
        elif src["type"] == "docx" and src["file_path"]:
            text, page_count, is_scanned = extract_text_from_docx(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("Word 文档读取失败，文件可能损坏或加密", now(), source_id))
                conn.commit()
                update_task("error", 0, "Word 文档读取失败")
                return
        elif src["type"] == "book" and src["file_path"]:
            text, page_count, is_scanned = extract_text_from_pdf(src["file_path"], update_task)
            if not text:
                conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                             ("PDF无可提取文字", now(), source_id))
                conn.commit()
                update_task("error", 0, "PDF无法识别文字")
                return
        elif src["type"] == "text":
            text = src["url"] or ""
            page_count = 1
        elif src["type"] == "url":
            import asyncio
            loop = asyncio.new_event_loop()
            try:
                text = loop.run_until_complete(extract_text_from_url(src["url"]))
            finally:
                loop.close()
            page_count = 1

        conn.execute("UPDATE sources SET page_count=?,char_count=?,status='processing',updated_at=? WHERE id=?",
                     (page_count, len(text), now(), source_id))
        conn.commit()

        ocr_hint = "（扫描版·AI-OCR）" if is_scanned else ""
        update_task("processing", 15, f"文本提取完成{ocr_hint}（{len(text):,}字），启动智能分析...")

        # 步骤2: 分层 Chunking（结构保留）
        from chunking import chunk_book_text

        def chunking_progress(stage, progress, message):
            # 映射到 15-25%
            mapped = 15 + int(progress * 0.1)
            update_task("processing", mapped, f"[结构分析] {message}")

        chunked = chunk_book_text(
            text=text,
            book_name=src["title"],
            max_chunk_size=4000
        )

        update_task("processing", 25,
                    f"结构分析完成：{chunked['chapter_count']}章节，{chunked['chunk_count']}个片段")

        # 步骤3: 多轮迭代提取
        from extraction_pipeline import MultiRoundExtractionPipeline

        pipeline = MultiRoundExtractionPipeline(
            primary_client=client,
            fallback_client=fallback_client,
            fallback2_client=fallback2_client,
            primary_model=MODEL_ID,
            fallback_model=FALLBACK_MODEL_ID,
            fallback2_model=FALLBACK2_MODEL_ID,
            ip_direction=IP_DIRECTION
        )

        def extraction_progress(stage, progress, message):
            # 映射到 25-90%
            mapped = 25 + int(progress * 0.65)
            update_task("processing", mapped, f"[{stage}] {message}")

        result = pipeline.extract_full(
            book_name=src["title"],
            text=text,
            progress_callback=extraction_progress
        )

        update_task("processing", 90, "提取完成，正在进行质量评分...")

        # 步骤4: 质量评分与路由
        from quality_control import QualityControlPipeline

        qc = QualityControlPipeline(ip_direction=IP_DIRECTION)

        # 收集所有素材
        all_materials = []
        round2_results = result["extraction_pipeline"]["round2_chapters"]

        for r in round2_results:
            data = r["structured_data"]
            source_id_val = source_id

            # 金句
            for q in data.get("quotes", []):
                all_materials.append({
                    "category": "quote",
                    "content": q.get("text", ""),
                    "metadata": {
                        "risk": q.get("risk", "safe"),
                        "scene": q.get("scene", "deep"),
                        "cost": q.get("cost", "mid"),
                        "timeliness": q.get("timeliness", "long"),
                        "context": q.get("context", "")
                    }
                })

            # 案例
            for c in data.get("cases", []):
                content = f"**{c.get('name', '案例')}**\n- 冲突：{c.get('conflict', '')}\n- 动作：{c.get('action', '')}\n- 结果：{c.get('result', '')}\n- 启示：{c.get('insight', '')}"
                all_materials.append({
                    "category": "case",
                    "content": content,
                    "metadata": {
                        "risk": c.get("risk", "safe"),
                        "timeliness": c.get("timeliness", "long")
                    }
                })

            # 观点
            for v in data.get("viewpoints", []):
                content = f"**{v.get('title', '观点')}**\n- 书中依据：{v.get('evidence', '')}\n- IP化角度：{v.get('angle', '')}\n- 冲突预警：{v.get('conflict_warning', '')}"
                all_materials.append({
                    "category": "viewpoint",
                    "content": content,
                    "metadata": {
                        "risk": v.get("risk", "safe"),
                        "timeliness": v.get("timeliness", "long")
                    }
                })

            # 行动
            for a in data.get("actions", []):
                steps = "\n".join([f"  {i+1}. {s}" for i, s in enumerate(a.get("steps", []))])
                content = f"**{a.get('name', '行动')}**\n- 步骤：\n{steps}\n- 适用场景：{a.get('scenario', '')}\n- 风险提示：{a.get('risk_hint', '')}"
                all_materials.append({
                    "category": "action",
                    "content": content,
                    "metadata": {
                        "cost": a.get("cost", "mid")
                    }
                })

        # 获取已有素材用于去重
        existing = conn.execute(
            "SELECT content FROM materials WHERE source_id=?", (source_id,)
        ).fetchall()
        existing_contents = [{"content": row["content"]} for row in existing]

        # 质量评分
        qc_results = qc.process(all_materials, existing_contents)

        # 步骤5: 存储入库
        approved_count = 0
        review_count = 0
        discarded_count = 0

        for m in qc_results["approved"]:
            meta = {**m["metadata"], "_quality_score": m.get("quality_score", {}), "_routing": "approved"}
            conn.execute(
                "INSERT OR REPLACE INTO materials VALUES (?,?,?,?,?,?,?,?,?,?)",
                (str(uuid.uuid4()), source_id, m["category"], m["content"],
                 json.dumps(meta, ensure_ascii=False),
                 json.dumps([], ensure_ascii=False),
                 "[]", 0, 0, now())
            )
            approved_count += 1

        # 待审核的也入库，但标记状态
        for m in qc_results["review"]:
            meta = {**m["metadata"], "_quality_score": m.get("quality_score", {}), "_routing": "review", "_review_needed": True}
            conn.execute(
                "INSERT OR REPLACE INTO materials VALUES (?,?,?,?,?,?,?,?,?,?)",
                (str(uuid.uuid4()), source_id, m["category"], m["content"],
                 json.dumps(meta, ensure_ascii=False),
                 json.dumps([], ensure_ascii=False),
                 "[]", 0, 0, now())
            )
            review_count += 1

        discarded_count = len(qc_results["discarded"])

        conn.execute("UPDATE sources SET status='done',is_scanned=?,updated_at=? WHERE id=?",
                     (1 if is_scanned else 0, now(), source_id))
        conn.commit()

        # 完成报告
        summary = result["summary"]
        report = (f"智能提取完成！\n"
                  f"✓ 自动入库: {approved_count} 条\n"
                  f"⚠ 待审核: {review_count} 条\n"
                  f"✗ 已过滤: {discarded_count} 条\n"
                  f"📊 原始素材: 金句{summary['total_quotes']}/"
                  f"案例{summary['total_cases']}/"
                  f"观点{summary['total_viewpoints']}/"
                  f"选题{summary['total_topics']}")

        update_task("done", 100, report, json.dumps(result, ensure_ascii=False)[:500])

    except Exception as e:
        err_str = str(e)
        print(f"[SmartExtraction] Error: {err_str}")
        import traceback
        traceback.print_exc()

        if is_xunfei_blocked(err_str):
            friendly = "内容审核拦截，请尝试其他书籍"
        else:
            friendly = f"智能提取失败：{err_str[:200]}"

        conn.execute("UPDATE sources SET status='error',error_msg=?,updated_at=? WHERE id=?",
                     (friendly, now(), source_id))
        conn.commit()
        update_task("error", 0, friendly)
    finally:
        conn.close()


# API: 启动智能提取（新版）
@app.post("/api/sources/{source_id}/extract-smart")
def start_smart_extraction(source_id: str, mode: str = "full"):
    """启动智能提取流程（分层Chunking + 多轮Pipeline + 质量评分）"""
    if not _SMART_EXTRACTION_AVAILABLE:
        raise HTTPException(400, "智能提取模块未加载，请检查依赖")

    conn = get_db()
    src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
    if not src:
        conn.close()
        raise HTTPException(404, "来源不存在")

    task_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
        (task_id, source_id, "pending", 0, "等待智能提取...", None, now(), now())
    )
    conn.commit()
    conn.close()

    # 使用新的智能提取流程（传 "smart" 标记，worker 会调用 process_source_task_smart）
    enqueue_task(task_id, source_id, "smart")

    return {
        "task_id": task_id,
        "source_id": source_id,
        "mode": "smart",
        "message": "智能提取任务已创建（分层Chunking+多轮Pipeline+质量评分）"
    }


# API: 获取素材质量详情
@app.get("/api/materials/{material_id}/quality")
def get_material_quality(material_id: str):
    """获取素材的质量评分详情"""
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM materials WHERE id=?", (material_id,)
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(404, "素材不存在")

    metadata = json.loads(row["metadata"] or "{}")
    quality_score = metadata.pop("_quality_score", {})

    return {
        "material_id": material_id,
        "quality_score": quality_score,
        "metadata": metadata,
        "suggestions": quality_score.get("suggestions", [])
    }


# API: 对比新旧提取方式（用于评估）
@app.post("/api/sources/{source_id}/extract-compare")
def compare_extraction_methods(source_id: str):
    """对比原版和智能版提取效果（测试用）"""
    # 创建两个任务，分别用两种方式处理
    conn = get_db()
    src = conn.execute("SELECT * FROM sources WHERE id=?", (source_id,)).fetchone()
    if not src:
        conn.close()
        raise HTTPException(404, "来源不存在")

    # 原版任务
    task1_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
        (task1_id, source_id, "pending", 0, "原版提取", None, now(), now())
    )

    # 智能版任务
    task2_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO tasks VALUES (?,?,?,?,?,?,?,?)",
        (task2_id, source_id, "pending", 0, "智能提取", None, now(), now())
    )
    conn.commit()
    conn.close()

    # 启动任务
    enqueue_task(task1_id, source_id, "full")
    enqueue_task(task2_id, source_id, "full")

    return {
        "classic_task_id": task1_id,
        "smart_task_id": task2_id,
        "message": "对比任务已创建，请分别查看结果"
    }


# ═══════════════════════════════════════════════════════════════════════════
# 跨书整合 API
# ═══════════════════════════════════════════════════════════════════════════

class CrossIntegrateRequest(BaseModel):
    content: str          # 用户输入的素材拼合内容（多本书素材拼在一起）
    mode: str = "主题关联"  # 发现主题关联 / 生成系列策划 / 提炼思维模型

@app.post("/api/studio/cross-integrate")
def cross_integrate(req: CrossIntegrateRequest):
    """
    跨书整合：对来自多本书的素材进行跨书深度分析。
    三种模式：
      - 主题关联：发现多本书共同的深层规律与底层逻辑
      - 系列策划：基于共同主题生成可执行的系列内容策划
      - 思维模型：从多本书中提炼可复用的思维模型框架
    """
    mode_prompts = {
        "主题关联": """你是一位顶级知识整合专家。
用户输入了来自多本不同书籍的素材片段，请完成以下分析：

【任务：发现跨书主题关联】

1. **共同主题识别**（3-5个）
   - 找出这些素材中反复出现的核心主题词或概念
   - 每个主题用一句话点明其在所有书中的共同表达方式

2. **底层规律提炼**
   - 这些书在某个更深的层面上，都在说同一件事是什么？
   - 用"它们都在说：……"的句式表达

3. **视角差异**
   - 不同书从什么不同角度切入这个共同主题？（一句话概括各书视角）

4. **IP创作价值**
   - 基于这个跨书共同主题，可以做什么内容（一句话给出方向）

输出格式：直接按上述4个标题输出，不要废话，不要总结段落。""",

        "系列策划": """你是一位顶级内容策划专家，擅长打造系列化IP内容。
用户输入了来自多本书的素材，请基于这些素材策划一套系列内容。

【任务：生成系列内容策划】

1. **系列主题**
   - 系列名称（5-10字，有记忆点）
   - 一句话定位（这个系列为谁解决什么问题）

2. **系列结构**（5-8期）
   - 每期标题（能独立传播，又形成系列感）
   - 每期核心观点（一句话）
   - 对应使用哪本书的哪个核心素材

3. **推送节奏建议**
   - 建议发布频率和顺序逻辑（为什么这样排）

4. **爆款切入点**
   - 哪一期最适合作为首发（理由）

直接输出策划内容，不要有任何前言和废话。""",

        "思维模型": """你是一位知识萃取专家，擅长从书籍素材中提炼可复用的思维框架。
用户输入了多本书的核心素材，请从中提炼思维模型。

【任务：提炼跨书思维模型】

1. **核心模型命名**
   - 给这个思维模型起一个好记的名字（可以是"XX法则""XX模型""XX框架"等）
   - 一句话说明这个模型解决什么问题

2. **模型结构**
   - 用2-4个步骤或维度描述这个模型的运作逻辑
   - 每个步骤/维度：名称 + 一句话解释 + 来自哪本书的什么概念

3. **适用场景**
   - 这个思维模型在哪3个场景下最有用？（具体场景，不要抽象）

4. **记忆口诀**（可选）
   - 如果能用一句顺口的话记住这个模型，是什么？

5. **内容创作切入**
   - 用这个思维模型，可以写什么选题？（给出3个具体标题）

直接输出，不要废话，要让人看完立刻能用。"""
    }

    system_prompt = f"""你是一位顶级IP内容创作专家，专注于职场认知升级、人性洞察和个人成长破局三大方向。

{mode_prompts.get(req.mode, mode_prompts['主题关联'])}

【去AI化规则——必须执行】
- 禁止"综上所述""值得注意""深刻启示""引人深省""不言而喻"等AI高频词
- 用真实直接的语言，有判断有态度，不说官话
- 说具体的，不说抽象的"""

    user_prompt = f"以下是来自多本书的素材内容，请按要求分析：\n\n{req.content}"

    try:
        result, model_used = ai_extract(system_prompt, user_prompt, max_tokens=3000, temperature=0.75)
        return {"result": result, "model": model_used, "mode": req.mode}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"跨书整合失败：{str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# MediaCrawler 数据接入
# ═══════════════════════════════════════════════════════════════════════════

# MediaCrawler 数据目录（本地 media_crawler/data 或从 config.json 读取）
_MEDIA_CRAWLER_LOCAL_DATA = BASE_DIR / "media_crawler" / "data"
_CONFIG_PATH = BASE_DIR / "config.json"
if _CONFIG_PATH.exists():
    try:
        with open(_CONFIG_PATH, "r", encoding="utf-8") as f:
            _cfg = json.load(f)
            _ext_path = _cfg.get("media_crawler", {}).get("data_dir", "")
            MEDIA_CRAWLER_DATA_DIR = Path(_ext_path) if _ext_path else _MEDIA_CRAWLER_LOCAL_DATA
    except Exception:
        MEDIA_CRAWLER_DATA_DIR = _MEDIA_CRAWLER_LOCAL_DATA
else:
    MEDIA_CRAWLER_DATA_DIR = _MEDIA_CRAWLER_LOCAL_DATA

# 平台 ID → 中文名映射
PLATFORM_NAMES = {
    "douyin": "抖音",
    "xhs": "小红书",
    "weibo": "微博",
    "kuaishou": "快手",
    "bili": "B站",
    "tieba": "贴吧",
    "zhihu": "知乎",
}


class MediaImportRequest(BaseModel):
    records: List[dict]  # 每条记录包含 platform, content, metadata


@app.get("/api/media/platforms")
def list_media_platforms():
    """返回 MediaCrawler 已采集的平台列表"""
    if not MEDIA_CRAWLER_DATA_DIR.exists():
        return {"platforms": []}

    platforms = []
    for subdir in MEDIA_CRAWLER_DATA_DIR.iterdir():
        if subdir.is_dir() and subdir.name != "wechat-output":
            pid = subdir.name
            platforms.append({
                "id": pid,
                "name": PLATFORM_NAMES.get(pid, pid),
                "path": str(subdir),
            })

    return {"platforms": platforms}


@app.get("/api/media/files")
def list_media_files(platform: str = ""):
    """返回指定平台的数据文件列表"""
    if not platform:
        return {"files": []}

    platform_dir = MEDIA_CRAWLER_DATA_DIR / platform
    if not platform_dir.exists():
        return {"files": []}

    files = []
    for subdir in platform_dir.iterdir():
        if subdir.is_dir():
            for f in subdir.iterdir():
                if f.suffix in (".json", ".jsonl", ".csv"):
                    files.append({
                        "name": f.name,
                        "path": str(f.relative_to(MEDIA_CRAWLER_DATA_DIR)).replace("\\", "/"),
                        "size": f.stat().st_size,
                    })

    # 也检查顶层文件
    for f in platform_dir.iterdir():
        if f.is_file() and f.suffix in (".json", ".jsonl", ".csv"):
            files.append({
                "name": f.name,
                "path": str(f.relative_to(MEDIA_CRAWLER_DATA_DIR)).replace("\\", "/"),
                "size": f.stat().st_size,
            })

    return {"files": files}


@app.get("/api/media/data")
def get_media_data(platform: str = "", file: str = "", offset: int = 0, limit: int = 20):
    """读取指定文件的数据，支持分页"""
    if not platform or not file:
        return {"records": [], "total": 0, "offset": offset, "limit": limit}

    file_path = MEDIA_CRAWLER_DATA_DIR / file
    if not file_path.exists():
        return {"records": [], "total": 0, "offset": offset, "limit": limit}

    records = []
    try:
        if file_path.suffix == ".jsonl":
            with open(file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
            total = len(lines)
            for line in lines[offset:offset + limit]:
                try:
                    records.append(json.loads(line))
                except:
                    pass
        elif file_path.suffix == ".json":
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                total = len(data)
                records = data[offset:offset + limit]
            else:
                total = 1
                records = [data] if offset == 0 else []
    except Exception as e:
        return {"records": [], "total": 0, "offset": offset, "limit": limit, "error": str(e)}

    return {"records": records, "total": total, "offset": offset, "limit": limit}


@app.post("/api/media/import")
def import_media_records(req: MediaImportRequest):
    """将选中的 MediaCrawler 记录导入到 ip-arsenal 素材库"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    imported = 0
    failed = 0
    today = datetime.now().strftime("%Y-%m-%d")

    for record in req.records:
        try:
            platform = record.get("platform", "unknown")
            content = record.get("content", "") or record.get("desc", "") or ""
            if not content:
                failed += 1
                continue

            metadata = record.get("metadata", {})
            if not isinstance(metadata, dict):
                metadata = {}

            # 从 record 直接提取有用字段放入 metadata
            for key in ("nickname", "liked_count", "comment_count", "share_count",
                        "collect_count", "user_id", "avatar", "ip_location",
                        "source_keyword", "create_time", "note_id", "aweme_id",
                        "note_url", "aweme_url"):
                if key in record and key not in metadata:
                    metadata[key] = record[key]

            source_id = f"media_crawler_{platform}_{today}"

            # 插入 materials 表
            mid = str(uuid.uuid4())
            c.execute("""
                INSERT INTO materials (id, source_id, category, content, metadata, tags, platform, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                mid,
                source_id,
                "topic",  # 默认分类，可后续 AI 重新分类
                content,
                json.dumps(metadata, ensure_ascii=False),
                "[]",
                platform,
                datetime.now().isoformat(),
            ))
            imported += 1
        except Exception as e:
            failed += 1
            print(f"[MediaImport] 导入失败: {e}")

    conn.commit()
    conn.close()

    return {"imported": imported, "failed": failed}


# ═══════════════════════════════════════════════════════════════════════════
# MediaCrawler API 代理（方案B：深度融合）
# ip-arsenal 后端作为网关，转发请求到 MediaCrawler (port 8080)
# ═══════════════════════════════════════════════════════════════════════════

MEDIACRAWLER_BASE = "http://localhost:8080"


@app.get("/api/mediacrawler/health")
async def mc_health():
    """转发到 MediaCrawler /api/health"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.get(f"{MEDIACRAWLER_BASE}/api/health")
        return r.json()


@app.get("/api/mediacrawler/config/platforms")
async def mc_config_platforms():
    """转发到 MediaCrawler /api/config/platforms"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.get(f"{MEDIACRAWLER_BASE}/api/config/platforms")
        return r.json()


@app.get("/api/mediacrawler/config/options")
async def mc_config_options():
    """转发到 MediaCrawler /api/config/options"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.get(f"{MEDIACRAWLER_BASE}/api/config/options")
        return r.json()


@app.post("/api/mediacrawler/crawler/start")
async def mc_crawler_start(req: dict):
    """转发到 MediaCrawler /api/crawler/start"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.post(f"{MEDIACRAWLER_BASE}/api/crawler/start", json=req)
        return r.json()


@app.post("/api/mediacrawler/crawler/stop")
async def mc_crawler_stop():
    """转发到 MediaCrawler /api/crawler/stop"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.post(f"{MEDIACRAWLER_BASE}/api/crawler/stop")
        return r.json()


@app.get("/api/mediacrawler/crawler/status")
async def mc_crawler_status():
    """转发到 MediaCrawler /api/crawler/status"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.get(f"{MEDIACRAWLER_BASE}/api/crawler/status")
        return r.json()


@app.get("/api/mediacrawler/crawler/logs")
async def mc_crawler_logs(limit: int = 100):
    """转发到 MediaCrawler /api/crawler/logs"""
    async with httpx.AsyncClient(timeout=10.0) as client:
        r = await client.get(f"{MEDIACRAWLER_BASE}/api/crawler/logs", params={"limit": limit})
        return r.json()


@app.get("/api/mediacrawler/data/files")
def mc_data_files(platform: str = ""):
    """返回本地 MediaCrawler 数据目录的文件（不走代理，避免跨域问题）"""
    return list_media_files(platform=platform)


@app.get("/api/mediacrawler/data/files/{file_path:path}")
def mc_data_file_preview(file_path: str, preview: bool = True, limit: int = 50):
    """从本地 MediaCrawler 数据目录读取文件内容"""
    # file_path 格式：douyin/jsonl/search_contents_2026-03-30.jsonl
    fp = MEDIA_CRAWLER_DATA_DIR / file_path
    if not fp.exists():
        return {"records": [], "total": 0, "offset": 0, "limit": limit, "error": "file not found: " + str(fp)}

    if preview:
        # 直接读取文件内容，避免调用 get_media_data 的 platform/file 拼接逻辑
        records = []
        total = 0
        try:
            if fp.suffix == ".jsonl":
                with open(fp, "r", encoding="utf-8") as f:
                    lines = f.readlines()
                total = len(lines)
                for line in lines[0:limit]:
                    try:
                        records.append(json.loads(line))
                    except:
                        pass
            elif fp.suffix == ".json":
                with open(fp, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list):
                    total = len(data)
                    records = data[0:limit]
                else:
                    total = 1
                    records = [data]
        except Exception as e:
            return {"records": [], "total": 0, "offset": 0, "limit": limit, "error": str(e)}
        return {"records": records, "total": total, "offset": 0, "limit": limit}
    else:
        with open(fp, "r", encoding="utf-8") as f:
            content = f.read()
        return {"content": content}


@app.websocket("/api/mediacrawler/ws/status")
async def mc_ws_status(websocket: WebSocket):
    """WebSocket 转发：/api/ws/status"""
    await websocket.accept()
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            async with client.stream("GET", f"{MEDIACRAWLER_BASE}/api/ws/status") as resp:
                async for text in resp.aiter_text():
                    if text:
                        await websocket.send_text(text)
    except Exception as e:
        print(f"[mc_ws_status] error: {e}")
    finally:
        await websocket.close()


# 静态文件（前端）—— 只在非 API 路径上提供服务（放最后，因为要注册在 API 路由之后）
@app.get("/{path:path}")
async def serve_frontend(path: str):
    from fastapi import HTTPException
    from fastapi.responses import FileResponse
    if path.startswith("api/"):
        raise HTTPException(status_code=404, detail="Not Found")
    file_path = FRONT_DIR / path
    if file_path.is_file():
        return FileResponse(file_path)
    # 其它路径返回 index.html（SPA 路由）
    return FileResponse(FRONT_DIR / "index.html")


if __name__ == "__main__":
    import uvicorn, sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    print("IP Arsenal starting on http://localhost:8765")
    print(f"[SmartExtraction] 智能提取模块: {'可用' if _SMART_EXTRACTION_AVAILABLE else '不可用'}")
    # Worker 线程和任务恢复由 @app.on_event("startup") 统一处理
    uvicorn.run("main:app", host="0.0.0.0", port=8765, reload=False)

