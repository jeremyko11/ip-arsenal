# -*- coding: utf-8 -*-
"""
IP Arsenal - 文本提取模块
PDF/EPUB/TXT/DOCX/URL 文本提取、OCR 降级策略
"""
import os
import re
import base64
import threading
import tempfile
import time
from pathlib import Path
from typing import List, Optional

import fitz  # PyMuPDF
from html.parser import HTMLParser
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
import httpx
from bs4 import BeautifulSoup

from config import (
    ODL_AVAILABLE, get_paddle_ocr, MAX_CHARS, MAX_PROMPT_TEXT_CHARS,
    IP_DIRECTION, client, MODEL_ID,
)


# ── PDF 扫描版检测 ────────────────────────────────────────────────────
def _is_scanned_pdf(doc: fitz.Document, sample_pages: int = 10) -> bool:
    """检测PDF是否需要OCR（扫描版 或 低质量混合版）

    借鉴 opendataloader-pdf 的混合路由策略：
    - 扫描版：文字层基本为空 + 有大面积图像
    - 混合版（低质量）：文字层存在但密度低、乱码比例高、字符稀疏
    """
    total_pages = len(doc)
    if total_pages == 0:
        return False

    start_page = min(1, total_pages - 1)
    check_pages = min(sample_pages, total_pages)
    step = max(1, (total_pages - start_page) // check_pages)
    indices = list(range(start_page, total_pages, step))[:check_pages]
    if not indices:
        indices = list(range(total_pages))

    total_chars = 0
    image_heavy_pages = 0
    garbled_chars = 0
    total_valid_chars = 0

    for i in indices:
        page = doc[i]
        text = page.get_text().strip()
        total_chars += len(text)

        if len(text) < 50:
            img_list = page.get_images(full=False)
            if img_list:
                image_heavy_pages += 1

        for ch in text:
            code = ord(ch)
            if (0x4E00 <= code <= 0x9FFF) or ch.isalnum():
                total_valid_chars += 1
            elif code > 127 and not (0x4E00 <= code <= 0x9FFF):
                garbled_chars += 1

    avg = total_chars / len(indices) if indices else 0
    image_ratio = image_heavy_pages / len(indices) if indices else 0

    if avg < 30 or image_ratio > 0.6:
        return True

    if 30 <= avg <= 150 and image_ratio > 0.4:
        return True

    if total_valid_chars + garbled_chars > 0:
        valid_ratio = total_valid_chars / (total_valid_chars + garbled_chars + 1)
        if valid_ratio < 0.3 and avg < 200:
            return True

    return False


def _page_to_numpy(page: fitz.Page, dpi: int = 300, enhance: bool = True) -> np.ndarray:
    """将PDF页面渲染为numpy数组（供PaddleOCR使用）"""
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    if not enhance:
        return np.array(img)

    try:
        img_arr_check = np.array(img.convert("L"))
        mean_brightness = img_arr_check.mean()

        if mean_brightness < 180:
            brightness_factor = min(1.3, 220 / (mean_brightness + 1))
            img = ImageEnhance.Brightness(img).enhance(brightness_factor)
        elif mean_brightness > 240:
            img = ImageEnhance.Brightness(img).enhance(0.95)

        img = ImageEnhance.Contrast(img).enhance(1.5)
        img = img.filter(ImageFilter.UnsharpMask(radius=1.2, percent=130, threshold=2))
        img = img.filter(ImageFilter.MedianFilter(size=3))

        r, g, b = img.split()
        r_mean = np.array(r).mean()
        g_mean = np.array(g).mean()
        b_mean = np.array(b).mean()
        channel_diff = max(r_mean, g_mean, b_mean) - min(r_mean, g_mean, b_mean)
        if channel_diff > 15:
            target_mean = (r_mean + g_mean + b_mean) / 3
            r = ImageEnhance.Brightness(r).enhance(target_mean / (r_mean + 0.01))
            g = ImageEnhance.Brightness(g).enhance(target_mean / (g_mean + 0.01))
            b = ImageEnhance.Brightness(b).enhance(target_mean / (b_mean + 0.01))
            img = Image.merge("RGB", (r, g, b))
    except Exception:
        pass

    return np.array(img)


# ── PaddleOCR 线程本地实例 ──────────────────────────────────────────
_paddle_ocr_local = threading.local()


def _get_thread_paddle_ocr():
    """获取当前线程专属的 PaddleOCR 实例（懒初始化，避免跨线程共享）"""
    if not getattr(_paddle_ocr_local, "ocr", None):
        try:
            os.environ.setdefault("FLAGS_use_mkldnn", "0")
            os.environ.setdefault("PADDLE_DISABLE_ONEDNN", "1")
            from paddleocr import PaddleOCR
            import logging
            logging.getLogger("ppocr").setLevel(logging.ERROR)
            _paddle_ocr_local.ocr = PaddleOCR(lang="ch", use_angle_cls=True)
        except Exception as e:
            print(f"Thread PaddleOCR init failed: {e}")
            _paddle_ocr_local.ocr = False
    return _paddle_ocr_local.ocr if _paddle_ocr_local.ocr is not False else None


def _ocr_single_page_paddle(args):
    """单页 OCR 工作函数（在线程池中执行），返回 (page_num, text_or_none)"""
    page_num, img_arr = args
    ocr = _get_thread_paddle_ocr()
    if ocr is None:
        return page_num, None
    try:
        result = ocr.predict(img_arr)
        page_lines = []
        for res in (result or []):
            if isinstance(res, dict):
                texts = res.get("rec_texts") or []
                scores = res.get("rec_scores") or []
                for i, t in enumerate(texts):
                    if not t or len(t.strip()) <= 1:
                        continue
                    score = scores[i] if i < len(scores) else 1.0
                    if score >= 0.65:
                        page_lines.append(t.strip())
            elif isinstance(res, list):
                for item in res:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        txt_info = item[1]
                        if isinstance(txt_info, (list, tuple)) and len(txt_info) >= 2:
                            txt, score = txt_info[0], txt_info[1]
                            if txt and len(txt.strip()) > 1 and score >= 0.65:
                                page_lines.append(txt.strip())
        return page_num, "\n".join(page_lines) if page_lines else None
    except Exception as e:
        print(f"PaddleOCR page {page_num} error: {e}")
        return page_num, None


# ── AI Studio 云端 OCR ──────────────────────────────────────────────
AISTUDIO_TOKEN = ""  # 可在环境变量或配置中设置


def _ocr_pdf_with_aistudio(pdf_path: str, page_indices: list, task_updater=None) -> str:
    """百度 AI Studio PaddleOCR 云端 API"""
    AISTUDIO_BASE = "https://paddleocr.aistudio-app.com"
    SUBMIT_URL = f"{AISTUDIO_BASE}/api/v2/ocr/jobs"

    token = AISTUDIO_TOKEN
    if not token:
        raise ValueError("AISTUDIO_TOKEN not configured, skipping AI Studio OCR")

    headers = {"Authorization": f"token {token}"}
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    all_pages = list(range(total_pages))

    if set(page_indices) == set(all_pages):
        upload_path = pdf_path
        temp_pdf = None
    else:
        temp_doc = fitz.open()
        for pi in sorted(page_indices):
            temp_doc.insert_pdf(doc, from_page=pi, to_page=pi)
        tmp_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        temp_pdf = tmp_file.name
        tmp_file.close()
        temp_doc.save(temp_pdf)
        temp_doc.close()
        upload_path = temp_pdf
    doc.close()

    try:
        if task_updater:
            task_updater("processing", 28, "上传到百度AI Studio PaddleOCR云端识别（高精度）...")

        file_size_mb = Path(upload_path).stat().st_size / (1024 * 1024)
        if file_size_mb > 48:
            raise ValueError(f"PDF too large for AI Studio ({file_size_mb:.1f}MB > 48MB)")

        with open(upload_path, "rb") as f:
            files_data = {"file": (Path(upload_path).name, f, "application/pdf")}
            resp = httpx.post(
                SUBMIT_URL,
                data={"model": "PP-StructureV3",
                      "optionalPayload": '{"useDocOrientationClassify":true,"useDocUnwarping":true}'},
                files=files_data,
                headers=headers,
                timeout=60,
            )

        if resp.status_code != 200:
            raise ValueError(f"AI Studio submit failed: {resp.status_code} {resp.text[:200]}")

        job_resp = resp.json()
        job_id = job_resp.get("jobId") or (job_resp.get("data") or {}).get("jobId")
        if not job_id:
            raise ValueError(f"No jobId in response: {job_resp}")

        if task_updater:
            task_updater("processing", 35, f"AI Studio任务已提交（jobId={job_id[:8]}...），等待识别完成...")

        result_url = f"{AISTUDIO_BASE}/api/v2/ocr/jobs/{job_id}"
        max_wait = 120
        poll_interval = 5
        elapsed = 0
        while elapsed < max_wait:
            time.sleep(poll_interval)
            elapsed += poll_interval
            r = httpx.get(result_url, headers=headers, timeout=30)
            if r.status_code != 200:
                continue
            data = r.json()
            status = (data.get("status") or "").lower()
            if status in ("success", "finished", "completed", "done"):
                pages = data.get("result", {}).get("pages") or data.get("pages") or []
                if not pages and "markdownUrl" in (data.get("result") or {}):
                    md_url = data["result"]["markdownUrl"]
                    md_resp = httpx.get(md_url, headers=headers, timeout=30)
                    return md_resp.text
                texts = []
                for pg in pages:
                    pg_text = pg.get("text") or ""
                    if not pg_text:
                        for layout in (pg.get("layouts") or []):
                            pg_text += layout.get("text", "") + "\n"
                    texts.append(pg_text.strip())
                return "\n\n".join(t for t in texts if t)
            elif status in ("failed", "error"):
                raise ValueError(f"AI Studio job failed: {data}")
            if task_updater:
                task_updater("processing", 35 + min(int(elapsed / max_wait * 30), 30),
                             f"AI Studio识别中...（已等待{elapsed}秒）")

        raise ValueError(f"AI Studio timed out after {max_wait}s")

    finally:
        if temp_pdf and Path(temp_pdf).exists():
            Path(temp_pdf).unlink(missing_ok=True)


def _ocr_pages_with_ai_fallback(doc: fitz.Document, page_indices: list, task_updater=None) -> str:
    """降级方案：用讯飞AI视觉接口做OCR（当PaddleOCR不可用时）"""
    import concurrent.futures as _cf

    total = len(page_indices)
    ai_workers = min(4, total)

    rendered = []
    for page_num in page_indices:
        try:
            mat = fitz.Matrix(250 / 72, 250 / 72)
            pix = doc[page_num].get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            rendered.append((page_num, b64))
        except Exception as e:
            rendered.append((page_num, None))
            print(f"AI-OCR render page {page_num} failed: {e}")

    if task_updater:
        task_updater("processing", 30, f"讯飞AI-OCR并行识别{total}页（{ai_workers}线程）...")

    def _call_ai_ocr(item):
        page_num, b64 = item
        if b64 is None:
            return page_num, f"[第{page_num+1}页渲染失败]"
        try:
            resp = client.chat.completions.create(
                model=MODEL_ID,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    {"type": "text", "text": "请将这张图片中的所有文字完整识别出来，保持原有段落和分节结构，直接输出识别到的文字内容，不要添加任何说明。处理完成后输出【识别完成】。"}
                ]}],
                max_tokens=4000, temperature=0,
            )
            return page_num, resp.choices[0].message.content or ""
        except Exception as e:
            return page_num, f"[第{page_num+1}页OCR失败: {e}]"

    results_map = {}
    completed = 0
    with _cf.ThreadPoolExecutor(max_workers=ai_workers, thread_name_prefix="AIOCR") as executor:
        future_map = {executor.submit(_call_ai_ocr, item): item[0] for item in rendered}
        for future in _cf.as_completed(future_map):
            try:
                page_num, text = future.result(timeout=90)
                results_map[page_num] = text
                completed += 1
            except Exception as e:
                pn = future_map[future]
                results_map[pn] = f"[第{pn+1}页OCR超时: {e}]"
                completed += 1
            if task_updater:
                pct = 30 + int((completed / total) * 40)
                task_updater("processing", pct, f"讯飞AI-OCR识别中...（{completed}/{total}页完成）")

    return "\n\n".join(results_map.get(pn, "") for pn in page_indices)


def _try_aistudio_or_fallback(doc, page_indices, pdf_path=None, task_updater=None) -> str:
    """尝试百度 AI Studio 云端 OCR，失败则降到讯飞 AI-OCR"""
    if pdf_path:
        try:
            result = _ocr_pdf_with_aistudio(pdf_path, page_indices, task_updater)
            if result and len(result.strip()) >= 50:
                return result
        except Exception as e:
            print(f"AI Studio OCR failed: {e}, falling back to AI vision OCR")
            if task_updater:
                task_updater("processing", 45, f"AI Studio识别失败（{str(e)[:60]}），切换到讯飞AI-OCR...")
    return _ocr_pages_with_ai_fallback(doc, page_indices, task_updater)


def _ocr_pages_with_paddle(doc: fitz.Document, page_indices: list, task_updater=None,
                           pdf_path: str = None) -> str:
    """用PaddleOCR对指定页面做本地OCR，返回合并文本"""
    if get_paddle_ocr() is None:
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    import concurrent.futures as _cf

    total = len(page_indices)
    ocr_workers = min(8, max(1, total // 3 + 1))

    if task_updater:
        task_updater("processing", 28, f"PaddleOCR并行识别{total}页（{ocr_workers}线程）...")

    rendered = []
    for page_num in page_indices:
        try:
            img_arr = _page_to_numpy(doc[page_num], dpi=300)
            rendered.append((page_num, img_arr))
        except Exception as e:
            print(f"Page render failed {page_num}: {e}")
            rendered.append((page_num, None))

    results_map = {}
    fail_count = 0
    completed = 0

    def _do_ocr(args):
        pn, img = args
        if img is None:
            return pn, None
        return _ocr_single_page_paddle((pn, img))

    with _cf.ThreadPoolExecutor(max_workers=ocr_workers, thread_name_prefix="OCR") as executor:
        future_map = {executor.submit(_do_ocr, item): item[0] for item in rendered}
        for future in _cf.as_completed(future_map):
            try:
                page_num, text = future.result(timeout=60)
                results_map[page_num] = text
                if text:
                    completed += 1
                else:
                    fail_count += 1
            except Exception as e:
                pn = future_map[future]
                print(f"OCR future error page {pn}: {e}")
                results_map[pn] = None
                fail_count += 1
            if task_updater:
                done_cnt = len(results_map)
                pct = 30 + int((done_cnt / total) * 40)
                task_updater("processing", pct, f"PaddleOCR识别中...（{done_cnt}/{total}页完成）")

    ordered_texts = [results_map.get(pn) for pn in page_indices]
    results = [t for t in ordered_texts if t]

    if fail_count > total * 0.5:
        if task_updater:
            task_updater("processing", 28, f"PaddleOCR识别率低（{fail_count}/{total}页失败），切换到AI Studio云端...")
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    combined = "\n\n".join(results)
    if len(combined.strip()) < 100 and total > 2:
        if task_updater:
            task_updater("processing", 28, "PaddleOCR输出为空，切换到AI Studio云端...")
        return _try_aistudio_or_fallback(doc, page_indices, pdf_path, task_updater)

    return combined


def _extract_with_opendataloader(pdf_path: str, max_chars: int = None) -> Optional[str]:
    """用 opendataloader-pdf (Java XY-Cut++ 引擎) 提取PDF全文"""
    if not ODL_AVAILABLE:
        return None

    from config import _odl_convert as odl_convert  # 延迟导入

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            odl_convert(
                input_path=pdf_path,
                output_dir=tmpdir,
                format="text",
                reading_order="xycut",
                keep_line_breaks=True,
                quiet=True,
            )

            txt_files = [f for f in os.listdir(tmpdir) if f.endswith(".txt")]
            if not txt_files:
                return None

            txt_path = os.path.join(tmpdir, txt_files[0])
            with open(txt_path, encoding="utf-8", errors="replace") as fh:
                text = fh.read()

            if not text or len(text.strip()) < 100:
                return None

            if max_chars:
                return text[:max_chars]
            return text

    except Exception as e:
        print(f"[opendataloader] 提取失败: {e}")
        return None


def _extract_text_layout_aware(page) -> str:
    """布局感知文本提取（XY-Cut++ 阅读顺序策略）"""
    try:
        blocks = page.get_text("blocks")
        if not blocks:
            return page.get_text()

        text_blocks = [(b[0], b[1], b[4]) for b in blocks if b[6] == 0 and b[4].strip()]

        if not text_blocks:
            return page.get_text()

        page_rect = page.rect
        page_width = page_rect.width if page_rect.width > 0 else 595

        right_blocks = sum(1 for b in text_blocks if b[0] > page_width * 0.55)
        is_multi_column = right_blocks > len(text_blocks) * 0.2

        if is_multi_column:
            left_col = [(x, y, t) for x, y, t in text_blocks if x < page_width * 0.55]
            right_col = [(x, y, t) for x, y, t in text_blocks if x >= page_width * 0.55]
            left_sorted = sorted(left_col, key=lambda b: b[1])
            right_sorted = sorted(right_col, key=lambda b: b[1])
            sorted_blocks = left_sorted + right_sorted
        else:
            sorted_blocks = sorted(text_blocks, key=lambda b: (round(b[1] / 20) * 20, b[0]))

        return "\n".join(t.strip() for _, _, t in sorted_blocks if t.strip())

    except Exception:
        return page.get_text()


# ── 各类文本提取入口 ─────────────────────────────────────────────────
def extract_text_from_epub(path: str, task_updater=None) -> tuple:
    """从 EPUB 文件直接提取纯文本，不走 OCR。返回 (text, chapter_count, False)"""
    try:
        import zipfile

        class _StripHTML(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []
            def handle_data(self, data):
                self.parts.append(data)
            def get_text(self):
                return " ".join(self.parts)

        def _html_to_text(html_bytes):
            try:
                html = html_bytes.decode("utf-8", errors="replace")
            except Exception:
                html = str(html_bytes)
            p = _StripHTML()
            p.feed(html)
            return p.get_text()

        if task_updater:
            task_updater("processing", 20, "正在解析 EPUB 文件...")

        chapters = []
        with zipfile.ZipFile(path, "r") as zf:
            names = sorted(zf.namelist())
            html_files = [n for n in names if n.lower().endswith((".xhtml", ".html", ".htm", ".xml"))
                          and "ncx" not in n.lower() and "opf" not in n.lower() and "toc" not in n.lower()]
            for n in html_files:
                try:
                    raw = zf.read(n)
                    chapters.append(_html_to_text(raw))
                except Exception:
                    pass

        if not chapters:
            return "", 0, False

        nc = len(chapters)
        text = "\n\n".join(chapters)
        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS]

        if task_updater:
            task_updater("processing", 60, f"EPUB 解析完成（{nc} 章，{len(text):,} 字），正在 AI 提炼...")

        return text, nc, False

    except Exception as e:
        print(f"[EPUB] 提取失败: {e}")
        return "", 0, False


def extract_text_from_txt(path: str, task_updater=None) -> tuple:
    """从 TXT/MD 文件直接读取文本。返回 (text_or_chunks, chunk_count, False)"""
    try:
        if task_updater:
            task_updater("processing", 20, "正在读取文本文件...")

        for enc in ("utf-8", "gbk", "utf-16", "latin-1"):
            try:
                with open(path, encoding=enc, errors="replace") as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            return "", 0, False

        if len(text) > MAX_CHARS:
            chunks = []
            paragraphs = text.split("\n\n")
            current_chunk = ""
            for para in paragraphs:
                if len(current_chunk) + len(para) + 2 > MAX_CHARS:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para
                else:
                    current_chunk = (current_chunk + "\n\n" + para) if current_chunk else para
            if current_chunk.strip():
                chunks.append(current_chunk.strip())

            if task_updater:
                task_updater("processing", 60, f"文件读取完成（{len(chunks)} 个分段，共 {len(text):,} 字），正在 AI 提炼...")
            return chunks, len(chunks), False
        else:
            if task_updater:
                task_updater("processing", 60, f"文件读取完成（{len(text):,} 字），正在 AI 提炼...")
            return text, 1, False

    except Exception as e:
        print(f"[TXT] 提取失败: {e}")
        return "", 0, False


def extract_text_from_docx(path: str, task_updater=None) -> tuple:
    """从 DOCX（Word）文件提取纯文本，不走 OCR"""
    try:
        import docx as _docx_mod
        if task_updater:
            task_updater("processing", 20, "正在读取 Word 文档...")

        doc = _docx_mod.Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            s = para.text.strip()
            if s:
                paragraphs.append(s)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        text = "\n".join(paragraphs)
        text = text[:MAX_CHARS]
        para_count = len(paragraphs)

        if task_updater:
            task_updater("processing", 60, f"Word 文档读取完成（{para_count} 段，{len(text):,} 字），正在 AI 提炼...")
        return text, para_count, False

    except Exception as e:
        print(f"[DOCX] 提取失败: {e}")
        return "", 0, False


def extract_text_from_pdf(path: str, task_updater=None) -> tuple:
    """
    提取PDF文本，三级精度递进策略：
    Level 1 — opendataloader-pdf（Java XY-Cut++ 引擎）
    Level 2 — PaddleOCR（本地GPU加速）
    Level 3 — 讯飞 AI-OCR（视觉大模型）
    返回 (text, page_count, is_scanned)
    """
    doc = fitz.open(path)
    page_count = len(doc)

    # Level 1: opendataloader-pdf
    if ODL_AVAILABLE:
        if task_updater:
            task_updater("processing", 15, f"正在用 opendataloader-pdf 高精度解析（{page_count}页）...")

        odl_text = _extract_with_opendataloader(path, max_chars=MAX_CHARS)

        if odl_text and len(odl_text.strip()) >= 200:
            doc.close()
            print(f"[OCR] Level1 opendataloader 成功，字符数: {len(odl_text)}")
            is_scanned_hint = len(odl_text) < 1000 and page_count > 10
            return odl_text, page_count, is_scanned_hint
        else:
            print(f"[OCR] Level1 opendataloader 输出不足，降级到 PaddleOCR")

    # 检测扫描版
    segments = []
    total_chars = 0
    for page in doc:
        t = _extract_text_layout_aware(page)
        total_chars += len(t)
        segments.append(t)

    is_scanned = _is_scanned_pdf(doc)

    if not is_scanned and total_chars >= 500:
        doc.close()
        n = len(segments)
        front = segments[:int(n * 0.4)]
        mid = segments[int(n * 0.35):int(n * 0.65)]
        back = segments[int(n * 0.7):]
        all_text = "\n".join(front + mid + back)
        return all_text[:MAX_CHARS], page_count, False

    # Level 2: PaddleOCR
    MAX_OCR_PAGES = 60
    n = page_count
    if n <= MAX_OCR_PAGES:
        page_indices = list(range(n))
    else:
        front_n = int(MAX_OCR_PAGES * 0.4)
        mid_n = int(MAX_OCR_PAGES * 0.3)
        back_n = MAX_OCR_PAGES - front_n - mid_n
        front_idx = list(range(front_n))
        mid_idx = list(range(int(n * 0.35), int(n * 0.35) + mid_n))
        back_idx = list(range(n - back_n, n))
        page_indices = sorted(set(front_idx + mid_idx + back_idx))

    if task_updater:
        task_updater("processing", 25,
                     f"检测到扫描版PDF（{page_count}页），启动PaddleOCR识别{len(page_indices)}页...")

    ocr_text = _ocr_pages_with_paddle(doc, page_indices, task_updater, pdf_path=path)
    doc.close()
    return ocr_text[:MAX_CHARS], page_count, True


def _url_login_check(url: str):
    """检查 URL 是否需要登录（同步版本）"""
    LOGIN_REQUIRED_DOMAINS = {
        "weibo.com": "微博",
        "weibo.cn": "微博",
        "m.weibo.cn": "微博",
        "douyin.com": "抖音",
        "xhs.link": "小红书",
        "xiaohongshu.com": "小红书",
        "www.zhihu.com": "知乎（部分内容需登录）",
        "mp.weixin.qq.com": "微信公众号",
        "x.com": "X（Twitter）",
        "twitter.com": "X（Twitter）",
        "www.twitter.com": "X（Twitter）",
    }
    from urllib.parse import urlparse
    parsed = urlparse(url)
    domain = parsed.netloc.lower().lstrip("www.")
    for d, name in LOGIN_REQUIRED_DOMAINS.items():
        if d in domain or domain == d:
            raise ValueError(
                f"【{name}】需要登录才能获取内容，直接导入URL无效。\n"
                f"请使用「文字导入」方式：复制博主内容文本后粘贴到文字输入框。\n"
                f"或使用浏览器爬取工具先抓取内容，再通过文字方式导入。"
            )


def extract_text_from_url(url: str) -> str:
    """同步版本：从URL提取文本（Worker线程安全）"""
    _url_login_check(url)

    try:
        with httpx.Client(
            timeout=15,
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"},
        ) as c:
            resp = c.get(url)
            text = resp.text

            final_url = str(resp.url)
            LOGIN_REDIRECT_HINTS = ["passport.", "login", "sso.", "signin", "auth."]
            if any(hint in final_url for hint in LOGIN_REDIRECT_HINTS):
                raise ValueError(
                    f"访问被重定向到登录页（{final_url[:80]}），无法直接抓取内容。\n"
                    "请复制网页正文后使用「文字导入」方式导入。"
                )

            soup = BeautifulSoup(text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
                tag.decompose()
            extracted = soup.get_text(separator="\n", strip=True)

            if len(extracted) < 500:
                raise ValueError(
                    f"提取内容过短（{len(extracted)}字），网页可能需要登录或JS渲染才能显示内容。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            lines = [l.strip() for l in extracted.split("\n") if l.strip()]
            if len(lines) < 3:
                raise ValueError(
                    f"提取内容无效（仅{len(lines)}行有效文字），网页可能需要JS渲染。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            return extracted[:MAX_CHARS]
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"抓取失败：{e}")


async def extract_text_from_url_async(url: str) -> str:
    """异步版本：从URL提取文本（仅用于 FastAPI 路由中）"""
    _url_login_check(url)

    try:
        async with httpx.AsyncClient(
            timeout=15,
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"},
        ) as c:
            resp = await c.get(url)
            text = resp.text

            final_url = str(resp.url)
            LOGIN_REDIRECT_HINTS = ["passport.", "login", "sso.", "signin", "auth."]
            if any(hint in final_url for hint in LOGIN_REDIRECT_HINTS):
                raise ValueError(
                    f"访问被重定向到登录页（{final_url[:80]}），无法直接抓取内容。\n"
                    "请复制网页正文后使用「文字导入」方式导入。"
                )

            soup = BeautifulSoup(text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
                tag.decompose()
            extracted = soup.get_text(separator="\n", strip=True)

            if len(extracted) < 500:
                raise ValueError(
                    f"提取内容过短（{len(extracted)}字），网页可能需要登录或JS渲染才能显示内容。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            lines = [l.strip() for l in extracted.split("\n") if l.strip()]
            if len(lines) < 3:
                raise ValueError(
                    f"提取内容无效（仅{len(lines)}行有效文字），网页可能需要JS渲染。\n"
                    "建议：复制网页文字后使用「文字导入」方式导入。"
                )

            return extracted[:MAX_CHARS]
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"抓取失败：{e}")
